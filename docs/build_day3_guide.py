"""Run this script to generate day-3-learning-guide.docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Style helpers ──────────────────────────────────────────────────────────────

def set_font(run, bold=False, italic=False, size=None, color=None, font_name=None):
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, bold=True, size=20, color=(44, 62, 80))
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, bold=True, size=15, color=(44, 62, 80))
    return p


def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, bold=True, size=12, color=(52, 73, 94))
    return p


def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def code_block(lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 30, 30)
        if i < len(lines) - 1:
            run.add_break()
    doc.add_paragraph()


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def callout(label, text, bg="FFF3CD"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg)
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    r1 = p.add_run(f"{label}  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    doc.add_paragraph()


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "2C3E50")
        tcPr.append(shd)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        fill = "F8F9FA" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("─" * 80)
    run.font.color.rgb = RGBColor(189, 195, 199)
    run.font.size = Pt(8)


# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run("AI Executive Reporting System")
set_font(run, bold=True, size=26, color=(44, 62, 80))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Day 3 Learning Guide")
set_font(run2, bold=True, size=20, color=(52, 152, 219))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("KPI Aggregation + Pipeline Orchestrator")
set_font(run3, italic=True, size=13, color=(127, 140, 141))

doc.add_paragraph()
divider()
doc.add_paragraph()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run(
    "This guide explains every SQL query in the KPI aggregator, the pipeline orchestrator pattern,\n"
    "the conftest.py module resolution fix, and all 6 new tests added on Day 3."
)
set_font(run4, size=11, color=(52, 73, 94))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — WHAT WE BUILT
# ══════════════════════════════════════════════════════════════════════════════

h1("1. What Day 3 Built")

body(
    "Days 1 and 2 built three DuckDB tables: bronze (raw), silver (clean), gold (KPIs). "
    "Day 3 adds two more layers: a Python aggregation function that rolls up the gold table "
    "into portfolio-level numbers, and a pipeline orchestrator that runs every stage with a single call."
)

body("The complete pipeline after Day 3:")

code_block([
    "CSV file",
    "  → csv_loader.py         (Day 1)  → bronze_project_raw",
    "  → bronze_to_silver.py   (Day 2)  → silver_project_clean",
    "  → silver_to_gold.py     (Day 2)  → gold_project_kpi",
    "  → kpi_aggregator.py     (Day 3)  → KPISummary dict (Python)",
    "  → pipeline.py           (Day 3)  → orchestrates all of the above",
])

body("Four new files were created:")
bullet("src/transforms/kpi_aggregator.py — rolls up gold table into a Python dict")
bullet("src/pipeline.py — orchestrates all five stages in order")
bullet("tests/test_kpi_aggregator.py — 4 tests for portfolio aggregations")
bullet("tests/test_pipeline.py — 2 end-to-end pipeline tests")
body("One existing file was created at project root: conftest.py — fixes Python module resolution for tests.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — KPI AGGREGATION CONCEPT
# ══════════════════════════════════════════════════════════════════════════════

h1("2. What KPI Aggregation Means")

body(
    "The gold table has one row per project. Fifteen rows. Each row has computed KPI columns like "
    "budget_variance_pct and is_late. But an executive report does not say '15 rows of data'. "
    "It says: '40% of projects are delayed, average overrun is 17%, top risk is Mobile App Launch.'"
)

body(
    "KPI aggregation is the step that collapses 15 rows into one Python dictionary of summary statistics. "
    "That dictionary is what the AI model will narrate — it never sees the raw rows."
)

add_table(
    ["Input (gold table, 15 rows)", "Output (KPI summary dict, 1 Python object)"],
    [
        ["Row: Mobile App Launch, budget_variance_pct=32.1, is_late=True",  "total_projects: 15"],
        ["Row: ERP Upgrade, budget_variance_pct=-2.8, is_late=False",        "pct_delayed: 40.0"],
        ["Row: Supply Chain Optimization, ...",                               "top_3_risks: [Mobile App Launch, ...]"],
        ["... 12 more rows ...",                                              "total_budget_variance: $183,000"],
        ["",                                                                  "avg_schedule_delay_days: 52.3"],
    ],
    col_widths=[3.2, 3.4],
)

callout(
    "The key insight:",
    "The AI model never sees raw project rows. It only sees the aggregated summary — "
    "numbers and lists that were computed from real data. This is the anti-hallucination boundary. "
    "The model cannot invent a project name that is not in the top_3_risks list, "
    "because that list was built from your actual data.",
    bg="E8F5E9",
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — kpi_aggregator.py QUERY BY QUERY
# ══════════════════════════════════════════════════════════════════════════════

h1("3. File: src/transforms/kpi_aggregator.py")

body(
    "This function executes eight SQL queries against the gold table and assembles the results "
    "into a Python dictionary. Let's walk through every query."
)

h2("3.1 Total Project Count")

code_block([
    "total = con.execute('SELECT COUNT(*) FROM gold_project_kpi').fetchone()[0]",
])

body(
    "COUNT(*) counts every row. fetchone() returns a tuple like (15,). "
    "The [0] index extracts just the integer. This is the denominator for all percentage calculations."
)

h2("3.2 Status Percentages — the pct() helper")

code_block([
    "def pct(status_val):",
    "    n = con.execute(",
    "        'SELECT COUNT(*) FROM gold_project_kpi WHERE status = ?', [status_val]",
    "    ).fetchone()[0]",
    "    return round((n / total * 100) if total else 0, 1)",
    "",
    "pct_on_track  = pct('on_track')",
    "pct_at_risk   = pct('at_risk')",
    "pct_delayed   = pct('delayed')",
    "pct_completed = pct('completed')",
])

body(
    "The ? is a parameterized query placeholder — DuckDB replaces it with the value in the list [status_val]. "
    "Never use f-strings to inject values into SQL. Parameterized queries prevent SQL injection and "
    "handle special characters in values automatically."
)
body(
    "The (n / total * 100) if total else 0 guard prevents ZeroDivisionError when the table is empty. "
    "round(..., 1) keeps one decimal place: 40.0%, not 40.000000001%."
)

callout(
    "Why four separate queries instead of one GROUP BY?",
    "A single GROUP BY would return a dynamic number of rows. The KPI dict needs four named keys "
    "(pct_on_track, pct_at_risk, ...) regardless of which statuses exist in the data. "
    "If 'at_risk' has zero projects, GROUP BY omits it — the pct() helper returns 0.0 instead, "
    "which is what the AI prompt builder and UI expect.",
    bg="E3F2FD",
)

h2("3.3 Portfolio Budget Rollup")

code_block([
    "budget_row = con.execute(\"\"\"",
    "    SELECT",
    "        SUM(budget_variance)                                        AS total_variance,",
    "        ROUND(SUM(budget_variance) / SUM(planned_cost) * 100, 2)   AS total_variance_pct",
    "    FROM gold_project_kpi",
    "\"\"\").fetchone()",
    "total_budget_variance     = budget_row[0] or 0",
    "total_budget_variance_pct = budget_row[1] or 0",
])

body(
    "SUM(budget_variance) adds up all dollar overruns and underruns across the portfolio. "
    "A project $60K over budget and a project $10K under budget = $50K net overrun."
)
body(
    "The variance percentage is computed as total variance / total planned spend — not the average "
    "of individual percentages. This is the correct portfolio-level metric. Averaging percentages "
    "would give the wrong answer when projects have very different budgets."
)
code_block([
    "Example with 3 projects:",
    "  Project A: planned $100K, actual $110K → +$10K overrun",
    "  Project B: planned $1M,   actual $950K → -$50K underrun",
    "  Project C: planned $200K, actual $240K → +$40K overrun",
    "",
    "SUM(budget_variance) = 10000 - 50000 + 40000 = 0  (net zero!)",
    "SUM(budget_variance) / SUM(planned_cost) * 100",
    "  = 0 / 1300000 * 100 = 0.0%",
    "",
    "Averaging individual pcts: (+10%, -5%, +20%) / 3 = +8.3%  ← WRONG for portfolio view",
])
body(
    "The 'or 0' fallback handles the case where SUM returns NULL (empty table). "
    "Python's 'None or 0' evaluates to 0."
)

h2("3.4 Average Schedule Delay (Late Projects Only)")

code_block([
    "avg_delay = con.execute(\"\"\"",
    "    SELECT ROUND(AVG(schedule_delay_days), 1)",
    "    FROM gold_project_kpi",
    "    WHERE is_late = TRUE",
    "\"\"\").fetchone()[0] or 0",
])

body(
    "This query only averages delay days for projects where is_late = TRUE. "
    "Including on-time and early projects in the average would dilute the number and "
    "make it meaningless for an executive who wants to know 'how late are our late projects?'"
)
body(
    "A project that finished 10 days early (schedule_delay_days = -10) should not reduce "
    "the average delay shown in a risk report. Filtering to is_late = TRUE excludes those."
)

h2("3.5 Top 3 High-Risk Projects — CASE ORDER")

code_block([
    "top_3_risks = con.execute(\"\"\"",
    "    SELECT project_name, risk_level, budget_variance_pct",
    "    FROM gold_project_kpi",
    "    ORDER BY",
    "        CASE risk_level",
    "            WHEN 'critical' THEN 4",
    "            WHEN 'high'     THEN 3",
    "            WHEN 'medium'   THEN 2",
    "            ELSE 1",
    "        END DESC,",
    "        budget_variance_pct DESC",
    "    LIMIT 3",
    "\"\"\").df().to_dict('records')",
])

body(
    "SQL cannot ORDER BY a text column in a meaningful business order (alphabetical would put "
    "'critical' before 'high' — wrong). The CASE statement converts the text label to an integer "
    "rank, then orders by that integer descending: critical(4) > high(3) > medium(2) > low(1)."
)
body(
    "The secondary sort (budget_variance_pct DESC) breaks ties within the same risk level: "
    "two 'high' risk projects are sorted by who is more over budget."
)
body(
    ".df() converts the DuckDB result to a pandas DataFrame. "
    ".to_dict('records') converts it to a list of dicts — one dict per row. "
    "This is the format the prompt_builder expects."
)
code_block([
    "# .df().to_dict('records') output:",
    "[",
    "  {'project_name': 'Mobile App Launch',          'risk_level': 'critical', 'budget_variance_pct': 32.14},",
    "  {'project_name': 'Financial System Integration','risk_level': 'high',     'budget_variance_pct': 21.43},",
    "  {'project_name': 'Supply Chain Optimization',  'risk_level': 'high',     'budget_variance_pct': 9.71},",
    "]",
])

h2("3.6 FIRST() Aggregation for Department Issues")

code_block([
    "issues_by_dept = con.execute(\"\"\"",
    "    SELECT",
    "        department,",
    "        COUNT(*) AS issue_count,",
    "        FIRST(issue_description) AS sample_issue",
    "    FROM gold_project_kpi",
    "    WHERE issue_description != ''",
    "    GROUP BY department",
    "    ORDER BY issue_count DESC",
    "\"\"\").df().to_dict('records')",
])

body(
    "FIRST() is a DuckDB aggregate function that returns the first non-null value in the group. "
    "Since each department has multiple projects with different issues, we cannot select "
    "issue_description without aggregating it — SQL would not know which row to pick."
)
body(
    "FIRST() says: 'just give me any one issue description from this department as a representative sample'. "
    "The WHERE issue_description != '' excludes projects with no logged issues before grouping, "
    "so the count reflects only projects that actually reported issues."
)

add_table(
    ["Function", "What it does", "When to use it"],
    [
        ["COUNT(*)",     "Count rows in group",                     "Always needed for group totals"],
        ["SUM(col)",     "Add all values in group",                 "Totals (budget, headcount)"],
        ["AVG(col)",     "Average of values in group",              "Means (delay days, cost ratio)"],
        ["MAX/MIN(col)", "Highest / lowest value in group",         "Ranges, extremes"],
        ["FIRST(col)",   "Any one value from the group (DuckDB)",   "Representative sample when exact choice does not matter"],
    ],
    col_widths=[1.3, 2.5, 2.8],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — pipeline.py: THE ORCHESTRATOR PATTERN
# ══════════════════════════════════════════════════════════════════════════════

h1("4. File: src/pipeline.py — The Orchestrator Pattern")

body(
    "An orchestrator is a file that knows the execution order of a system. "
    "It does not implement business logic — it just calls the right functions in the right sequence."
)

h2("4.1 The Architectural Rule")

callout(
    "Critical rule:",
    "src/pipeline.py is the ONLY file that imports across layer boundaries. "
    "csv_loader.py only knows about db.py. bronze_to_silver.py only knows about db.py. "
    "kpi_aggregator.py only knows about db.py. Only pipeline.py imports all of them together.",
    bg="FCE4EC",
)

body(
    "Why does this matter? If every file imported from every other file, you would create a "
    "tangled web of dependencies. A change to kpi_aggregator.py would potentially break "
    "csv_loader.py if they shared imports. By isolating the orchestration in one file:"
)
bullet("Each layer module is independently testable (no hidden dependencies)")
bullet("The execution order is documented in one place, not scattered")
bullet("You can replace any layer (e.g., swap CSV for API) without touching other layers")

h2("4.2 The Five Stages")

code_block([
    "def run_pipeline(csv_path: str) -> dict:",
    "    start = time.time()",
    "",
    "    # Stage 1: Load CSV strings into bronze table",
    "    loaded, skipped = load_csv_to_bronze(csv_path)",
    "",
    "    # Stage 2: Clean and normalize → silver table",
    "    n = bronze_to_silver()",
    "",
    "    # Stage 3: Compute KPI columns → gold table",
    "    n = silver_to_gold()",
    "",
    "    # Stage 4: Roll up portfolio statistics → Python dict",
    "    kpi_summary = compute_kpi_summary()",
    "",
    "    # Stage 5: [placeholder — LLM wired on Day 4]",
    "    ai_report = '[AI report will appear here after Day 4]'",
    "",
    "    elapsed = round(time.time() - start, 2)",
    "    return {'kpi_summary': kpi_summary, 'ai_report': ai_report}",
])

body(
    "Each stage returns something that the logger records. The pipeline returns a dict with two keys: "
    "kpi_summary (the aggregated data) and ai_report (the LLM text — placeholder until Day 4). "
    "The UI and tests always use these keys to access results."
)

h2("4.3 Why loguru Instead of print()")

code_block([
    "from loguru import logger",
    "",
    "logger.info('Stage 1: Ingestion')",
    "logger.info(f'  → {loaded} rows loaded, {skipped} skipped')",
])

body(
    "loguru is a structured logging library. Unlike print(), it adds a timestamp, log level, "
    "and source location to every message automatically."
)
code_block([
    "# print() output:",
    "Stage 1: Ingestion",
    "",
    "# loguru output:",
    "2024-05-07 14:23:01.412 | INFO | src.pipeline:run_pipeline:13 - Stage 1: Ingestion",
    "2024-05-07 14:23:01.445 | INFO | src.pipeline:run_pipeline:15 -   → 15 rows loaded, 0 skipped",
])
body(
    "In a production system, these logs are shipped to a monitoring service (Datadog, CloudWatch). "
    "When a pipeline run fails at 3am, the logs tell you exactly which stage failed and why. "
    "For this demo, they make the pipeline feel professional and give instant feedback in the terminal."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — conftest.py AND PYTHONPATH
# ══════════════════════════════════════════════════════════════════════════════

h1("5. conftest.py and the PYTHONPATH Fix")

body(
    "Day 3 introduced a module resolution error that blocked both pytest and direct script execution. "
    "Understanding this fix teaches you how Python finds modules."
)

h2("5.1 The Error")

code_block([
    "$ PYTHONPATH=. python src/pipeline.py",
    "ModuleNotFoundError: No module named 'src'",
])

body(
    "The error appeared when running 'python src/pipeline.py' directly. "
    "pipeline.py starts with 'from src.ingestion.csv_loader import load_csv_to_bronze'. "
    "Python tried to find a package called 'src' but could not find it."
)

h2("5.2 How Python Finds Modules")

body("When Python runs 'from src.ingestion.csv_loader import ...', it searches these locations in order:")
bullet("The directory containing the script being run")
bullet("Directories in the PYTHONPATH environment variable")
bullet("Standard library directories")
bullet("Site-packages (installed packages)")

body(
    "When you run 'python src/pipeline.py', the script directory is 'my-ai-project/src/'. "
    "Python adds 'src/' to its search path — but 'src' is not a package inside 'src/'. "
    "The 'src' package lives one level up, at 'my-ai-project/src/__init__.py'. "
    "Python was looking in the wrong place."
)

h2("5.3 Solution 1 — PYTHONPATH Environment Variable")

code_block([
    "# Add project root to Python's module search path before running:",
    "PYTHONPATH=. python src/pipeline.py",
    "",
    "# PYTHONPATH=. means: 'also look in the current directory (.) for modules'",
    "# Current directory = my-ai-project/",
    "# From there, Python can find src/__init__.py and the whole src package",
])

body(
    "This works for running scripts manually. But it requires you to always remember the prefix. "
    "pytest has its own module discovery mechanism that also needs the path fix."
)

h2("5.4 Solution 2 — conftest.py (the Better Fix)")

code_block([
    "# conftest.py at project root:",
    "import sys",
    "import os",
    "",
    "sys.path.insert(0, os.path.dirname(__file__))",
])

body(
    "conftest.py is a special pytest configuration file. pytest automatically discovers and runs "
    "conftest.py files before running any tests — no import required. "
    "os.path.dirname(__file__) is the absolute path to the directory containing conftest.py, "
    "which is the project root. sys.path.insert(0, ...) adds that path to the front of Python's "
    "module search list."
)

add_table(
    ["Approach", "When it works", "Limitation"],
    [
        ["PYTHONPATH=. python ...",     "Manually running scripts",          "Must remember the prefix every time"],
        ["conftest.py with sys.path",   "pytest discovers it automatically", "Only active during pytest sessions"],
        ["Both together",               "All scenarios covered",             "None — this is the correct setup"],
    ],
    col_widths=[2.2, 2.4, 2.0],
)

callout(
    "Why not just put src/ on the permanent PYTHONPATH?",
    "That would work but it couples your project to your machine's environment configuration. "
    "conftest.py is checked into the repository — it works on every developer's machine "
    "without any setup. PYTHONPATH in shell profile is personal and not reproducible.",
    bg="E3F2FD",
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — THE 6 NEW TESTS
# ══════════════════════════════════════════════════════════════════════════════

h1("6. The Six New Tests")

body(
    "Day 3 adds six tests: four in test_kpi_aggregator.py and two in test_pipeline.py. "
    "All six use the same isolation pattern as Day 2 tests — a fresh DuckDB file per test."
)

h2("6.1 Shared Fixture: gold_data")

code_block([
    "@pytest.fixture",
    "def gold_data(tmp_path):",
    "    from src.ingestion.csv_loader import load_csv_to_bronze",
    "    from src.transforms.bronze_to_silver import bronze_to_silver",
    "    from src.transforms.silver_to_gold import silver_to_gold",
    "",
    "    csv_path = tmp_path / 'projects.csv'",
    "    with open(csv_path, 'w', newline='') as f:",
    "        writer = csv.DictWriter(f, fieldnames=SAMPLE_PROJECTS[0].keys())",
    "        writer.writeheader()",
    "        writer.writerows(SAMPLE_PROJECTS)",
    "",
    "    load_csv_to_bronze(str(csv_path))",
    "    bronze_to_silver()",
    "    silver_to_gold()",
])

body(
    "The gold_data fixture runs the full pipeline through the gold layer. "
    "Tests that depend on this fixture receive a populated gold table with all 15 projects. "
    "Because set_test_db (autouse=True) ran first, this data goes into a temp DuckDB file "
    "that is deleted after the test."
)
body(
    "Notice the fixture returns nothing (no 'return' statement). The test does not need the fixture "
    "value — it just needs the side effect: a populated database. The fixture is listed as a "
    "parameter to signal the dependency: 'run gold_data setup before this test'."
)

h2("6.2 test_total_projects")

code_block([
    "def test_total_projects(gold_data):",
    "    from src.transforms.kpi_aggregator import compute_kpi_summary",
    "    summary = compute_kpi_summary()",
    "    assert summary['total_projects'] == 15",
])

body(
    "The most basic correctness test: 15 projects were loaded, 15 projects must be counted. "
    "If this fails, something in the ingestion or aggregation pipeline dropped rows."
)

h2("6.3 test_pct_sums_to_100")

code_block([
    "def test_pct_sums_to_100(gold_data):",
    "    summary = compute_kpi_summary()",
    "    total = (",
    "        summary['pct_on_track']",
    "        + summary['pct_at_risk']",
    "        + summary['pct_delayed']",
    "        + summary['pct_completed']",
    "    )",
    "    assert abs(total - 100.0) < 0.5, f'Percentages sum to {total}, expected ~100'",
])

body(
    "Every project must fall into exactly one status category. The four percentages must sum to 100%. "
    "abs(total - 100.0) < 0.5 uses a tolerance of 0.5 to absorb floating-point rounding "
    "(e.g. 33.3 + 33.3 + 33.4 = 100.0 exactly, but rounding could give 99.9 or 100.1)."
)

h2("6.4 test_top_3_risks_length")

code_block([
    "def test_top_3_risks_length(gold_data):",
    "    summary = compute_kpi_summary()",
    "    assert len(summary['top_3_risks']) == 3",
])

body(
    "The prompt builder and UI always expect exactly 3 items in top_3_risks. "
    "The SQL uses LIMIT 3 to guarantee this. This test confirms the LIMIT works correctly "
    "and that .df().to_dict('records') preserves the count."
)

h2("6.5 test_issues_by_dept_has_3_departments")

code_block([
    "def test_issues_by_dept_has_3_departments(gold_data):",
    "    summary = compute_kpi_summary()",
    "    depts = {row['department'] for row in summary['issues_by_dept']}",
    "    assert len(depts) == 3",
])

body(
    "The sample dataset has three departments (IT, Finance, Operations). "
    "This test uses a set comprehension to extract unique department names from the list of dicts, "
    "then asserts there are 3 unique departments. A set automatically removes duplicates, "
    "though there should only be one row per department due to GROUP BY."
)

h2("6.6 test_pipeline_returns_expected_keys")

code_block([
    "def test_pipeline_returns_expected_keys(sample_csv):",
    "    from src.pipeline import run_pipeline",
    "    result = run_pipeline(sample_csv)",
    "    assert 'kpi_summary' in result",
    "    assert 'ai_report' in result",
    "    assert result['kpi_summary']['total_projects'] == 15",
])

body(
    "This is an integration test — it runs the entire pipeline from CSV to final dict. "
    "It verifies the return value has the two expected keys and that the data arrived correctly. "
    "The ai_report value is a placeholder string (Day 5 will replace it with a real LLM call)."
)

h2("6.7 test_pipeline_completes_fast")

code_block([
    "def test_pipeline_completes_fast(sample_csv):",
    "    import time",
    "    start = time.time()",
    "    run_pipeline(sample_csv)",
    "    elapsed = time.time() - start",
    "    assert elapsed < 10, f'Pipeline took {elapsed:.1f}s — should be under 10s without LLM'",
])

body(
    "A performance regression test. The data pipeline (without LLM) should run in under 10 seconds. "
    "In practice it runs in about 0.08 seconds on most machines. If something unexpectedly slows it "
    "down (e.g. accidentally querying a remote database, or a loop over 15 rows in Python), "
    "this test will catch it before it reaches production."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — READING THE PIPELINE LOG
# ══════════════════════════════════════════════════════════════════════════════

h1("7. Reading the Pipeline Log Output")

body(
    "Running 'PYTHONPATH=. python src/pipeline.py' produces output like this:"
)

code_block([
    "2024-05-07 14:23:01.412 | INFO | src.pipeline:run_pipeline:13 - Stage 1: Ingestion",
    "2024-05-07 14:23:01.445 | INFO | src.ingestion.csv_loader:load_csv_to_bronze:22 - Loaded 15 rows into bronze_project_raw (0 skipped)",
    "2024-05-07 14:23:01.447 | INFO | src.pipeline:run_pipeline:18 -   → 15 rows loaded, 0 skipped",
    "2024-05-07 14:23:01.449 | INFO | src.pipeline:run_pipeline:21 - Stage 2: Bronze → Silver",
    "2024-05-07 14:23:01.461 | INFO | src.transforms.bronze_to_silver:bronze_to_silver:31 - Silver layer: 15 rows in silver_project_clean",
    "2024-05-07 14:23:01.462 | INFO | src.pipeline:run_pipeline:23 -   → 15 rows in silver",
    "2024-05-07 14:23:01.464 | INFO | src.pipeline:run_pipeline:26 - Stage 3: Silver → Gold",
    "2024-05-07 14:23:01.481 | INFO | src.transforms.silver_to_gold:silver_to_gold:21 - Gold layer: 15 rows in gold_project_kpi",
    "2024-05-07 14:23:01.482 | INFO | src.pipeline:run_pipeline:28 -   → 15 rows in gold",
    "2024-05-07 14:23:01.483 | INFO | src.pipeline:run_pipeline:31 - Stage 4: KPI Aggregation",
    "2024-05-07 14:23:01.511 | INFO | src.pipeline:run_pipeline:33 -   → 15 projects summarized",
    "2024-05-07 14:23:01.512 | INFO | src.pipeline:run_pipeline:35 - Stage 5: LLM Report [placeholder — wired on Day 4]",
    "2024-05-07 14:23:01.513 | INFO | src.pipeline:run_pipeline:40 - Pipeline complete in 0.1s",
    "15",
])

add_table(
    ["Log field", "What it means"],
    [
        ["2024-05-07 14:23:01.412",                    "Timestamp — when this log line was emitted"],
        ["INFO",                                         "Log level — INFO is routine progress; WARNING means unexpected; ERROR means failure"],
        ["src.pipeline:run_pipeline:13",                "Module : function : line number — tells you exactly where in the code this ran"],
        ["Stage 1: Ingestion",                           "The message text — written by logger.info()"],
        ["15 (last line)",                               "print(result['kpi_summary']['total_projects']) from __main__ block"],
    ],
    col_widths=[2.8, 4.0],
)

body(
    "The 0.1s total time (0.08s typically) is the benchmark for the data pipeline without LLM. "
    "Day 4 adds the LLM call. With claude-haiku-4-5, total pipeline time will be ~2-4 seconds including the API call."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — INSPECTING THE DATA
# ══════════════════════════════════════════════════════════════════════════════

h1("8. Inspecting the Data After Day 3")

body("Run the full pipeline once to populate the database, then inspect the results:")

h3("Run the pipeline")
code_block([
    "PYTHONPATH=. python src/pipeline.py",
    "# Expected output: 15",
])

h3("Inspect the KPI summary dict")
code_block([
    "PYTHONPATH=. python -c \"",
    "from dotenv import load_dotenv; load_dotenv()",
    "from src.pipeline import run_pipeline",
    "import json",
    "result = run_pipeline('data/input/project_data.csv')",
    "kpi = result['kpi_summary']",
    "print('Total projects:', kpi['total_projects'])",
    "print('On track:', kpi['pct_on_track'], '%')",
    "print('Delayed:', kpi['pct_delayed'], '%')",
    "print('Budget variance: $', kpi['total_budget_variance'])",
    "print('Top risk:', kpi['top_3_risks'][0]['project_name'])",
    "\"",
])

h3("Inspect top 3 risks")
code_block([
    "PYTHONPATH=. python -c \"",
    "from dotenv import load_dotenv; load_dotenv()",
    "from src.pipeline import run_pipeline",
    "result = run_pipeline('data/input/project_data.csv')",
    "for r in result['kpi_summary']['top_3_risks']:",
    "    print(r['project_name'], '|', r['risk_level'], '|', r['budget_variance_pct'], '%')",
    "\"",
    "",
    "# Expected output:",
    "Mobile App Launch | critical | 32.14 %",
    "Financial System Integration | high | 21.43 %",
    "Supply Chain Optimization | high | 9.71 %",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — KEY CONCEPTS
# ══════════════════════════════════════════════════════════════════════════════

h1("9. Key Concepts to Remember")

add_table(
    ["Concept", "What it means", "Why it matters"],
    [
        ["KPI aggregation",             "Collapse N rows → 1 summary dict",                    "AI gets clean numbers, not raw rows"],
        ["Parameterized queries",        "? placeholder instead of f-strings in SQL",           "Prevents SQL injection; handles special chars"],
        ["Orchestrator pattern",         "One file knows the execution order",                  "Isolated layers, single place to change sequence"],
        ["CASE ORDER trick",             "Map text labels to integers for ORDER BY",            "SQL cannot order text labels by business meaning"],
        ["FIRST() aggregation",          "Pick any one value from a group",                     "Required when GROUP BY has multi-valued columns"],
        ["Portfolio vs average pct",     "SUM/SUM, not AVG(pct)",                               "Averaging percentages gives wrong portfolio view"],
        ["conftest.py sys.path fix",     "Add project root to Python's module search path",    "pytest finds 'src' package on all machines"],
        ["PYTHONPATH=. prefix",          "Same fix for manual script execution",                 "No global environment changes needed"],
        ["Performance regression test",  "Assert pipeline runs under N seconds",                "Catches accidental slowness before it ships"],
    ],
    col_widths=[1.8, 2.4, 2.4],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — TEST RESULTS
# ══════════════════════════════════════════════════════════════════════════════

h1("10. Final Test Results After Day 3")

body("All 15 tests pass — 3 ingestion + 6 transforms + 4 aggregator + 2 pipeline:")

code_block([
    "$ pytest tests/ -v",
    "",
    "tests/test_ingestion.py::test_load_happy_path                         PASSED",
    "tests/test_ingestion.py::test_missing_required_column                  PASSED",
    "tests/test_ingestion.py::test_empty_project_name_rows_skipped          PASSED",
    "tests/test_transforms.py::test_silver_date_type                        PASSED",
    "tests/test_transforms.py::test_silver_status_normalization              PASSED",
    "tests/test_transforms.py::test_silver_null_issue_description            PASSED",
    "tests/test_transforms.py::test_gold_budget_variance_math               PASSED",
    "tests/test_transforms.py::test_gold_is_late_flag                        PASSED",
    "tests/test_transforms.py::test_gold_on_time_not_late                    PASSED",
    "tests/test_kpi_aggregator.py::test_total_projects                       PASSED",
    "tests/test_kpi_aggregator.py::test_pct_sums_to_100                      PASSED",
    "tests/test_kpi_aggregator.py::test_top_3_risks_length                   PASSED",
    "tests/test_kpi_aggregator.py::test_issues_by_dept_has_3_departments     PASSED",
    "tests/test_pipeline.py::test_pipeline_returns_expected_keys             PASSED",
    "tests/test_pipeline.py::test_pipeline_completes_fast                    PASSED",
    "",
    "15 passed in 1.43s",
])

h2("Files Created on Day 3")

add_table(
    ["File", "Purpose"],
    [
        ["src/transforms/kpi_aggregator.py", "8 SQL queries → portfolio KPI dict"],
        ["src/pipeline.py",                  "Orchestrates all 5 stages, Stage 5 is placeholder"],
        ["tests/test_kpi_aggregator.py",     "4 tests for aggregation correctness"],
        ["tests/test_pipeline.py",           "2 end-to-end pipeline tests"],
        ["conftest.py",                      "sys.path fix for pytest module resolution"],
    ],
    col_widths=[2.8, 3.8],
)

h1("11. What Day 4 Builds On This")

body("Day 3 left Stage 5 of the pipeline as a placeholder string. Day 4 will:")

bullet("Create src/llm/prompt_builder.py — inject KPI dict values into a structured prompt template")
bullet("Create src/llm/claude_client.py — call the Anthropic API and return the generated text")
bullet("Update src/pipeline.py — wire Stage 5 to call build_prompt() then generate_report()")
bullet("Add tests/test_prompt_builder.py — 3 tests verifying prompt correctness without API calls")
bullet("Add a mocked LLM test to test_pipeline.py — verifies pipeline still passes with mocker.patch()")
bullet("Checkpoint: PYTHONPATH=. python src/pipeline.py prints a 5-section report with real project names")

body(
    "After Day 4, running the pipeline will produce a professional executive report that "
    "mentions actual project names like 'Mobile App Launch' and 'Financial System Integration' — "
    "narrated by Claude claude-haiku-4-5 from the injected KPI data."
)

divider()

p_final = doc.add_paragraph()
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_final.paragraph_format.space_before = Pt(20)
run_final = p_final.add_run("Day 3 complete — 15 tests passing, full data pipeline running in under 0.1s.")
set_font(run_final, bold=True, size=12, color=(39, 174, 96))

# ── Save ───────────────────────────────────────────────────────────────────────
out_path = "docs/day-3-learning-guide.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
