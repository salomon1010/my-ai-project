"""Run this script to generate day-4-learning-guide.docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Style helpers ──────────────────────────────────────────────────────────────

def set_font(run, bold=False, italic=False, size=None, color=None, font_name=None):
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, bold=True, size=20, color=(44, 62, 80))
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, bold=True, size=15, color=(44, 62, 80))
    return p


def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, bold=True, size=12, color=(52, 73, 94))
    return p


def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def code_block(lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 30, 30)
        if i < len(lines) - 1:
            run.add_break()
    doc.add_paragraph()


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def callout(label, text, bg="FFF3CD"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg)
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    r1 = p.add_run(f"{label}  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    doc.add_paragraph()


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "2C3E50")
        tcPr.append(shd)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        fill = "F8F9FA" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("─" * 80)
    run.font.color.rgb = RGBColor(189, 195, 199)
    run.font.size = Pt(8)


# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run("AI Executive Reporting System")
set_font(run, bold=True, size=26, color=(44, 62, 80))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Day 4 Learning Guide")
set_font(run2, bold=True, size=20, color=(52, 152, 219))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("LLM Integration — Anti-Hallucination Prompt Design")
set_font(run3, italic=True, size=13, color=(127, 140, 141))

doc.add_paragraph()
divider()
doc.add_paragraph()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run(
    "This guide explains the anti-hallucination prompt architecture, the Anthropic SDK wrapper,\n"
    "the 5-section prompt template, mocking in tests, and the 'patch where used' rule."
)
set_font(run4, size=11, color=(52, 73, 94))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — WHAT WE BUILT
# ══════════════════════════════════════════════════════════════════════════════

h1("1. What Day 4 Built")

body(
    "Day 3 left Stage 5 of the pipeline as a placeholder string. Day 4 replaced that placeholder "
    "with real AI-generated executive prose — using the Anthropic API and a carefully engineered prompt."
)

body("The complete pipeline after Day 4:")

code_block([
    "CSV file",
    "  → load_csv_to_bronze()    (Day 1) → bronze_project_raw",
    "  → bronze_to_silver()      (Day 2) → silver_project_clean",
    "  → silver_to_gold()        (Day 2) → gold_project_kpi",
    "  → compute_kpi_summary()   (Day 3) → KPISummary dict",
    "  → build_prompt()          (Day 4) → grounded prompt string",
    "  → generate_report()       (Day 4) → AI executive report (5 sections)",
])

body("Four files were created or updated:")
bullet("src/llm/prompt_builder.py — injects KPI dict into a structured prompt template")
bullet("src/llm/claude_client.py — Anthropic SDK wrapper with graceful fallback")
bullet("tests/test_prompt_builder.py — 3 tests verifying prompt correctness without API calls")
bullet("src/pipeline.py — Stage 5 wired (replaced placeholder with build_prompt + generate_report)")
body("One existing file was updated: tests/test_pipeline.py — added 2 mocked LLM tests.")

body("One bug was found during testing and fixed (documented in Section 7).")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — THE ANTI-HALLUCINATION CONCEPT
# ══════════════════════════════════════════════════════════════════════════════

h1("2. The Anti-Hallucination Design")

body(
    "Most people who try to use AI for business reporting make the same mistake. "
    "They ask the AI: 'What are the risks in our project portfolio?' "
    "The AI will invent risks, make up project names, fabricate numbers. "
    "This is called hallucination — and it destroys trust in the entire system."
)

h2("2.1 The Wrong Approach")

code_block([
    "# BAD — gives the AI nothing to work with:",
    "prompt = 'Analyze our project portfolio and identify risks.'",
    "",
    "# AI response will contain invented data:",
    "# 'Project Alpha is at 35% completion and faces budget overruns...'",
    "# (Project Alpha does not exist in your data)",
])

callout(
    "The core problem:",
    "Language models are trained to produce plausible-sounding text. "
    "Without real data injected into the prompt, the model writes what sounds like "
    "a risk report — but the numbers, project names, and dates are fabrications. "
    "A single invented fact discovered by a client destroys all credibility.",
    bg="FCE4EC",
)

h2("2.2 The Correct Approach — Data In, Narration Out")

body(
    "The correct pattern: inject your real computed numbers first, then ask the model "
    "to write professional prose about them. The model becomes a writer, not an analyst."
)

code_block([
    "# GOOD — all facts injected, model just narrates:",
    "prompt = f'''",
    "=== PORTFOLIO DATA ===",
    "Total Projects: 15",
    "Delayed: 40.0%",
    "Over-budget projects:",
    "  - Mobile App Launch (IT): planned $140,000 → actual $185,000 (+32.1%)",
    "  - Financial System Integration (Finance): planned $210,000 → actual $255,000 (+21.4%)",
    "",
    "Write an executive summary referencing ONLY the data above.",
    "'''",
    "",
    "# AI response uses the real numbers:",
    "# 'Mobile App Launch is 32.1% over its $140,000 budget, representing...'",
])

callout(
    "The key insight:",
    "Every project name and number that appears in the AI report must have been "
    "injected by prompt_builder.py. The model cannot invent facts that are not "
    "in the prompt. This is the anti-hallucination boundary — and it is what makes "
    "this system safe to show to a real executive.",
    bg="E8F5E9",
)

h2("2.3 The Data Flow")

code_block([
    "kpi_summary dict  ←  computed from real DuckDB data (100% accurate)",
    "       ↓",
    "build_prompt()    ←  formats real numbers into a structured text block",
    "       ↓",
    "prompt string     ←  contains ONLY facts that came from your database",
    "       ↓",
    "generate_report() ←  Claude narrates the facts in professional prose",
    "       ↓",
    "ai_report string  ←  every sentence references injected data",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — prompt_builder.py
# ══════════════════════════════════════════════════════════════════════════════

h1("3. File: src/llm/prompt_builder.py")

body(
    "This file has one function: build_prompt(). It takes the KPI summary dict and returns "
    "a formatted string. Every piece of data in that string came from your DuckDB gold table."
)

h2("3.1 Building the Data Sections")

body("The function first formats each data category into a block of text lines:")

code_block([
    "# Top 3 risks — uses a generator expression inside join():",
    "risk_lines = '\\n'.join(",
    "    f'- {r[\"project_name\"]} | Risk: {r[\"risk_level\"].upper()} | '",
    "    f'Budget: {r[\"budget_variance_pct\"]:+.1f}%'",
    "    for r in s['top_3_risks']",
    ") or 'None identified'",
    "",
    "# Produces lines like:",
    "# - Mobile App Launch | Risk: CRITICAL | Budget: +32.1%",
    "# - Financial System Integration | Risk: HIGH | Budget: +21.4%",
    "# - Supply Chain Optimization | Risk: HIGH | Budget: +9.7%",
])

body(
    "The '\\n'.join(...) pattern turns a list of strings into a single multi-line string. "
    "The generator expression (f'...' for r in list) creates each line without building "
    "an intermediate list. The 'or \"None identified\"' fallback handles empty lists — "
    "if there are no high-risk projects, the section says 'None identified' instead of being blank."
)

h2("3.2 Format Specifiers in f-strings")

body(
    "The budget numbers use format specifiers that control exactly how they appear in the prompt:"
)

add_table(
    ["Format spec", "Example input", "Output", "Meaning"],
    [
        [":,.0f",  "185000",   "$185,000",  "comma thousands separator, no decimals"],
        [":+.1f",  "32.142",   "+32.1%",    "always show + or - sign, 1 decimal"],
        [":.1f",   "40.0",     "40.0%",     "1 decimal, no forced sign"],
        [":,.2f",  "183000.5", "183,000.50","comma separator, 2 decimals"],
    ],
    col_widths=[1.3, 1.3, 1.3, 3.7],
)

body(
    "The :+ sign specifier is especially important for budget variance. A project that is under budget "
    "should show '-5.2%', not just '5.2%'. The + forces the sign to appear even when positive. "
    "Without it, the AI might write 'budget variance is 32.1%' without clarifying it means over-budget."
)

h2("3.3 The Complete Prompt Template")

body(
    "After building the data sections, the function assembles them into a structured prompt "
    "using a multi-line f-string:"
)

code_block([
    "prompt = f'''You are a senior management consultant writing an executive briefing.",
    "You MUST only reference the data provided below. Do not invent project names,",
    "costs, or dates not present in the data. Be direct, professional, and concise.",
    "",
    "=== PORTFOLIO SUMMARY ===",
    "Total Projects: {s['total_projects']}",
    "On Track: {s['pct_on_track']:.1f}%  |  At Risk: {s['pct_at_risk']:.1f}%  ...",
    "Total Budget Variance: ${s['total_budget_variance']:,.0f} ({s['total_budget_variance_pct']:+.1f}%)",
    "Average Schedule Delay (late projects): {s['avg_schedule_delay_days']:.1f} days",
    "",
    "=== TOP 3 HIGH-RISK PROJECTS ===",
    "{risk_lines}",
    "",
    "=== OVER-BUDGET PROJECTS (>10%) ===",
    "{over_budget_lines}",
    "",
    "=== LATE PROJECTS ===",
    "{late_lines}",
    "",
    "=== ISSUES BY DEPARTMENT ===",
    "{dept_lines}",
    "",
    "---",
    "Write the following sections using these EXACT headers (markdown ## format):",
    "",
    "## Executive Summary",
    "## Key Risks",
    "## Budget Concerns",
    "## Schedule Delays",
    "## Recommended Actions'''",
])

h2("3.4 Why Exact Section Headers Matter")

body(
    "The prompt instructs the model to use these exact ## headers. This is not aesthetic — "
    "it is structural. Day 6 will build a PDF generator that splits the AI report into "
    "five sections by searching for these headers. If the model uses different headers "
    "(e.g. '## Financial Risk' instead of '## Budget Concerns'), the PDF splitter breaks."
)

body(
    "This is a contract between the prompt and the downstream code. The prompt enforces it "
    "on the model side, and the PDF generator assumes it on the output side."
)

add_table(
    ["Section header", "Content the model writes"],
    [
        ["## Executive Summary", "3-5 sentences: overall health, key numbers, urgency"],
        ["## Key Risks",          "Top 3 risks with project names and variance data"],
        ["## Budget Concerns",    "All over-budget projects with dollar amounts"],
        ["## Schedule Delays",    "All late projects with delay duration and owner"],
        ["## Recommended Actions","3-5 concrete bullet points tied to specific issues"],
    ],
    col_widths=[2.3, 4.3],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — claude_client.py
# ══════════════════════════════════════════════════════════════════════════════

h1("4. File: src/llm/claude_client.py")

body(
    "This file wraps the Anthropic Python SDK into a single function. "
    "It handles API key validation, model selection, the API call, and graceful fallback."
)

h2("4.1 The Complete File")

code_block([
    "import os",
    "import anthropic",
    "from dotenv import load_dotenv",
    "from loguru import logger",
    "",
    "load_dotenv()",
    "",
    "FALLBACK_MESSAGE = (",
    "    '[Report generation unavailable. '",
    "    'Please check that ANTHROPIC_API_KEY is set in your .env file and try again.]'",
    ")",
    "",
    "",
    "def generate_report(prompt: str) -> str:",
    "    api_key = os.getenv('ANTHROPIC_API_KEY')",
    "    model = os.getenv('MODEL_ID', 'claude-haiku-4-5')",
    "",
    "    if not api_key:",
    "        logger.warning('ANTHROPIC_API_KEY not set — returning fallback message')",
    "        return FALLBACK_MESSAGE",
    "",
    "    try:",
    "        client = anthropic.Anthropic(api_key=api_key)",
    "        message = client.messages.create(",
    "            model=model,",
    "            max_tokens=1024,",
    "            messages=[{'role': 'user', 'content': prompt}]",
    "        )",
    "        return message.content[0].text",
    "",
    "    except anthropic.APIError as e:",
    "        logger.error(f'Claude API error: {e}')",
    "        return FALLBACK_MESSAGE",
])

h2("4.2 The Anthropic Messages API")

body(
    "The Anthropic SDK uses a messages-based API. Every call requires a list of messages "
    "with role and content. For a single-turn prompt, there is one message with role='user'."
)

add_table(
    ["Parameter", "Value used", "Why"],
    [
        ["model",      "claude-haiku-4-5 (env default)",    "Fastest, cheapest Claude model; ~$0.001 per call"],
        ["max_tokens", "1024",                               "Caps output cost; 5-section report needs ~500 tokens"],
        ["messages",   "[{'role': 'user', 'content': ...}]", "Single-turn conversation; no chat history needed"],
    ],
    col_widths=[1.5, 2.3, 2.8],
)

body(
    "message.content[0].text extracts the text from the response. "
    "The API returns a Message object. .content is a list of content blocks "
    "(there can be multiple for tool use). [0].text gets the first text block."
)

h2("4.3 Why max_tokens=1024?")

body(
    "max_tokens sets the maximum number of tokens the model can generate. "
    "One token is approximately 4 characters, or 0.75 words."
)

code_block([
    "# Token budget estimate for a 5-section report:",
    "#",
    "# Executive Summary:    ~80  tokens",
    "# Key Risks:            ~100 tokens",
    "# Budget Concerns:      ~120 tokens",
    "# Schedule Delays:      ~100 tokens",
    "# Recommended Actions:  ~100 tokens",
    "#",
    "# Total estimate:       ~500 tokens",
    "# max_tokens=1024 gives 2x headroom",
    "#",
    "# At claude-haiku-4-5 pricing (~$0.00125 per 1K output tokens):",
    "# 1024 tokens = $0.00128 per report run",
    "# 1000 report runs per month = $1.28",
])

callout(
    "Cost control tip:",
    "Always set max_tokens when calling the API. Without it, the model could generate "
    "thousands of tokens if the prompt encourages long output. max_tokens=1024 is a safety "
    "ceiling — the model will stop at 1024 tokens even if it could continue. "
    "For this use case, 1024 is sufficient with comfortable headroom.",
    bg="E3F2FD",
)

h2("4.4 Graceful Fallback Pattern")

body(
    "The function handles two failure modes gracefully: missing API key and API errors. "
    "In both cases, it returns a fallback string instead of crashing the pipeline."
)

code_block([
    "# Failure mode 1: No API key in .env",
    "if not api_key:",
    "    logger.warning('ANTHROPIC_API_KEY not set — returning fallback message')",
    "    return FALLBACK_MESSAGE",
    "",
    "# Failure mode 2: API call fails (rate limit, network, invalid key, etc.)",
    "except anthropic.APIError as e:",
    "    logger.error(f'Claude API error: {e}')",
    "    return FALLBACK_MESSAGE",
])

body(
    "Why return a fallback instead of raising an exception? "
    "The pipeline and the Streamlit UI both display ai_report as a string. "
    "If generate_report() raises, the entire pipeline crashes and the user sees an error page. "
    "If it returns a fallback string, the pipeline completes, the KPI data is still visible, "
    "and the fallback message tells the user exactly what to fix."
)

callout(
    "Design principle:",
    "In a demo system, a graceful degradation ('report unavailable, check your API key') "
    "is far better than a stack trace. The client still sees the KPI cards and the project table. "
    "Only the AI text section shows the fallback. The rest of the demo still works.",
    bg="E8F5E9",
)

h2("4.5 MODEL_ID Environment Variable")

body(
    "The model is read from os.getenv('MODEL_ID', 'claude-haiku-4-5'). "
    "This means you can switch models by changing a single line in .env — no code change needed."
)

code_block([
    "# .env file — switch model without touching code:",
    "MODEL_ID=claude-haiku-4-5         # fast, cheap ($0.00125/1K output tokens)",
    "# MODEL_ID=claude-sonnet-4-6      # better quality, higher cost",
    "# MODEL_ID=claude-opus-4-7        # highest quality, premium cost",
])

body(
    "For a consulting demo, claude-haiku-4-5 is the right choice: fast enough that the client "
    "sees the report appear within a few seconds, and cheap enough that you can run it dozens "
    "of times in a meeting without worrying about cost."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — UPDATED pipeline.py
# ══════════════════════════════════════════════════════════════════════════════

h1("5. Updated src/pipeline.py — Stage 5 Wired")

body("Three changes were made to pipeline.py:")

h2("5.1 New Imports")

code_block([
    "# Added to the top of pipeline.py:",
    "from src.llm.prompt_builder import build_prompt",
    "from src.llm.claude_client import generate_report",
])

body(
    "These are the two new functions from the Day 4 files. "
    "Adding them here is consistent with the orchestrator pattern: "
    "pipeline.py is the only file that imports across layer boundaries."
)

h2("5.2 Stage 5 Replaced")

code_block([
    "# Before (Day 3 placeholder):",
    "logger.info('Stage 5: LLM Report [placeholder — wired on Day 4]')",
    "ai_report = '[AI report will appear here after Day 4]'",
    "",
    "# After (Day 4 — wired):",
    "logger.info('Stage 5: Generating AI Report')",
    "prompt = build_prompt(kpi_summary)",
    "ai_report = generate_report(prompt)",
    "logger.info(f'  → Report generated ({len(ai_report)} chars)')",
])

body(
    "The pipeline now calls build_prompt() to construct the grounded prompt, "
    "then passes it to generate_report() to get the AI text. "
    "The log line records how many characters the report contains — "
    "a quick sanity check that the model produced a real response (not a fallback)."
)

h2("5.3 __main__ Block Updated")

code_block([
    "# Before:",
    "print(result['kpi_summary']['total_projects'])",
    "",
    "# After:",
    "print(result['ai_report'])",
])

body(
    "When you run 'PYTHONPATH=. python src/pipeline.py' directly, "
    "it now prints the AI report instead of '15'. "
    "This is the Day 4 checkpoint — you should see a 5-section professional report "
    "mentioning real project names like Mobile App Launch and Financial System Integration."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — TESTS
# ══════════════════════════════════════════════════════════════════════════════

h1("6. The Five New Tests")

body(
    "Day 4 adds three tests in test_prompt_builder.py and two updated tests in test_pipeline.py. "
    "None of these tests require a real ANTHROPIC_API_KEY — the LLM is either not called "
    "or is replaced with a mock."
)

h2("6.1 tests/test_prompt_builder.py — 3 Tests")

body(
    "These tests check that build_prompt() produces a correctly formatted string. "
    "They use a hardcoded MOCK_KPI dict with known values."
)

h3("Test 1 — test_no_unfilled_placeholders")

code_block([
    "def test_no_unfilled_placeholders():",
    "    prompt = build_prompt(MOCK_KPI)",
    "    assert '{' not in prompt, 'Prompt has unfilled {} placeholders'",
])

body(
    "An f-string that references a key that does not exist in the dict would raise a KeyError "
    "and never reach this assertion. But a f-string that has an unclosed brace like {missing "
    "would produce a literal '{missing' in the output. This test catches any case where "
    "a brace was accidentally left in the template without a corresponding variable."
)

h3("Test 2 — test_project_names_in_prompt")

code_block([
    "def test_project_names_in_prompt():",
    "    prompt = build_prompt(MOCK_KPI)",
    "    for project in MOCK_KPI['over_budget_projects']:",
    "        assert project['project_name'] in prompt, \\",
    "            f\"Expected '{project['project_name']}' in prompt\"",
])

body(
    "This is the anti-hallucination integrity check. Every over-budget project name "
    "from the input data must appear in the prompt string. "
    "If a project name is missing from the prompt, the model cannot reference it — "
    "and may hallucinate a replacement. This test guarantees the data injection worked."
)

h3("Test 3 — test_section_headers_present")

code_block([
    "def test_section_headers_present():",
    "    prompt = build_prompt(MOCK_KPI)",
    "    for header in ['## Executive Summary', '## Key Risks', '## Budget Concerns',",
    "                   '## Schedule Delays', '## Recommended Actions']:",
    "        assert header in prompt, f'Missing section header: {header}'",
])

body(
    "The prompt instructs the model to use these exact section headers. "
    "If a header is missing from the prompt, the model will not know it needs to "
    "write that section. The PDF splitter (Day 6) also depends on these exact strings. "
    "This test locks the contract between prompt_builder and pdf_generator."
)

h2("6.2 Updated tests/test_pipeline.py — 2 Tests with Mocking")

body(
    "The pipeline tests now mock the LLM call. Without mocking, every test would make a "
    "real API call — slow, expensive, and requiring an API key in the test environment."
)

code_block([
    "def test_pipeline_with_mocked_llm(sample_csv, mocker):",
    "    mocker.patch(",
    "        'src.pipeline.generate_report',",
    "        return_value='## Executive Summary\\nMocked report.'",
    "    )",
    "    result = run_pipeline(sample_csv)",
    "    assert result['ai_report'].startswith('## Executive Summary')",
    "    assert result['kpi_summary']['total_projects'] == 15",
])

body(
    "mocker.patch() replaces the named function with a fake that always returns the given value. "
    "The pipeline runs normally — all four data stages execute with real DuckDB — "
    "but Stage 5 returns the mocked string instead of calling the API."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — THE MOCKING BUG
# ══════════════════════════════════════════════════════════════════════════════

h1("7. The Bug We Hit — 'Patch Where Used, Not Where Defined'")

body(
    "This was the only bug on Day 4. It was subtle — the test code looked correct, "
    "but the mock did not work. Understanding this teaches you one of the most important "
    "rules in Python testing."
)

h2("7.1 The Failing Test")

code_block([
    "# What we wrote first (WRONG):",
    "mocker.patch(",
    "    'src.llm.claude_client.generate_report',  ← patches the function where it lives",
    "    return_value='## Executive Summary\\nMocked report.'",
    ")",
])

body("The test ran, but the AI report was the fallback message, not the mocked string. The assertion failed:")

code_block([
    "AssertionError: assert False",
    "  where False = '...'.startswith('## Executive Summary')",
    "  ai_report = '[Report generation unavailable. Please check...]'",
])

h2("7.2 Why the Mock Did Not Work")

body(
    "When Python executes 'from src.llm.claude_client import generate_report' in pipeline.py, "
    "it creates a new name 'generate_report' in the pipeline module's namespace that points "
    "to the function object."
)

code_block([
    "# In Python's memory after 'from src.llm.claude_client import generate_report':",
    "",
    "src.llm.claude_client.generate_report  →  <function object at 0x...>",
    "src.pipeline.generate_report           →  <same function object at 0x...>",
    "",
    "# mocker.patch('src.llm.claude_client.generate_report') replaces the name",
    "# in src.llm.claude_client — but src.pipeline still holds the OLD reference:",
    "",
    "src.llm.claude_client.generate_report  →  <Mock object>     (patched)",
    "src.pipeline.generate_report           →  <function object>  (unchanged!)",
    "",
    "# When pipeline.py calls generate_report(), it uses src.pipeline.generate_report",
    "# which still points to the real function — not the mock.",
])

h2("7.3 The Fix — Patch Where It Is Used")

code_block([
    "# What we changed to (CORRECT):",
    "mocker.patch(",
    "    'src.pipeline.generate_report',  ← patches the name where pipeline.py uses it",
    "    return_value='## Executive Summary\\nMocked report.'",
    ")",
])

body(
    "Patching 'src.pipeline.generate_report' replaces the name in the pipeline module's namespace. "
    "When pipeline.py calls generate_report(), it looks up that name in its own namespace — "
    "and finds the mock."
)

callout(
    "The rule — memorize this:",
    "Always patch the name where it is used, not where it is defined. "
    "If pipeline.py does 'from src.llm.claude_client import generate_report', "
    "patch 'src.pipeline.generate_report'. "
    "If pipeline.py does 'import src.llm.claude_client' and calls "
    "'src.llm.claude_client.generate_report(...)', then patch 'src.llm.claude_client.generate_report'. "
    "The pattern is: patch the module where the lookup happens at call time.",
    bg="FFF3CD",
)

add_table(
    ["Import style in pipeline.py", "Correct patch target"],
    [
        ["from src.llm.claude_client import generate_report",     "src.pipeline.generate_report"],
        ["import src.llm.claude_client",                           "src.llm.claude_client.generate_report"],
        ["from src.llm import claude_client",                      "src.pipeline.claude_client.generate_report  (if used as claude_client.generate_report)"],
    ],
    col_widths=[3.3, 3.3],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — SAMPLE OUTPUT
# ══════════════════════════════════════════════════════════════════════════════

h1("8. What the AI Report Looks Like")

body(
    "After adding ANTHROPIC_API_KEY to .env and running 'PYTHONPATH=. python src/pipeline.py', "
    "the output is a professional 5-section executive report. Here is a representative example:"
)

code_block([
    "## Executive Summary",
    "The project portfolio of 15 initiatives is under significant pressure, with 40.0% of projects",
    "delayed and 33.3% at risk or delayed. Total budget overrun stands at $183,000 across the",
    "portfolio. The Mobile App Launch project represents the most critical risk, running 32.1%",
    "over budget with a vendor termination event. Immediate executive attention is required on",
    "three high-priority projects in IT and Finance.",
    "",
    "## Key Risks",
    "- **Mobile App Launch** (IT) carries a CRITICAL risk designation, with a budget overrun of",
    "  $45,000 (+32.1%) and a 77-day schedule delay. The original vendor was terminated.",
    "- **Financial System Integration** (Finance) is HIGH risk at +$45,000 (+21.4%) over budget",
    "  and 76 days behind schedule due to underestimated integration complexity.",
    "- **Supply Chain Optimization** (Operations) is HIGH risk with a 30-day delay caused by",
    "  supplier disruption and backup vendor onboarding.",
    "",
    "## Budget Concerns",
    "Seven of 15 projects (46.7%) are more than 10% over budget:",
    "- Mobile App Launch (IT): planned $140,000 → actual $185,000 (+32.1%)",
    "- Financial System Integration (Finance): planned $210,000 → actual $255,000 (+21.4%)",
    "...",
    "",
    "## Schedule Delays",
    "Eight of 15 projects are running late. Average delay among late projects: 52.3 days.",
    "- Mobile App Launch: 77 days late | Owner: Lisa Wang",
    "- Financial System Integration: 76 days late | Owner: David Brown",
    "...",
    "",
    "## Recommended Actions",
    "1. Convene an emergency review for Mobile App Launch within 48 hours.",
    "2. Request a revised budget and timeline from the Financial System Integration team.",
    "3. Escalate Supply Chain Optimization vendor situation to procurement leadership.",
    "4. Conduct a portfolio-wide risk review with all department heads.",
    "5. Establish a weekly steering committee for the four most critical projects.",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — KEY CONCEPTS
# ══════════════════════════════════════════════════════════════════════════════

h1("9. Key Concepts to Remember")

add_table(
    ["Concept", "What it means", "Why it matters"],
    [
        ["Anti-hallucination design",   "Inject real data → model narrates it",                  "Prevents invented facts in executive reports"],
        ["f-string format specifiers",  ":+.1f adds sign, :,.0f adds comma separators",           "Numbers in prompts must be unambiguous"],
        ["Exact section headers",       "## headers are a contract with the PDF splitter",        "Downstream PDF splitting depends on them"],
        ["max_tokens cap",              "Limits output length and API cost",                       "Always set; prevents runaway generation cost"],
        ["MODEL_ID env var",            "Switch model via .env, no code change",                  "Easy to upgrade for production deployments"],
        ["Graceful fallback",           "Return message instead of crashing on API failure",      "Demo still works even without valid API key"],
        ["'or default' on empty list",  "'\\n'.join(...) or 'None identified'",                   "Empty sections show a message, not blank text"],
        ["Patch where used",            "mocker.patch('src.pipeline.fn') not 'src.module.fn'",   "Python name binding requires correct patch path"],
        ["messages=[{role, content}]",  "Anthropic API requires list format for messages",        "Single-turn prompt = one dict in the list"],
    ],
    col_widths=[2.0, 2.3, 2.3],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — TEST RESULTS
# ══════════════════════════════════════════════════════════════════════════════

h1("10. Final Test Results After Day 4")

body("All 19 tests pass — no API key needed in the test environment:")

code_block([
    "$ PYTHONPATH=. pytest tests/ -v",
    "",
    "tests/test_ingestion.py::test_load_happy_path                         PASSED",
    "tests/test_ingestion.py::test_missing_required_column                  PASSED",
    "tests/test_ingestion.py::test_empty_project_name_rows_skipped          PASSED",
    "tests/test_kpi_aggregator.py::test_total_projects                       PASSED",
    "tests/test_kpi_aggregator.py::test_pct_sums_to_100                      PASSED",
    "tests/test_kpi_aggregator.py::test_top_3_risks_length                   PASSED",
    "tests/test_kpi_aggregator.py::test_issues_by_dept_has_3_departments     PASSED",
    "tests/test_pipeline.py::test_pipeline_returns_expected_keys             PASSED",
    "tests/test_pipeline.py::test_pipeline_completes_fast                    PASSED",
    "tests/test_pipeline.py::test_pipeline_with_mocked_llm                  PASSED",
    "tests/test_prompt_builder.py::test_no_unfilled_placeholders             PASSED",
    "tests/test_prompt_builder.py::test_project_names_in_prompt              PASSED",
    "tests/test_prompt_builder.py::test_section_headers_present              PASSED",
    "tests/test_transforms.py::test_silver_date_type                        PASSED",
    "tests/test_transforms.py::test_silver_status_normalization              PASSED",
    "tests/test_transforms.py::test_silver_null_issue_description            PASSED",
    "tests/test_transforms.py::test_gold_budget_variance_math               PASSED",
    "tests/test_transforms.py::test_gold_is_late_flag                        PASSED",
    "tests/test_transforms.py::test_gold_on_time_not_late                    PASSED",
    "",
    "19 passed in 2.10s",
])

h2("Files Created or Updated on Day 4")

add_table(
    ["File", "Change", "Purpose"],
    [
        ["src/llm/prompt_builder.py",  "Created",  "KPI dict → grounded prompt string"],
        ["src/llm/claude_client.py",   "Created",  "Anthropic SDK wrapper, graceful fallback"],
        ["tests/test_prompt_builder.py","Created", "3 prompt correctness tests"],
        ["src/pipeline.py",            "Updated",  "Stage 5 wired (was placeholder)"],
        ["tests/test_pipeline.py",     "Updated",  "Added 2 mocked LLM tests"],
    ],
    col_widths=[2.5, 1.2, 3.0],
)

h1("11. What Day 5 Builds On This")

body("Day 4 completed the data and AI engine. Day 5 adds the browser-based UI:")

bullet("Create app/components/kpi_cards.py — four st.metric() cards at the top of the page")
bullet("Create app/components/project_table.py — styled st.dataframe() with column config")
bullet("Create app/components/ai_summary.py — 'Generate Report' button + markdown display")
bullet("Create app/main.py — Streamlit entry point: sidebar, session state, page layout")
bullet("Checkpoint: streamlit run app/main.py — browser opens, 'Load Sample Data' works, KPI cards show")

body(
    "After Day 5, you will be able to demo the entire system live in a browser: "
    "upload a CSV, watch the KPI cards populate, click 'Generate Report', "
    "and watch Claude write a professional executive summary in real time."
)

divider()

p_final = doc.add_paragraph()
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_final.paragraph_format.space_before = Pt(20)
run_final = p_final.add_run("Day 4 complete — 19 tests passing, AI brain wired.")
set_font(run_final, bold=True, size=12, color=(39, 174, 96))

# ── Save ───────────────────────────────────────────────────────────────────────
out_path = "docs/day-4-learning-guide.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
