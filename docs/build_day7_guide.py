"""Run this script to generate day-7-learning-guide.docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

def set_font(run, bold=False, italic=False, size=None, color=None, font_name=None):
    run.bold = bold
    run.italic = italic
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    if font_name: run.font.name = font_name

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, bold=True, size=20, color=(44, 62, 80))
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, bold=True, size=15, color=(44, 62, 80))
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, bold=True, size=12, color=(52, 73, 94))
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if indent: p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def code_block(lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 30, 30)
        if i < len(lines) - 1:
            run.add_break()
    doc.add_paragraph()

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def callout(label, text, bg="FFF3CD"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg)
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    r1 = p.add_run(f"{label}  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    doc.add_paragraph()

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "2C3E50")
        tcPr.append(shd)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        fill = "F8F9FA" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("─" * 80)
    run.font.color.rgb = RGBColor(189, 195, 199)
    run.font.size = Pt(8)

# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run("AI Executive Reporting System")
set_font(run, bold=True, size=26, color=(44, 62, 80))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Day 7 Learning Guide")
set_font(run2, bold=True, size=20, color=(52, 152, 219))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Polish, Tests, and Demo Rehearsal — Finishing a Consulting-Ready System")
set_font(run3, italic=True, size=13, color=(127, 140, 141))

doc.add_paragraph()
divider()
doc.add_paragraph()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run(
    "This guide covers the two final tests, assert_called_once() mechanics,\n"
    "README structure, cold start rehearsal, and the 3-minute demo script."
)
set_font(run4, size=11, color=(52, 73, 94))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — WHAT WE BUILT
# ══════════════════════════════════════════════════════════════════════════════

h1("1. What Day 7 Built")

body(
    "Days 1–6 built a complete working system. Day 7 was not about adding features — "
    "it was about making the system trustworthy: additional tests that verify "
    "behavior under stress, a README that lets anyone run it from scratch, "
    "and a rehearsed demo that runs in under 3 minutes."
)

body("Three deliverables:")
bullet("tests/test_pipeline.py — 2 new tests: stage call verification and LLM failure resilience")
bullet("README.md — 5-minute setup guide, demo flow, architecture overview, consulting offer")
bullet("Cold start rehearsal — delete DB → demo_reset → full UI flow verified end-to-end")

body("Final state: 21 tests, all green, no API key required for any test.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — TEST: STAGES CALLED IN ORDER
# ══════════════════════════════════════════════════════════════════════════════

h1("2. Test: test_pipeline_stages_called_in_order")

body(
    "This test mocks every single function that pipeline.py calls and then asserts "
    "that each one was called exactly once. It is the most complete test of the "
    "orchestrator's structure."
)

h2("2.1 The Full Test")

code_block([
    "def test_pipeline_stages_called_in_order(sample_csv, mocker):",
    "    mock_ingest = mocker.patch('src.pipeline.load_csv_to_bronze', return_value=(15, 0))",
    "    mock_silver = mocker.patch('src.pipeline.bronze_to_silver',   return_value=15)",
    "    mock_gold   = mocker.patch('src.pipeline.silver_to_gold',     return_value=15)",
    "    mock_kpi    = mocker.patch('src.pipeline.compute_kpi_summary', return_value={",
    "        'total_projects': 15, 'pct_on_track': 40.0, ...",
    "    })",
    "    mock_prompt = mocker.patch('src.pipeline.build_prompt',    return_value='test prompt')",
    "    mock_llm    = mocker.patch('src.pipeline.generate_report', return_value='## Executive Summary\\nTest.')",
    "",
    "    result = run_pipeline(sample_csv)",
    "",
    "    mock_ingest.assert_called_once()",
    "    mock_silver.assert_called_once()",
    "    mock_gold.assert_called_once()",
    "    mock_kpi.assert_called_once()",
    "    mock_prompt.assert_called_once()",
    "    mock_llm.assert_called_once()",
    "    assert result['kpi_summary']['total_projects'] == 15",
])

h2("2.2 Why Mock All 6 Functions?")

body(
    "Previous pipeline tests only mocked generate_report (the LLM call) because that "
    "was the only function that required an API key. This test goes further: "
    "it replaces every function with a mock. The result is a pure unit test of "
    "the orchestrator's logic — completely isolated from all data and I/O."
)

add_table(
    ["What gets mocked", "Why it's mocked here", "What the mock returns"],
    [
        ["load_csv_to_bronze", "Avoid actual file I/O and DB writes",       "(15, 0) — 15 loaded, 0 skipped"],
        ["bronze_to_silver",   "Avoid actual SQL transforms",               "15 — row count"],
        ["silver_to_gold",     "Avoid actual SQL transforms",               "15 — row count"],
        ["compute_kpi_summary","Avoid actual DB queries",                   "Full KPI dict with known values"],
        ["build_prompt",       "Avoid building a real prompt string",       "'test prompt' — placeholder"],
        ["generate_report",    "Avoid API call",                            "5-section mock report string"],
    ],
    col_widths=[2.0, 2.3, 2.3],
)

body(
    "When all six are mocked, run_pipeline() executes only the orchestration logic itself: "
    "calling functions in sequence, passing return values to the next stage, and assembling "
    "the final dict. That is exactly what we want to test."
)

h2("2.3 assert_called_once() vs assert_called_once_with()")

code_block([
    "# assert_called_once() — was it called exactly one time? (any arguments)",
    "mock_silver.assert_called_once()",
    "",
    "# assert_called_once_with(...) — was it called once WITH these specific arguments?",
    "mock_ingest.assert_called_once_with(sample_csv)",
    "",
    "# assert_called() — was it called at least once? (risky — misses double-calls)",
    "# assert_not_called() — was it never called?",
    "# call_count — integer, how many times it was called",
    "# assert mock_llm.call_count == 1",
])

body(
    "We use assert_called_once() (without arguments) for the transform functions "
    "because we care that they ran, not what arguments they received — their arguments "
    "are internal to the pipeline. We could use assert_called_once_with() for load_csv_to_bronze "
    "to verify the correct CSV path was passed, but the test's goal here is structural "
    "coverage, not argument verification."
)

callout(
    "What this test catches:",
    "If someone accidentally deletes a stage from pipeline.py "
    "(e.g. removes the silver_to_gold() call), this test fails immediately. "
    "Without this test, the pipeline would silently skip a stage and produce wrong results "
    "that only become visible when comparing output numbers.",
    bg="FCE4EC",
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — TEST: SURVIVES LLM FAILURE
# ══════════════════════════════════════════════════════════════════════════════

h1("3. Test: test_pipeline_survives_llm_failure")

code_block([
    "def test_pipeline_survives_llm_failure(sample_csv, mocker):",
    "    mocker.patch(",
    "        'src.pipeline.generate_report',",
    "        return_value='[Report generation unavailable. Please check API key.]'",
    "    )",
    "    result = run_pipeline(sample_csv)",
    "",
    "    assert 'kpi_summary' in result",
    "    assert 'ai_report' in result",
    "    assert result['kpi_summary']['total_projects'] == 15",
    "    assert 'unavailable' in result['ai_report'].lower()",
])

h2("3.1 What It Verifies")

body(
    "This test simulates the fallback scenario: generate_report() returns the "
    "fallback message string instead of a real report. The test verifies four things:"
)
bullet("The pipeline does not raise an exception — it completes normally")
bullet("The return dict still contains both expected keys")
bullet("The KPI data is correct — the data pipeline ran successfully despite AI failure")
bullet("The fallback string is in ai_report — the failure was captured, not silently swallowed")

h2("3.2 Why This Matters for a Demo")

body(
    "During a live client demo, internet can drop, the API key can expire, "
    "or Anthropic can have a brief outage. Without this behavior, the demo crashes "
    "mid-presentation with an unhandled exception and a stack trace. "
    "With graceful degradation, the KPI cards and project table remain visible — "
    "only the AI text section shows the fallback message. "
    "You can continue the demo manually and regenerate the report once connectivity is restored."
)

callout(
    "Design principle tested here:",
    "A system that degrades gracefully under failure is more valuable than a system "
    "that is perfect when everything works. Real enterprise software earns trust "
    "by handling failures cleanly, not by pretending failures won't happen.",
    bg="E8F5E9",
)

h2("3.3 Testing Graceful Degradation vs Testing Success")

add_table(
    ["Test type", "What it checks", "Example"],
    [
        ["Happy path",      "Does it work when everything is correct?",        "test_pipeline_returns_expected_keys"],
        ["Edge case",       "Does it work with unusual but valid input?",       "test_empty_project_name_rows_skipped"],
        ["Failure path",    "Does it fail gracefully when something goes wrong?","test_pipeline_survives_llm_failure"],
        ["Regression",      "Does it still work after a change?",              "All tests re-run after every edit"],
        ["Contract",        "Does it call its dependencies as expected?",       "test_pipeline_stages_called_in_order"],
    ],
    col_widths=[1.5, 2.8, 2.3],
)

body(
    "A mature test suite has all five types. Days 1–6 covered happy paths and edge cases. "
    "Day 7 added failure path and contract tests, completing the suite."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — README.md
# ══════════════════════════════════════════════════════════════════════════════

h1("4. README.md — The First Thing Anyone Reads")

body(
    "A README is the front door of your project. "
    "Anyone who receives your code — a client's IT team, a future collaborator, "
    "or yourself six months from now — will read the README first. "
    "A good README answers four questions in under two minutes: "
    "what does this do, how do I run it, how do I use it, and why does it matter."
)

h2("4.1 Structure of This README")

add_table(
    ["Section", "Purpose", "Why it's there"],
    [
        ["Tagline",           "One sentence: 'Turn raw project data into AI reports'",     "Hooks the reader immediately"],
        ["Consulting quote",  "'I help organizations automate reporting...'",               "Sets context as a product pitch"],
        ["What it does",      "6-step numbered list of the pipeline",                       "Shows scope without overwhelming detail"],
        ["Prerequisites",     "Python 3.11+, Anthropic API key",                           "Prevents wasted setup time"],
        ["Setup (5 minutes)", "3 commands: install, configure, reset",                     "The actual getting-started path"],
        ["Run the demo",      "1 command + 4-step demo flow",                              "The thing the client actually sees"],
        ["Reset",             "1 command for pre-meeting cleanup",                          "Reliability guarantee"],
        ["Tests",             "1 command to verify everything works",                       "Trust signal for technical clients"],
        ["Architecture",      "ASCII diagram + link to design.md",                          "Technical credibility"],
        ["Consulting offer",  "Price, timeline, deliverables",                              "The call to action"],
    ],
    col_widths=[1.7, 2.4, 2.5],
)

h2("4.2 The Consulting Offer Section")

code_block([
    "## Consulting offer",
    "",
    "**AI Reporting Automation Pilot**",
    "- Price: $5,000–$10,000",
    "- Timeline: 30 days",
    "- Deliverables: data pipeline, KPI model, AI report, dashboard, executive demo",
])

body(
    "Putting the consulting offer directly in the README makes a deliberate statement: "
    "this is not a hobby project, it is a product. When you show this repository to a client "
    "or share a link, they immediately see the commercial framing. "
    "The price range is intentionally wide to anchor high while leaving room to negotiate."
)

h2("4.3 Why '5 Minutes' in Setup")

body(
    "The setup section header says '5 minutes' for a reason. "
    "Every minute of setup friction reduces the chance someone actually runs it. "
    "Three commands — pip install, cp .env.example, demo_reset.py — is achievable in 5 minutes "
    "on any machine with Python. If setup takes longer, something is wrong with the instructions."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — COLD START REHEARSAL
# ══════════════════════════════════════════════════════════════════════════════

h1("5. The Cold Start Rehearsal")

body(
    "A cold start test simulates what happens when someone runs your system for the very first time "
    "on a machine that has never seen it before — no database, no cached data, nothing pre-populated."
)

h2("5.1 The Procedure")

code_block([
    "# Step 1: Simulate a fresh machine",
    "rm -f data/db/reporting.duckdb",
    "",
    "# Step 2: Reset (should complete in < 30 seconds)",
    "python demo/demo_reset.py",
    "",
    "# Expected output:",
    "# Resetting demo environment...",
    "#   Database cleared",
    "#   Sample data written → data/input/project_data.csv",
    "#   Pipeline complete — 15 projects loaded",
    "#",
    "# Reset complete in 0.9s",
    "",
    "# Step 3: Start the UI",
    "streamlit run app/main.py",
])

h2("5.2 The Manual Verification Checklist")

add_table(
    ["Step", "Action", "Pass condition"],
    [
        ["1", "Page loads",                       "Empty state shown: 'Use the sidebar...'"],
        ["2", "Click 'Load Sample Data'",          "Spinner, then KPI cards: Total = 15"],
        ["3", "Check At Risk / Delayed card",      "Value > 0% (not all projects on track)"],
        ["4", "Scroll project table",              "15 rows, checkmarks in is_late / is_over_budget"],
        ["5", "Click 'Generate Report'",           "Spinner, then 5-section report appears"],
        ["6", "Report mentions a real project",    "'Mobile App Launch' or 'Financial System Integration' visible"],
        ["7", "Click 'Export PDF'",                "Download dialog appears"],
        ["8", "Open the PDF",                      "5 sections, page numbers, company name"],
        ["9", "Close browser, stop with Ctrl+C",   "Terminal returns to prompt cleanly"],
    ],
    col_widths=[0.4, 2.4, 3.8],
)

body(
    "Total time for this checklist: under 3 minutes. "
    "If any step fails, find and fix the root cause before any client meeting. "
    "A broken demo is worse than no demo."
)

h2("5.3 Why Delete the Database?")

body(
    "Running demo_reset.py while a database already exists tests the reset, not the cold start. "
    "The database might have tables from a previous run with different data. "
    "Deleting it first guarantees you are testing the scenario a client's IT team would experience "
    "when they run the setup instructions on a new machine for the first time."
)

callout(
    "The reliability principle:",
    "If you cannot delete the database, run demo_reset.py, and have everything work perfectly "
    "in under 60 seconds — you are not ready to demo. Run this procedure the morning before "
    "every client meeting, not as an afterthought.",
    bg="E3F2FD",
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — THE 3-MINUTE DEMO SCRIPT
# ══════════════════════════════════════════════════════════════════════════════

h1("6. The 3-Minute Demo Script")

body(
    "A great demo is not showing features — it is telling a story about a problem being solved. "
    "Every sentence should either establish the pain or demonstrate the relief. "
    "Practice this script out loud until it feels natural."
)

h2("6.1 The Script")

code_block([
    "'Most organizations waste hours every week manually gathering data",
    " and building reports in PowerPoint.'",
    "",
    "[pause — let that land]",
    "",
    "'What you're seeing here is a system that does that automatically.",
    " I connect to your raw data — project files, spreadsheets, whatever you have —",
    " and in under a minute it cleans the data, calculates your KPIs,",
    " and uses AI to write this executive summary.'",
    "",
    "[Click 'Load Sample Data' — KPI cards appear]",
    "",
    "'You can see the portfolio health immediately: 15 projects, 40% delayed,",
    " $183,000 in budget overruns. No one had to build a spreadsheet for this.'",
    "",
    "[Scroll to the project table]",
    "",
    "'The table shows every project — which are late, which are over budget.",
    " Click any column to sort. Your leadership team can drill in themselves.'",
    "",
    "[Click 'Generate Report']",
    "",
    "'Now the AI writes the executive summary — using your actual data,",
    " your actual project names, your actual numbers.",
    " It names the projects at risk, gives the real budget figures,",
    " and tells you exactly where to focus. No manual work. No copy-paste errors.'",
    "",
    "[Wait for report — point out a real project name in the output]",
    "",
    "[Click 'Export PDF']",
    "",
    "'And it exports a ready-to-send PDF — this is what you would email",
    " to your leadership team on Monday morning. Automatically.'",
    "",
    "[pause]",
    "",
    "'We can have a version of this running on your actual data in 30 days.'",
])

h2("6.2 Anatomy of the Script")

add_table(
    ["Line / action", "Purpose"],
    [
        ["'Most organizations waste hours...'",    "Open with the pain — not your solution"],
        ["[pause — let that land]",                "Give them time to recognize themselves in the problem"],
        ["'I connect to your raw data...'",        "Make it concrete: 'your data', not 'any data'"],
        ["[Load sample data]",                     "Action while talking — keeps attention"],
        ["'You can see the portfolio health...'",  "Narrate what they're seeing — don't stay silent"],
        ["[Click Generate Report]",                "Build tension: something is about to happen"],
        ["'using your actual data, your actual...'","Repeat 'actual' to reinforce anti-hallucination"],
        ["[Point to real project name]",           "The 'aha' moment — AI mentioned a real name"],
        ["[Export PDF]",                           "Show the deliverable — something they can keep"],
        ["'30 days' close",                        "Concrete, time-bound call to action"],
    ],
    col_widths=[2.8, 3.8],
)

body(
    "Notice what is absent from the script: no mention of DuckDB, no mention of medallion architecture, "
    "no mention of Python or Streamlit. The client does not care about the technology. "
    "They care about the outcome: less manual work, better decisions, faster reporting."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — COMPLETE SYSTEM SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

h1("7. The Complete System — All 7 Days at a Glance")

h2("7.1 The Full Data Flow")

code_block([
    "CSV file (15 projects, 10 columns)",
    "  ↓  load_csv_to_bronze()       Day 1  → bronze_project_raw  (raw strings)",
    "  ↓  bronze_to_silver()         Day 2  → silver_project_clean (typed, normalized)",
    "  ↓  silver_to_gold()           Day 2  → gold_project_kpi     (+ 5 KPI columns)",
    "  ↓  compute_kpi_summary()      Day 3  → KPISummary dict       (portfolio stats)",
    "  ↓  build_prompt()             Day 4  → grounded prompt string",
    "  ↓  generate_report()          Day 4  → 5-section AI report",
    "  ↓  Streamlit UI               Day 5  → browser dashboard",
    "  ↓  generate_pdf()             Day 6  → 3-page PDF file",
    "  ↓  demo_reset.py              Day 6  → clean state in < 1s",
])

h2("7.2 All Files Created")

add_table(
    ["File", "Day", "Purpose"],
    [
        ["src/storage/db.py",                "1", "Single DuckDB connection factory"],
        ["src/ingestion/csv_loader.py",       "1", "CSV → bronze_project_raw"],
        ["demo/sample_data.py",              "1", "15-project dataset + write_sample_csv()"],
        ["src/transforms/bronze_to_silver.py","2", "Raw strings → clean types, normalized values"],
        ["src/transforms/silver_to_gold.py",  "2", "Clean data → 5 computed KPI columns"],
        ["src/transforms/kpi_aggregator.py",  "3", "Gold table → portfolio KPI dict (8 SQL queries)"],
        ["src/pipeline.py",                  "3", "Orchestrator: runs all 5 stages in order"],
        ["conftest.py",                      "3", "sys.path fix so pytest finds src package"],
        ["src/llm/prompt_builder.py",        "4", "KPI dict → anti-hallucination prompt string"],
        ["src/llm/claude_client.py",         "4", "Anthropic SDK wrapper with graceful fallback"],
        ["app/components/kpi_cards.py",      "5", "Four st.metric() cards"],
        ["app/components/project_table.py",  "5", "st.dataframe() with column config"],
        ["app/components/ai_summary.py",     "5", "Generate button + markdown display"],
        ["app/main.py",                      "5", "Streamlit entry point, session state, layout"],
        ["src/reporting/pdf_generator.py",   "6", "ReportLab 3-page PDF with 5 content sections"],
        ["demo/demo_reset.py",               "6", "Pre-demo cleanup: drop tables, regenerate data"],
        ["README.md",                        "7", "5-minute setup guide + consulting offer"],
    ],
    col_widths=[2.8, 0.5, 3.3],
)

h2("7.3 All 21 Tests")

add_table(
    ["Test file", "Tests", "What they cover"],
    [
        ["test_ingestion.py",       "3", "CSV loading: happy path, missing columns, blank rows"],
        ["test_transforms.py",      "6", "Silver: date type, status normalization, null handling; Gold: budget math, is_late, on-time boundary"],
        ["test_kpi_aggregator.py",  "4", "Portfolio totals, pct sums to 100, top-3 length, dept count"],
        ["test_prompt_builder.py",  "3", "No {} placeholders, project names present, 5 headers present"],
        ["test_pipeline.py",        "5", "Expected keys, fast without LLM, mocked LLM, stage call verification, LLM failure resilience"],
    ],
    col_widths=[2.3, 0.7, 3.6],
)

h2("7.4 Performance Characteristics")

add_table(
    ["Stage", "Typical time", "Bottleneck"],
    [
        ["Stages 1–4 (data pipeline)",    "0.05–0.10s",   "DuckDB SQL — extremely fast"],
        ["Stage 5 without LLM (mocked)",  "< 0.01s",      "No I/O"],
        ["Stage 5 with real LLM",         "3–8s",         "Anthropic API round-trip"],
        ["PDF generation",                "0.2–0.5s",     "ReportLab rendering"],
        ["Demo reset (cold start)",       "0.8–1.2s",     "DuckDB file creation"],
        ["pytest tests/ -v (21 tests)",   "2–3s",         "DuckDB in-memory operations"],
    ],
    col_widths=[2.5, 1.4, 2.7],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — KEY CONCEPTS
# ══════════════════════════════════════════════════════════════════════════════

h1("8. Key Concepts to Remember")

add_table(
    ["Concept", "What it means", "Why it matters"],
    [
        ["assert_called_once()",    "Mock was invoked exactly once, any args",              "Verifies orchestrator calls every stage"],
        ["assert_called_once_with()","Mock was invoked once with specific args",            "Verifies correct arguments passed to dependency"],
        ["Contract test",           "Verify a function calls its dependencies correctly",   "Catches accidental stage removal from pipeline"],
        ["Failure path test",       "Verify graceful behavior when a dependency fails",     "Prevents crashes during live demos"],
        ["Cold start test",         "Delete all state, run from scratch, verify result",    "The ultimate reliability check"],
        ["README as pitch deck",    "Front page doubles as consulting proposal",             "Technical and commercial audiences both served"],
        ["30-day pilot framing",    "Specific time-bound deliverable in README",             "Concrete enough to trigger a buying decision"],
        ["Demo script structure",   "Open with pain, show relief, close with offer",         "Keeps attention, creates narrative arc"],
        ["'Actual data' repetition","Reinforce that AI uses real numbers, not invented",     "Counteracts hallucination skepticism"],
        ["3-minute rule",           "If it takes longer than 3 min, something is wrong",    "Client attention fades after 3 minutes"],
    ],
    col_widths=[2.0, 2.3, 2.3],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — THE COMPLETE CHECKLIST
# ══════════════════════════════════════════════════════════════════════════════

h1("9. The 7-Day Completion Checklist")

code_block([
    "Day 1 ✓  requirements.txt, .env, .gitignore, db.py, csv_loader.py, sample_data.py",
    "         3 tests passing",
    "",
    "Day 2 ✓  bronze_to_silver.py, silver_to_gold.py",
    "         9 tests passing",
    "         Bugs fixed: TRY_STRPTIME→DATE cast, keep_default_na=False",
    "",
    "Day 3 ✓  kpi_aggregator.py, pipeline.py (Stage 5 placeholder), conftest.py",
    "         15 tests passing",
    "         Bug fixed: patch where used not where defined",
    "",
    "Day 4 ✓  prompt_builder.py, claude_client.py, pipeline.py (Stage 5 wired)",
    "         19 tests passing",
    "",
    "Day 5 ✓  kpi_cards.py, project_table.py, ai_summary.py, app/main.py",
    "         19 tests passing (UI not tested by pytest — verified manually)",
    "",
    "Day 6 ✓  pdf_generator.py, demo_reset.py",
    "         19 tests passing",
    "         PDF generates in < 0.5s, demo reset runs in < 1s",
    "",
    "Day 7 ✓  README.md, 2 new pipeline tests",
    "         21 tests passing",
    "         Cold start rehearsal complete",
    "         3-minute demo script memorized",
])

callout(
    "What comes next:",
    "Record a 3-minute demo video. This is the single highest-leverage activity you can do now. "
    "A video lets you share the demo asynchronously — prospects can watch it before a call, "
    "forward it to their team, or replay it while writing up a budget request. "
    "Record it with Loom or QuickTime, share the link, and put it at the top of your README.",
    bg="E8F5E9",
)

divider()

p_final = doc.add_paragraph()
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_final.paragraph_format.space_before = Pt(20)
run_final = p_final.add_run("Day 7 complete — 21 tests passing, cold start verified, demo ready.")
set_font(run_final, bold=True, size=12, color=(39, 174, 96))

# ── Save ───────────────────────────────────────────────────────────────────────
out_path = "docs/day-7-learning-guide.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
