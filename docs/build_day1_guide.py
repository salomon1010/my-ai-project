"""Run this script to generate day-1-learning-guide.docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Style helpers ──────────────────────────────────────────────────────────────

def set_font(run, bold=False, italic=False, size=None, color=None, font_name=None):
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, bold=True, size=20, color=(44, 62, 80))
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, bold=True, size=15, color=(44, 62, 80))
    return p


def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, bold=True, size=12, color=(52, 73, 94))
    return p


def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def code_block(lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 30, 30)
        if i < len(lines) - 1:
            run.add_break()
    doc.add_paragraph()


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def callout(label, text, bg="FFF3CD"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg)
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    r1 = p.add_run(f"{label}  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    doc.add_paragraph()


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "2C3E50")
        tcPr.append(shd)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        fill = "F8F9FA" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
    if col_widths:
        for ri, row in enumerate(t.rows):
            for ci, width in enumerate(col_widths):
                row.cells[ci].width = Inches(width)
    doc.add_paragraph()
    return t


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Cover page ─────────────────────────────────────────────────────────────────

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(40)
p.paragraph_format.space_after = Pt(8)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AI Executive Reporting System")
set_font(run, bold=True, size=28, color=(44, 62, 80))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(4)
run2 = p2.add_run("Day 1 Learning Guide")
set_font(run2, bold=True, size=18, color=(52, 152, 219))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(30)
run3 = p3.add_run("Foundation: Database, CSV Ingestion & Sample Data")
set_font(run3, italic=True, size=13, color=(100, 100, 100))

callout("What this guide covers:",
        "Every file, every decision, and every line of code built on Day 1. "
        "Read this to understand not just WHAT was built, but WHY — so you can "
        "explain it to a client, debug it yourself, and extend it later.",
        bg="D6EAF8")

doc.add_page_break()

# ── Section 1: The Big Picture ─────────────────────────────────────────────────

h1("The Big Picture First")
body("Before touching any code, you need to understand what you're building and why "
     "it's structured this way.")

h2("What the system does")
code_block([
    "Your messy CSV file",
    "        ↓",
    "    Python reads it",
    "        ↓",
    "    Saves it to DuckDB (a database that lives in a single file)",
    "        ↓",
    "    Later: cleans it, calculates KPIs, sends to AI, generates report",
])
body("Day 1 only builds the first two arrows: read the CSV → save to database. "
     "Everything else (cleaning, AI, reports) comes in later days. This is intentional — "
     "you build one solid layer at a time.")

h2("Why a database instead of just keeping data in Python memory?")
body("You could just read the CSV and keep it in a Python variable. But databases have "
     "real advantages for a multi-component system:")

add_table(
    ["Memory (Python variable)", "Database (DuckDB)"],
    [
        ["Gone when script stops", "Persists on disk"],
        ["Hard to query with SQL", "Full SQL support"],
        ["One script can't share with another",
         "Streamlit app, reset script, and tests can all read the same data"],
        ["No history of what came in", "You can inspect it anytime"],
    ],
    col_widths=[2.8, 3.5],
)
body("Think of the database as a shared filing cabinet. Every part of your system — "
     "the pipeline, the UI, the tests — goes to the same cabinet to get data.")

h2("Why DuckDB specifically?")
body("DuckDB is what's called an 'embedded database.' It's just a single .duckdb file "
     "on your computer — no server to install, no network connection needed. It's like "
     "SQLite but built for analytics (fast SQL, good Pandas support).")
callout("Demo advantage:",
        "A client doesn't need to set up anything. One file, everything works. "
        "Perfect for a laptop demo.",
        bg="D5F5E3")

divider()

# ── Section 2: Files Created ───────────────────────────────────────────────────

h1("The Files You Created")
body("Here is every file and why it exists:")

code_block([
    "my-ai-project/",
    "├── requirements.txt          ← 'Shopping list' of Python packages",
    "├── .env.example              ← Template showing what secrets are needed",
    "├── .gitignore                ← What NOT to commit to git (secrets, temp files)",
    "├── data/",
    "│   ├── input/",
    "│   │   └── project_data.csv  ← The demo dataset (15 fake projects)",
    "│   └── db/                   ← DuckDB database file lives here",
    "├── src/",
    "│   ├── __init__.py           ← Tells Python 'this folder is a package'",
    "│   ├── storage/",
    "│   │   ├── __init__.py",
    "│   │   └── db.py             ← The ONE place DuckDB connection is made",
    "│   ├── ingestion/",
    "│   │   ├── __init__.py",
    "│   │   └── csv_loader.py     ← Reads CSV → writes to bronze table",
    "│   ├── transforms/           ← Empty today, used Day 2+",
    "│   ├── llm/                  ← Empty today, used Day 4+",
    "│   └── reporting/            ← Empty today, used Day 6+",
    "├── app/                      ← Empty today, Streamlit UI lives here Day 5+",
    "├── demo/",
    "│   ├── __init__.py",
    "│   └── sample_data.py        ← The 15 fake projects as Python data",
    "└── tests/",
    "    ├── __init__.py",
    "    └── test_ingestion.py     ← Automated tests for csv_loader.py",
])

divider()

# ── Section 3: requirements.txt ───────────────────────────────────────────────

h1("File 1: requirements.txt")

h2("What it is")
body("A plain text file that lists every Python package your project needs. "
     "Think of it as a shopping list for your Python environment.")

h2("The file")
code_block([
    "duckdb==0.10.3",
    "pandas==2.2.2",
    "anthropic==0.28.0",
    "streamlit==1.35.0",
    "reportlab==4.2.2",
    "python-dotenv==1.0.1",
    "loguru==0.7.2",
    "pytest==8.2.2",
    "pytest-mock==3.14.0",
])

h2("What each package does")
add_table(
    ["Package", "What it does", "Used when"],
    [
        ["duckdb", "The database engine (runs SQL on local files)", "Day 1+"],
        ["pandas", "Reads CSV files, manipulates data in Python", "Day 1+"],
        ["anthropic", "Official SDK to call Claude AI", "Day 4"],
        ["streamlit", "Builds the web UI in pure Python", "Day 5"],
        ["reportlab", "Generates PDF files", "Day 6"],
        ["python-dotenv", "Reads your .env file for secrets", "Day 1+"],
        ["loguru", "Better logging (replaces Python's print for status messages)", "Day 1+"],
        ["pytest", "Runs your automated tests", "Day 1+"],
        ["pytest-mock", "Lets tests fake (mock) the Claude AI call so tests don't cost money", "Day 4+"],
    ],
    col_widths=[1.4, 3.2, 1.2],
)

h2("Why pin exact versions (==0.10.3 not just duckdb)?")
body("If you just write duckdb, pip installs the latest version. In 6 months that version "
     "might have breaking changes and your code breaks. Pinning ==0.10.3 guarantees the "
     "same version every time anyone installs it — on your laptop, a client's machine, "
     "or in the future.")

code_block(["pip install -r requirements.txt"])
body("This reads the file and installs all packages at once.")

divider()

# ── Section 4: .env.example ────────────────────────────────────────────────────

h1("File 2: .env.example")

h2("What it is")
body("A template showing what environment variables (secrets and settings) your app needs. "
     "The actual secrets go in .env — this file just shows the structure.")

h2("The file")
code_block([
    "ANTHROPIC_API_KEY=your_key_here",
    "MODEL_ID=claude-haiku-4-5",
    "DB_PATH=data/db/reporting.duckdb",
    "REPORTS_OUTPUT_DIR=outputs/reports",
    "COMPANY_NAME=Acme Corporation",
])

h2("What each variable does")
add_table(
    ["Variable", "Purpose"],
    [
        ["ANTHROPIC_API_KEY", "Your secret key to call Claude AI (never share this)"],
        ["MODEL_ID", "Which Claude model to use (claude-haiku-4-5 is fast and cheap)"],
        ["DB_PATH", "Where the DuckDB file lives on disk"],
        ["REPORTS_OUTPUT_DIR", "Where PDF reports get saved"],
        ["COMPANY_NAME", "Appears on the PDF cover page (change per client)"],
    ],
    col_widths=[2.0, 4.3],
)

h2("Why .env.example and not just .env?")
bullet(".env contains your real API key — it must NEVER be committed to git")
bullet(".env.example has fake placeholder values — it IS committed to git so other "
       "developers (or future you on a new computer) know what to fill in")

h2("The workflow")
code_block([
    "# Copy the template",
    "cp .env.example .env",
    "",
    "# Edit .env and add your real API key",
    "# (use any text editor — VS Code, nano, etc.)",
])
body("Then .gitignore makes sure .env is never accidentally committed.")

divider()

# ── Section 5: .gitignore ─────────────────────────────────────────────────────

h1("File 3: .gitignore")

h2("What it is")
body("Tells git which files to ignore — never track, never commit, never push.")

h2("The file")
code_block([
    ".env",
    "data/db/",
    "outputs/",
    "__pycache__/",
    ".pytest_cache/",
    "*.pyc",
    "*.pyo",
    ".DS_Store",
])

h2("Why each entry")
add_table(
    ["Entry", "Why ignored"],
    [
        [".env", "Contains real API keys — if this hits GitHub, someone steals your key"],
        ["data/db/", "The DuckDB file is generated — not source code, pointless to commit"],
        ["outputs/", "Generated PDF reports — not source code"],
        ["__pycache__/", "Python's compiled bytecode — auto-generated, clutters repo"],
        [".pytest_cache/", "Test run cache — auto-generated"],
        ["*.pyc, *.pyo", "Compiled Python files — auto-generated"],
        [".DS_Store", "macOS folder metadata — irrelevant to other developers"],
    ],
    col_widths=[1.6, 4.7],
)

callout("The rule:",
        "Commit source code and configuration templates. Never commit secrets, "
        "generated files, or OS noise.",
        bg="FADBD8")

divider()

# ── Section 6: project_data.csv ───────────────────────────────────────────────

h1("File 4: data/input/project_data.csv")

h2("What it is")
body("The fake dataset that powers all your demos. 15 projects across 3 departments, "
     "designed to have enough variety to make the AI report interesting.")

h2("How it was generated")
body("You created demo/sample_data.py with the data as a Python list, then ran:")
code_block(["python demo/sample_data.py"])
body("That script wrote the Python list to a CSV file. This way, the source of truth for "
     "the data is sample_data.py (Python code you can control), and the CSV is generated "
     "from it — never edited by hand.")

h2("The columns explained")
code_block([
    "project_name        - What the project is called",
    "department          - Which business unit owns it (IT, Finance, Operations)",
    "planned_cost        - Original budget ($)",
    "actual_cost         - What was actually spent ($)",
    "planned_finish_date - Original target completion date",
    "actual_finish_date  - Actual/forecast completion date",
    "status              - on_track | at_risk | delayed | completed",
    "risk_level          - low | medium | high | critical",
    "issue_description   - Free text notes about problems",
    "owner               - Project manager name",
])

h2("Why 15 projects with this specific mix?")
body("The AI report is only interesting if there's something to report. Your 15 projects "
     "were chosen to have:")

add_table(
    ["What", "Count", "Why it matters"],
    [
        ["Over budget (>10%)", "7 of 15", "Gives the AI budget concerns to write about"],
        ["Late projects", "8 of 15", "Gives the AI schedule delays to report"],
        ["Critical risk", "1 (Mobile App Launch)", "Gives the AI a dramatic story to tell"],
        ["Completed", "3 of 15", "So it's not all doom and gloom"],
        ["Departments", "3 (IT, Finance, Ops)", "Gives 'issues by department' analysis something to aggregate"],
    ],
    col_widths=[1.8, 1.8, 2.7],
)

callout("Demo insight:",
        "If all 15 projects were 'on track, on budget,' the AI would have nothing to say. "
        "Realistic messy data makes the demo compelling.",
        bg="D5F5E3")

divider()

# ── Section 7: db.py ──────────────────────────────────────────────────────────

h1("File 5: src/storage/db.py — The Connection Factory")

h2("What it is")
body("The single place in your entire codebase where DuckDB is opened. Every other file "
     "calls get_connection() instead of opening DuckDB directly.")

h2("The file")
code_block([
    "import os",
    "import duckdb",
    "from pathlib import Path",
    "from dotenv import load_dotenv",
    "",
    "load_dotenv()",
    "",
    "",
    "def get_connection() -> duckdb.DuckDBPyConnection:",
    "    db_path = os.getenv(\"DB_PATH\", \"data/db/reporting.duckdb\")",
    "    Path(db_path).parent.mkdir(parents=True, exist_ok=True)",
    "    return duckdb.connect(db_path)",
])

h2("Line by line explanation")

h3("import os")
body("Python's built-in module for reading environment variables and file paths. "
     "You use it to read DB_PATH from .env.")

h3("import duckdb")
body("The DuckDB library. This is what gives you the ability to run SQL against a local file.")

h3("from pathlib import Path")
body("Python's modern way to work with file paths. Path('data/db/').mkdir(parents=True, "
     "exist_ok=True) creates the directory if it doesn't exist, without crashing if it already does.")

h3("from dotenv import load_dotenv  +  load_dotenv()")
body("Reads your .env file and loads its contents into environment variables. This call at "
     "module level means: the moment this file is imported anywhere, your .env is loaded. "
     "You don't have to remember to call it from other files.")

h3("db_path = os.getenv(\"DB_PATH\", \"data/db/reporting.duckdb\")")
body("Read the DB_PATH environment variable. If it's not set, use the default path. "
     "This is the default value pattern — tests can override DB_PATH to use a temporary "
     "directory without changing any code.")

h3("Path(db_path).parent.mkdir(parents=True, exist_ok=True)")
body("Let's break this down:")
bullet("Path(db_path) — converts the string path to a Path object", level=0)
bullet(".parent — gets the directory containing the file (data/db/ from data/db/reporting.duckdb)", level=0)
bullet(".mkdir(parents=True, exist_ok=True) — creates that directory", level=0)
bullet("parents=True means 'create intermediate directories too' (like mkdir -p)", level=1)
bullet("exist_ok=True means 'don't crash if it already exists'", level=1)
body("Without this line, DuckDB would crash if data/db/ didn't exist yet.")

h3("return duckdb.connect(db_path)")
body("Opens (or creates) the DuckDB file and returns a connection object. "
     "If the file doesn't exist, DuckDB creates it.")

h2("Why this matters architecturally")
body("This is called the Single Responsibility Principle. Only one file knows where the "
     "database is. If you later want to change the file path, switch to Postgres, use an "
     "in-memory database for speed, or use a different database per environment — "
     "you change only this file. None of your transform files, aggregation files, or UI "
     "code need to change.")

h3("The wrong way (what NOT to do):")
code_block([
    "# In csv_loader.py — WRONG",
    "con = duckdb.connect(\"data/db/reporting.duckdb\")  # hardcoded path!",
    "",
    "# In bronze_to_silver.py — WRONG",
    "con = duckdb.connect(\"data/db/reporting.duckdb\")  # same hardcoded path again!",
])
body("If you ever need to change the path, you'd have to find and update every file. "
     "With db.py, you change one line, everything works.")

divider()

# ── Section 8: sample_data.py ─────────────────────────────────────────────────

h1("File 6: demo/sample_data.py — The Dataset")

h2("What it is")
body("The 15 fake projects stored as a Python list of dictionaries. It also includes a "
     "function to write that data to a CSV file.")

h2("Why store data as Python, not just a CSV?")
body("Two reasons:")
bullet("The reset script needs to regenerate the CSV from scratch. If you only have a CSV "
       "and someone edits it by hand, your demo is broken. With Python as the source of "
       "truth, you can always regenerate a clean CSV.")
bullet("Tests can import the data directly. Your test file imports SAMPLE_PROJECTS to "
       "build test CSV files — no file path needed.")

h2("The key parts")

h3("SAMPLE_PROJECTS list")
code_block([
    "SAMPLE_PROJECTS = [",
    "    {",
    "        \"project_name\": \"Cloud Migration Phase 1\",",
    "        \"department\": \"IT\",",
    "        \"planned_cost\": 250000,",
    "        \"actual_cost\": 310000,",
    "        ...",
    "    },",
    "    ...  # 14 more projects",
    "]",
])
body("This is a list of dictionaries. Each dictionary represents one project row. "
     "The keys match the CSV column headers exactly.")

h3("write_sample_csv function")
code_block([
    "def write_sample_csv(path: str = \"data/input/project_data.csv\") -> None:",
    "    Path(path).parent.mkdir(parents=True, exist_ok=True)",
    "    with open(path, \"w\", newline=\"\") as f:",
    "        writer = csv.DictWriter(f, fieldnames=SAMPLE_PROJECTS[0].keys())",
    "        writer.writeheader()",
    "        writer.writerows(SAMPLE_PROJECTS)",
])
bullet("csv.DictWriter — Python's built-in CSV writer that takes dictionaries")
bullet("fieldnames=SAMPLE_PROJECTS[0].keys() — gets the column names from the first project's keys")
bullet("writer.writeheader() — writes the header row (column names)")
bullet("writer.writerows(SAMPLE_PROJECTS) — writes all 15 project rows")

h3("if __name__ == '__main__': pattern")
code_block([
    "if __name__ == \"__main__\":",
    "    write_sample_csv()",
    "    print(\"Sample CSV written to data/input/project_data.csv\")",
])
body("This code only runs if you execute this file directly (python demo/sample_data.py). "
     "If another file imports sample_data, this block is skipped. "
     "It's the standard Python pattern for making a file both importable and runnable.")

divider()

# ── Section 9: csv_loader.py ──────────────────────────────────────────────────

h1("File 7: src/ingestion/csv_loader.py — Bronze Ingestion")

h2("What is medallion architecture?")
body("It's a data engineering pattern with three layers:")
code_block([
    "Bronze — Raw data, copied exactly as received (no changes)",
    "Silver — Cleaned and normalized (dates parsed, nulls handled, status normalized)",
    "Gold   — Business-ready KPIs (budget_variance, is_late, etc.)",
])
body("Why three layers instead of one? Because:")
bullet("Bronze gives you a recovery point. If your cleaning logic has a bug, you can "
       "re-clean from bronze without re-ingesting.")
bullet("Silver gives downstream users clean data without knowing the cleaning rules.")
bullet("Gold gives business users pre-computed metrics without knowing the calculation formulas.")

callout("Enterprise credibility:",
        "Companies like Databricks, Snowflake, and Delta Lake are built around this pattern. "
        "By naming your DuckDB tables bronze_project_raw, silver_project_clean, and "
        "gold_project_kpi, you're demonstrating the exact same concept — which impresses "
        "enterprise clients.",
        bg="D5F5E3")

h2("The file")
code_block([
    "import pandas as pd",
    "from loguru import logger",
    "from src.storage.db import get_connection",
    "",
    "",
    "def load_csv_to_bronze(file_path: str) -> tuple[int, int]:",
    "    df = pd.read_csv(file_path, dtype=str)",
    "",
    "    required = {\"project_name\", \"planned_cost\"}",
    "    if not required.issubset(df.columns):",
    "        missing = required - set(df.columns)",
    "        raise ValueError(f\"CSV is missing required columns: {missing}\")",
    "",
    "    initial_count = len(df)",
    "    df = df[df[\"project_name\"].notna() & (df[\"project_name\"].str.strip() != \"\")]",
    "    skipped = initial_count - len(df)",
    "",
    "    con = get_connection()",
    "    con.execute(\"CREATE OR REPLACE TABLE bronze_project_raw AS SELECT * FROM df\")",
    "    con.close()",
    "",
    "    loaded = len(df)",
    "    logger.info(f\"Loaded {loaded} rows into bronze_project_raw ({skipped} skipped)\")",
    "    return loaded, skipped",
])

h2("Line by line")

h3("df = pd.read_csv(file_path, dtype=str)")
body("Read the CSV into a pandas DataFrame. dtype=str is critical: it reads EVERY column "
     "as a string, no matter what the values look like.")
body("Why? Because in the bronze layer, you want raw data as-is. If pandas tries to "
     "auto-detect types, it might parse '250000' as integer 250000, '2024-03-31' as a "
     "datetime object using its own assumptions, or interpret 'N/A' as NaN. All of these "
     "'helpful' conversions happen before you've validated the data. Bronze = raw = strings only.")

h3("required = {\"project_name\", \"planned_cost\"}")
body("A Python set of column names that MUST exist. Using a set (with {}) instead of a "
     "list (with []) makes the issubset check efficient.")

h3("if not required.issubset(df.columns): raise ValueError(...)")
body("df.columns is the list of column names in the CSV. .issubset() checks if all "
     "required columns exist in that list. If any are missing, raise an error immediately. "
     "This is a boundary check — validate input at the entry point so errors are caught "
     "early with clear messages.")

h3("df = df[df[\"project_name\"].notna() & (df[\"project_name\"].str.strip() != \"\")]")
body("A pandas filter that keeps only rows where project_name is not null AND not "
     "just whitespace. Real client data often has blank rows at the bottom of a spreadsheet. "
     "This filter drops them silently.")

h3("con.execute(\"CREATE OR REPLACE TABLE bronze_project_raw AS SELECT * FROM df\")")
body("DuckDB magic: df is a pandas DataFrame, and DuckDB can reference it directly by name "
     "in a SQL statement. One line creates the table, infers all columns from the DataFrame, "
     "and populates it. CREATE OR REPLACE means: if the table already exists, drop it and "
     "recreate (safe for re-runs).")

h3("return loaded, skipped")
body("Returns a tuple. In Python, 'return 15, 0' is the same as 'return (15, 0)'. "
     "The caller can unpack it: loaded, skipped = load_csv_to_bronze(path).")

divider()

# ── Section 10: test_ingestion.py ─────────────────────────────────────────────

h1("File 8: tests/test_ingestion.py — Automated Tests")

h2("What tests are and why they matter")
body("A test is code that checks whether your other code works correctly. You run tests "
     "after making changes to make sure you didn't break anything.")

add_table(
    ["Without tests", "With tests"],
    [
        ["Manually click through the app after every change hoping nothing broke",
         "Run pytest and it tells you in 2 seconds if anything broke"],
        ["Bugs reach clients", "Bugs caught before they leave your machine"],
        ["'I think it works' confidence", "'Tests are green' confidence"],
    ],
    col_widths=[3.0, 3.3],
)

callout("Consulting value:",
        "For consulting work, tests are proof that your system is reliable. "
        "Clients with IT departments will ask: 'Do you have tests?' — and you can say yes.",
        bg="D5F5E3")

h2("Understanding pytest fixtures")
body("A fixture is a reusable piece of test setup. Instead of repeating setup code in "
     "every test, you define it once as a fixture and pytest injects it automatically.")

h3("Built-in: tmp_path")
body("A built-in pytest fixture that gives you a temporary directory unique to each test. "
     "Files created here are automatically deleted after the test. Your tests never "
     "interfere with each other or with your real data.")

h3("set_test_db (autouse=True)")
code_block([
    "@pytest.fixture(autouse=True)",
    "def set_test_db(tmp_path, monkeypatch):",
    "    monkeypatch.setenv(\"DB_PATH\", str(tmp_path / \"test.duckdb\"))",
])
body("autouse=True means this fixture runs automatically for EVERY test in the file — "
     "you don't have to request it. It overrides the DB_PATH environment variable to "
     "point to a fresh temporary DuckDB file. Each test gets a completely fresh, empty "
     "database and never writes to your real data/db/reporting.duckdb.")

h3("sample_csv fixture")
code_block([
    "@pytest.fixture",
    "def sample_csv(tmp_path):",
    "    from demo.sample_data import SAMPLE_PROJECTS",
    "    csv_path = tmp_path / \"test_projects.csv\"",
    "    with open(csv_path, \"w\", newline=\"\") as f:",
    "        writer = csv.DictWriter(f, fieldnames=SAMPLE_PROJECTS[0].keys())",
    "        writer.writeheader()",
    "        writer.writerows(SAMPLE_PROJECTS)",
    "    return str(csv_path)",
])
body("Creates a temporary CSV from SAMPLE_PROJECTS and returns its path. Any test that "
     "needs a ready-made CSV just declares sample_csv as a parameter and gets the path "
     "handed to it.")

h2("Test 1: The Happy Path")
code_block([
    "def test_load_happy_path(sample_csv):",
    "    from src.ingestion.csv_loader import load_csv_to_bronze",
    "    from src.storage.db import get_connection",
    "",
    "    loaded, skipped = load_csv_to_bronze(sample_csv)",
    "    assert loaded == 15",
    "    assert skipped == 0",
    "",
    "    con = get_connection()",
    "    count = con.execute(\"SELECT COUNT(*) FROM bronze_project_raw\").fetchone()[0]",
    "    con.close()",
    "    assert count == 15",
])
body("When you load a valid 15-row CSV, exactly 15 rows end up in the database. "
     "Notice two separate checks: the return value AND a SQL query against the database. "
     "The function could return 15 while secretly writing 0 rows — the SQL check catches that. "
     "fetchone() returns a tuple like (15,); [0] gets the count.")

h2("Test 2: Missing Required Column")
code_block([
    "def test_missing_required_column(tmp_path):",
    "    from src.ingestion.csv_loader import load_csv_to_bronze",
    "",
    "    bad_csv = tmp_path / \"bad.csv\"",
    "    bad_csv.write_text(\"project_name,department\\nProject A,IT\\n\")",
    "    with pytest.raises(ValueError, match=\"missing required columns\"):",
    "        load_csv_to_bronze(str(bad_csv))",
])
body("A CSV without planned_cost raises a clear error. pytest.raises() is a context "
     "manager that expects an exception — if no exception is raised, the test FAILS. "
     "This tests your boundary check: 'bad input produces a clear error, not a confusing crash.'")

h2("Test 3: Empty Project Names Skipped")
code_block([
    "def test_empty_project_name_rows_skipped(tmp_path):",
    "    # Build a CSV with 2 blank project_name rows and 1 valid row",
    "    ...",
    "    loaded, skipped = load_csv_to_bronze(str(csv_path))",
    "    assert loaded == 1",
    "    assert skipped == 2",
])
body("Rows with blank project_name are silently dropped and counted as skipped. "
     "This simulates real-world client data where someone exported a spreadsheet with "
     "blank trailing rows. Your system handles it gracefully.")

divider()

# ── Section 11: How It All Connects ───────────────────────────────────────────

h1("How It All Connects")

h2("Full execution flow")
code_block([
    "1. Someone calls: load_csv_to_bronze('data/input/project_data.csv')",
    "",
    "2. csv_loader.py reads the CSV with pandas",
    "   → All 15 rows, all columns as strings",
    "   → DataFrame has shape (15, 10)",
    "",
    "3. Validates required columns exist",
    "   → 'project_name' ✓  'planned_cost' ✓",
    "",
    "4. Filters out blank project names",
    "   → All 15 are valid, none dropped",
    "   → skipped = 0",
    "",
    "5. Calls get_connection() from db.py",
    "   → db.py reads DB_PATH from .env",
    "   → Creates data/db/ directory if needed",
    "   → Opens (or creates) data/db/reporting.duckdb",
    "   → Returns a DuckDB connection",
    "",
    "6. Executes SQL:",
    "   CREATE OR REPLACE TABLE bronze_project_raw AS SELECT * FROM df",
    "   → DuckDB reads the pandas DataFrame",
    "   → Creates a table with 15 rows, 10 columns, all strings",
    "",
    "7. Closes the connection",
    "   → File is written to disk",
    "",
    "8. Logs: 'Loaded 15 rows into bronze_project_raw (0 skipped)'",
    "",
    "9. Returns (15, 0)",
])

h2("How the tests work")
code_block([
    "1. pytest starts",
    "",
    "2. For each test, set_test_db runs automatically (autouse=True)",
    "   → DB_PATH is set to /tmp/pytest-xxx/test.duckdb (a fresh temp file)",
    "   → This means each test has a clean, empty database",
    "",
    "3. For tests that declare sample_csv:",
    "   → sample_csv fixture creates a temp CSV from SAMPLE_PROJECTS",
    "   → Returns the path to that CSV",
    "",
    "4. The test function runs",
    "   → Calls load_csv_to_bronze()",
    "   → Makes assertions about the results",
    "",
    "5. After the test:",
    "   → monkeypatch restores DB_PATH to original value",
    "   → tmp_path files are cleaned up",
    "   → No side effects remain",
])

divider()

# ── Section 12: What You Can Do Right Now ─────────────────────────────────────

h1("What You Can Do Right Now")

h2("1. Inspect the DuckDB table")
code_block([
    "python -c \"",
    "from src.storage.db import get_connection",
    "from dotenv import load_dotenv",
    "load_dotenv()",
    "",
    "con = get_connection()",
    "result = con.execute(",
    "    'SELECT project_name, department, planned_cost, actual_cost'",
    "    ' FROM bronze_project_raw LIMIT 5'",
    ").fetchdf()",
    "print(result)",
    "con.close()",
    "\"",
])
body("See the raw data exactly as it sits in the bronze table.")

h2("2. Check column types")
code_block([
    "python -c \"",
    "from src.storage.db import get_connection",
    "from dotenv import load_dotenv",
    "load_dotenv()",
    "",
    "con = get_connection()",
    "result = con.execute('DESCRIBE bronze_project_raw').fetchdf()",
    "print(result)",
    "con.close()",
    "\"",
])
body("You'll see all columns are VARCHAR (string) type. This confirms the bronze layer "
     "stores everything raw. Day 2 will convert them to proper types.")

h2("3. Count over-budget projects manually")
code_block([
    "python -c \"",
    "from src.storage.db import get_connection",
    "from dotenv import load_dotenv",
    "load_dotenv()",
    "",
    "con = get_connection()",
    "# Note: CAST needed because costs are stored as strings in bronze",
    "result = con.execute('''",
    "    SELECT project_name, planned_cost, actual_cost",
    "    FROM bronze_project_raw",
    "    WHERE CAST(actual_cost AS DOUBLE) > CAST(planned_cost AS DOUBLE)",
    "''').fetchdf()",
    "print(f'{len(result)} projects are over budget')",
    "print(result)",
    "con.close()",
    "\"",
])
body("A preview of what Day 2's gold layer will calculate automatically.")

divider()

# ── Section 13: Key Concepts ──────────────────────────────────────────────────

h1("Key Concepts to Remember")

add_table(
    ["Concept", "What it means", "Why it matters"],
    [
        ["Separation of concerns",
         "Each file does one thing: db.py connects, csv_loader.py loads, sample_data.py holds data",
         "Makes bugs easier to find and code easier to change"],
        ["Single connection point",
         "db.py is the only file that calls duckdb.connect(). Everything else calls get_connection()",
         "Makes the database location easy to change"],
        ["Bronze = raw",
         "The bronze layer stores data exactly as received — all strings, no transforms",
         "Your recovery point if anything goes wrong downstream"],
        ["dtype=str when reading CSV",
         "Never let pandas auto-detect types at the ingestion layer",
         "Parse types explicitly in silver where you control the logic"],
        ["Fixtures over setup methods",
         "pytest fixtures (tmp_path, monkeypatch, custom ones) are the modern way to set up test state",
         "autouse=True ensures isolation without repeating code"],
        ["Test the real thing",
         "test_load_happy_path checks the return value AND queries the database to verify",
         "Tests verify outcomes, not just function calls"],
    ],
    col_widths=[1.5, 2.5, 2.3],
)

divider()

# ── Section 14: Day 1 Checklist ───────────────────────────────────────────────

h1("Day 1 Completion Checklist")

add_table(
    ["Item", "Status"],
    [
        ["Folder structure created (src/, app/, demo/, tests/, data/)", "✓ Done"],
        ["Python packages installed (pip install -r requirements.txt)", "✓ Done"],
        ["Environment variable template (.env.example) created", "✓ Done"],
        [".env file created with real API key (done manually)", "✓ Done"],
        [".gitignore protecting secrets and generated files", "✓ Done"],
        ["db.py — single DuckDB connection factory", "✓ Done"],
        ["sample_data.py — 15 projects as Python data", "✓ Done"],
        ["project_data.csv — generated sample dataset", "✓ Done"],
        ["csv_loader.py — CSV → bronze_project_raw", "✓ Done"],
        ["test_ingestion.py — 3 tests, all passing", "✓ Done"],
    ],
    col_widths=[4.5, 1.2],
)

code_block(["pytest tests/test_ingestion.py -v    →  3 passed"])

divider()

# ── Section 15: What Day 2 Builds On This ─────────────────────────────────────

h1("What Day 2 Will Build On This")

body("Day 1 left bronze_project_raw with all strings. Day 2 will:")
bullet("Read bronze_project_raw")
bullet("Parse dates from strings to proper DATE type")
bullet("Cast costs from strings to DOUBLE type")
bullet("Normalize status values ('On Track' → 'on_track')")
bullet("Handle null values gracefully")
bullet("Write silver_project_clean — the cleaned version")
bullet("Then compute KPI columns and write gold_project_kpi")

callout("Why the separation matters:",
        "The cleaning intentionally happens in a separate step so you can inspect "
        "bronze_project_raw anytime and see exactly what came in — before any "
        "transformation touched it. This is your audit trail.",
        bg="D6EAF8")

# ── Save ──────────────────────────────────────────────────────────────────────

output_path = "docs/day-1-learning-guide.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
