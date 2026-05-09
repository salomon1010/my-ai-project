"""Run this script to generate day-2-learning-guide.docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Styles helpers ─────────────────────────────────────────────────────────────

def set_font(run, bold=False, italic=False, size=None, color=None, font_name=None):
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, bold=True, size=20, color=(44, 62, 80))
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, bold=True, size=15, color=(44, 62, 80))
    return p


def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, bold=True, size=12, color=(52, 73, 94))
    return p


def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def body_mixed(parts, indent=False):
    """parts = list of (text, bold, italic, mono)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    for text, bold, italic, mono in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(11)
        if mono:
            run.font.name = "Courier New"
            run.font.size = Pt(10)
    return p


def code_block(lines):
    """Render a code block as a shaded table."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 30, 30)
        if i < len(lines) - 1:
            run.add_break()
    doc.add_paragraph()  # spacing after block


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def callout(label, text, bg="FFF3CD", border="FFC107"):
    """Warning / tip box."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg)
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    r1 = p.add_run(f"{label}  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    doc.add_paragraph()


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "2C3E50")
        tcPr.append(shd)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    # Data rows
    for ri, row in enumerate(rows):
        fill = "F8F9FA" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("─" * 80)
    run.font.color.rgb = RGBColor(189, 195, 199)
    run.font.size = Pt(8)


# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run("AI Executive Reporting System")
set_font(run, bold=True, size=26, color=(44, 62, 80))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Day 2 Learning Guide")
set_font(run2, bold=True, size=20, color=(52, 152, 219))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Silver + Gold Transforms — Cleaning Data and Computing KPIs")
set_font(run3, italic=True, size=13, color=(127, 140, 141))

doc.add_paragraph()
divider()
doc.add_paragraph()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run(
    "This guide explains every file, every SQL statement, and every test built on Day 2.\n"
    "It also documents the two real bugs we hit — and exactly why they happened."
)
set_font(run4, size=11, color=(52, 73, 94))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — WHAT WE BUILT
# ══════════════════════════════════════════════════════════════════════════════

h1("1. What Day 2 Built")

body(
    "Day 1 gave you a bronze table with raw strings — data exactly as it came out of the CSV file. "
    "Day 2 added two more layers on top of that: a silver layer that cleans the data, and a gold layer "
    "that computes KPI columns used in the final report."
)

body("The pipeline after Day 2 looks like this:")

code_block([
    "bronze_project_raw      ← raw strings (Day 1)",
    "        ↓  bronze_to_silver.py",
    "silver_project_clean    ← clean types, normalized values",
    "        ↓  silver_to_gold.py",
    "gold_project_kpi        ← + budget_variance, is_late, is_over_budget, ...",
])

body("Three new files were created:")
bullet("src/transforms/bronze_to_silver.py — cleans the raw bronze data")
bullet("src/transforms/silver_to_gold.py   — calculates KPI columns")
bullet("tests/test_transforms.py            — 6 tests verifying all of the above")
body("Two bugs were found and fixed (documented in Section 5).")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — MEDALLION ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════

h1("2. The Medallion Architecture — Why Three Layers?")

body(
    "Medallion architecture is the industry standard pattern for organizing data pipelines. "
    "It was popularized by Databricks and Delta Lake, and is used by most enterprise data teams. "
    "By using this pattern in your demo, you speak the same language as data engineers at large companies."
)

add_table(
    ["Layer", "Table Name", "What it contains", "Who uses it"],
    [
        ["Bronze", "bronze_project_raw",   "Exact copy of the source — no transforms, all strings", "Recovery point if cleaning breaks"],
        ["Silver", "silver_project_clean", "Cleaned, typed, normalized — nulls handled, dates parsed", "Downstream transforms, analysts"],
        ["Gold",   "gold_project_kpi",     "Business-ready KPIs — budget_variance, is_late, etc.",   "Dashboards, AI reports, executives"],
    ],
    col_widths=[0.8, 1.8, 2.8, 2.2],
)

h3("Why not just transform everything in one step?")
body(
    "If you combine all three layers into one big transformation, you lose your recovery points. "
    "Imagine your cleaning logic has a bug that corrupts date parsing. With separate layers:"
)
bullet("Bronze is untouched — you can always re-clean from it")
bullet("Silver can be recomputed without re-fetching from the source")
bullet("Gold can be recomputed without re-running the cleaning step")

body(
    "In a real enterprise, source data might come from a live system you can only query once a day. "
    "The bronze layer is your insurance policy."
)

callout(
    "Consulting insight:",
    "When a client asks 'can you prove the numbers are correct?' — you open DuckDB, show them "
    "bronze_project_raw (their original data, unchanged), and trace every calculation through to "
    "gold. That traceability is what separates a real system from a dashboard that just shows numbers.",
    bg="E8F5E9",
    border="4CAF50",
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — bronze_to_silver.py
# ══════════════════════════════════════════════════════════════════════════════

h1("3. File: src/transforms/bronze_to_silver.py")

body(
    "This file reads the bronze table (raw strings) and writes a clean version to the silver table. "
    "All transformations are expressed as a single SQL statement — no Python loops, no row-by-row logic."
)

h2("3.1 The Full SQL Statement")

code_block([
    "CREATE OR REPLACE TABLE silver_project_clean AS",
    "SELECT",
    "    TRIM(project_name)                                    AS project_name,",
    "    TRIM(department)                                      AS department,",
    "    TRY_CAST(planned_cost AS DOUBLE)                     AS planned_cost,",
    "    TRY_CAST(actual_cost AS DOUBLE)                      AS actual_cost,",
    "    CAST(TRY_STRPTIME(planned_finish_date,'%Y-%m-%d') AS DATE)  AS planned_finish_date,",
    "    CAST(TRY_STRPTIME(actual_finish_date, '%Y-%m-%d') AS DATE)  AS actual_finish_date,",
    "    CASE LOWER(TRIM(status))",
    "        WHEN 'on track'  THEN 'on_track'",
    "        WHEN 'on_track'  THEN 'on_track'",
    "        WHEN 'green'     THEN 'on_track'",
    "        WHEN 'at risk'   THEN 'at_risk'",
    "        WHEN 'at_risk'   THEN 'at_risk'",
    "        WHEN 'amber'     THEN 'at_risk'",
    "        WHEN 'delayed'   THEN 'delayed'",
    "        WHEN 'red'       THEN 'delayed'",
    "        WHEN 'completed' THEN 'completed'",
    "        ELSE LOWER(TRIM(status))",
    "    END                                                    AS status,",
    "    LOWER(TRIM(risk_level))                               AS risk_level,",
    "    COALESCE(issue_description, '')                       AS issue_description,",
    "    COALESCE(owner, 'Unassigned')                         AS owner",
    "FROM bronze_project_raw",
    "WHERE project_name IS NOT NULL",
    "  AND planned_cost IS NOT NULL",
])

h2("3.2 Every Transformation Explained")

h3("TRIM — remove accidental whitespace")
body(
    "TRIM(project_name) removes leading and trailing spaces. Real client spreadsheets often have "
    "cells like '  Cloud Migration  ' (with spaces). Without TRIM, two projects with the same name "
    "would not match in joins or aggregations."
)
code_block(["TRIM('  Cloud Migration  ')  →  'Cloud Migration'"])

h3("TRY_CAST — convert strings to numbers safely")
body(
    "Bronze stores all values as strings (VARCHAR). To do math like actual_cost - planned_cost, "
    "the values must be numbers. CAST converts them. TRY_CAST (with TRY) returns NULL instead of "
    "crashing if a value cannot be converted."
)
code_block([
    "CAST('250000' AS DOUBLE)     →  250000.0   (works fine)",
    "TRY_CAST('N/A' AS DOUBLE)   →  NULL        (safe, no crash)",
    "CAST('N/A' AS DOUBLE)        →  ERROR       (crashes the whole query)",
])
callout(
    "Rule:",
    "Always use TRY_CAST at the silver layer. You do not know what garbage clients put in spreadsheets. "
    "A single bad cell with CAST would kill your entire pipeline.",
    bg="FFF3CD",
)

h3("TRY_STRPTIME + CAST AS DATE — parse date strings")
body(
    "TRY_STRPTIME parses a string into a datetime using a format pattern. '%Y-%m-%d' means "
    "four-digit year, two-digit month, two-digit day — matching '2024-03-31'."
)
body(
    "TRY_STRPTIME alone returns TIMESTAMP (date + time, e.g. '2024-03-31 00:00:00'). "
    "We wrap it in CAST(... AS DATE) to get a pure date with no time component. "
    "This matters for DATEDIFF in the gold layer — DATE arithmetic is cleaner than TIMESTAMP arithmetic."
)
code_block([
    "TRY_STRPTIME('2024-03-31', '%Y-%m-%d')           →  2024-03-31 00:00:00  (TIMESTAMP)",
    "CAST(TRY_STRPTIME('2024-03-31', '%Y-%m-%d') AS DATE)  →  2024-03-31           (DATE)",
    "TRY_STRPTIME('not-a-date', '%Y-%m-%d')           →  NULL   (safe, no crash)",
])

h3("CASE LOWER(TRIM(status)) — status normalization")
body(
    "Client data uses many spellings for the same status. A project manager might type 'On Track', "
    "another types 'on track', another uses a legacy system that says 'GREEN'. "
    "All mean the same thing. The CASE statement maps all variants to a controlled vocabulary."
)

add_table(
    ["Raw value in spreadsheet", "After normalization"],
    [
        ["'On Track'", "'on_track'"],
        ["'on track'", "'on_track'"],
        ["'GREEN'",    "'on_track'"],
        ["'At Risk'",  "'at_risk'"],
        ["'AMBER'",    "'at_risk'"],
        ["'Delayed'",  "'delayed'"],
        ["'RED'",      "'delayed'"],
        ["'Completed'","'completed'"],
    ],
    col_widths=[2.8, 2.8],
)

body(
    "LOWER(TRIM(status)) normalizes the input before the CASE runs, so you only need "
    "to handle lowercase variants in the WHEN clauses. The ELSE clause passes through "
    "anything unrecognized — you add new WHEN lines as you encounter new client systems."
)

h3("COALESCE — replace NULL with a default")
body(
    "COALESCE(value, default) returns the first non-NULL argument. "
    "It handles two cases: missing issue notes and missing owner names."
)
code_block([
    "COALESCE(NULL, '')            →  ''           (empty string, not NULL)",
    "COALESCE('vendor delay', '')  →  'vendor delay'  (real value kept)",
    "COALESCE(NULL, 'Unassigned')  →  'Unassigned'",
])
body(
    "Without COALESCE, NULL values would propagate through aggregations and cause "
    "the AI prompt builder to inject 'None' or 'null' into the executive report."
)

h3("WHERE clause — row-level filtering")
body(
    "Only rows with a non-null project_name AND planned_cost are included in silver. "
    "Rows that fail this check are silently excluded. This is acceptable because csv_loader.py "
    "already dropped rows with empty project_name — these WHERE conditions catch any edge cases "
    "that slipped through, like rows with a name but no budget at all."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — silver_to_gold.py
# ══════════════════════════════════════════════════════════════════════════════

h1("4. File: src/transforms/silver_to_gold.py")

body(
    "This file reads the clean silver table and adds five computed KPI columns. "
    "The SELECT * carries all existing columns forward — the five new columns are appended."
)

h2("4.1 The Full SQL Statement")

code_block([
    "CREATE OR REPLACE TABLE gold_project_kpi AS",
    "SELECT",
    "    *,",
    "    (actual_cost - planned_cost)                              AS budget_variance,",
    "    ROUND(((actual_cost - planned_cost) / planned_cost * 100.0), 2)  AS budget_variance_pct,",
    "    DATEDIFF('day', planned_finish_date, actual_finish_date)  AS schedule_delay_days,",
    "    CASE WHEN ((actual_cost - planned_cost) / planned_cost * 100.0) > 10",
    "         THEN TRUE ELSE FALSE END                             AS is_over_budget,",
    "    CASE WHEN actual_finish_date > planned_finish_date",
    "         THEN TRUE ELSE FALSE END                             AS is_late",
    "FROM silver_project_clean",
])

h2("4.2 Every KPI Column Explained")

h3("budget_variance — raw dollar difference")
body(
    "The simplest KPI: how many dollars over or under budget is this project?"
)
code_block([
    "budget_variance = actual_cost - planned_cost",
    "",
    "Example: planned $250,000, actual $310,000",
    "budget_variance = 310000 - 250000 = 60000",
    "",
    "Positive = over budget (bad)",
    "Negative = under budget (good)",
    "Zero     = exactly on budget",
])

h3("budget_variance_pct — percentage over/under budget")
body(
    "The dollar variance tells you the size of the problem. The percentage tells you the severity. "
    "A $60K overrun on a $250K project (24%) is worse than a $60K overrun on a $2M project (3%)."
)
code_block([
    "budget_variance_pct = (actual - planned) / planned × 100",
    "",
    "Example: planned $250,000, actual $310,000",
    "budget_variance_pct = (310000 - 250000) / 250000 × 100 = 24.0%",
    "",
    "ROUND(..., 2) keeps it to 2 decimal places: 24.00 not 24.000000001",
])

h3("schedule_delay_days — how many days late")
body(
    "DATEDIFF('day', start, end) counts calendar days between two dates. "
    "The argument order matters: (unit, start_date, end_date)."
)
code_block([
    "DATEDIFF('day', planned_finish_date, actual_finish_date)",
    "",
    "Planned: 2024-03-31,  Actual: 2024-05-15",
    "schedule_delay_days = DATEDIFF('day', 2024-03-31, 2024-05-15) = 45",
    "",
    "Positive = late (actual is after planned)",
    "Negative = early (actual is before planned)",
    "Zero     = on time",
])

h3("is_over_budget — boolean flag for >10% overrun")
body(
    "A simple TRUE/FALSE flag. The 10% threshold was defined in the requirements and "
    "agreed with the client concept. This flag powers filters like 'show me all over-budget projects' "
    "without requiring the dashboard to repeat the 10% logic."
)
code_block([
    "is_over_budget = TRUE  when budget_variance_pct > 10",
    "is_over_budget = FALSE otherwise",
    "",
    "Project at +9.9%:  is_over_budget = FALSE   (just under threshold)",
    "Project at +10.1%: is_over_budget = TRUE    (over threshold)",
])

h3("is_late — boolean flag for any delay")
body(
    "TRUE when the actual finish date is strictly after the planned finish date. "
    "A project that finished on the exact planned date is NOT late (is_late = FALSE)."
)
code_block([
    "is_late = TRUE  when actual_finish_date > planned_finish_date",
    "is_late = FALSE otherwise (on time or early)",
])

h2("4.3 Why SELECT * Instead of Listing Every Column?")
body(
    "Using SELECT * carries all silver columns forward automatically. "
    "If silver adds a new column later, gold gets it for free with no code change. "
    "The five KPI columns are appended after the *."
)
body(
    "This pattern works here because gold is always derived from silver and has no independent schema. "
    "In a production system with strict contracts, you would list columns explicitly. "
    "For this demo, SELECT * is the right tradeoff."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — BUGS AND FIXES
# ══════════════════════════════════════════════════════════════════════════════

h1("5. The Two Bugs We Hit — and Why They Happened")

body(
    "These were not mistakes in the design. They were real version-specific and library-specific "
    "behaviors that you can only discover by running the code. Understanding them makes you a "
    "better data engineer."
)

h2("Bug 1 — TRY_STRPTIME Returns TIMESTAMP, Not DATE")

h3("What happened")
body("The test expected the planned_finish_date column to have type DATE in the silver table.")
body("The test failed with:")
code_block(["AssertionError: Expected DATE type, got: TIMESTAMP"])

h3("Why it happened")
body(
    "In DuckDB 0.10.3, TRY_STRPTIME always returns a TIMESTAMP — a datetime value with both "
    "date and time components (e.g. '2024-03-31 00:00:00'). There is no version of TRY_STRPTIME "
    "that returns a pure DATE directly."
)
body(
    "DATE and TIMESTAMP are different types in DuckDB. TIMESTAMP stores '2024-03-31 00:00:00'. "
    "DATE stores '2024-03-31'. For date arithmetic (how many days between two dates), "
    "DATE is cleaner and less error-prone."
)

h3("The fix")
body("Wrap TRY_STRPTIME in an explicit CAST to DATE:")
code_block([
    "-- Before (returns TIMESTAMP):",
    "TRY_STRPTIME(planned_finish_date, '%Y-%m-%d')",
    "",
    "-- After (returns DATE):",
    "CAST(TRY_STRPTIME(planned_finish_date, '%Y-%m-%d') AS DATE)",
])

callout(
    "The lesson:",
    "Always check column types explicitly after transforms. The type in memory affects every "
    "downstream calculation. A TIMESTAMP date minus another TIMESTAMP gives an INTERVAL, "
    "while a DATE minus another DATE gives an INTEGER (days). DATEDIFF handles both, but "
    "explicit DATE types make intent clear to anyone reading the code.",
    bg="E3F2FD",
)

h2("Bug 2 — Empty CSV Cells Become NaN, Not Empty Strings")

h3("What happened")
body(
    "The test set issue_description to an empty string, wrote it to CSV, "
    "then ran the silver transform. The transform crashed with:"
)
code_block(["duckdb.ConversionException: Could not convert string '' to INT32"])

h3("Why it happened — step by step")

body("Step 1: The test wrote a CSV row with issue_description = '' (empty string).")
body("Step 2: pandas read the CSV with dtype=str.")
body("Step 3: Here is the problem. With dtype=str, pandas still converts empty CSV cells to NaN (not empty string).")

code_block([
    "# Empty cell in CSV:  project_name,planned_cost,...,issue_description,...",
    "#                     Cloud Migration,250000,...,                 ,...",
    "",
    "# What you expect:    df['issue_description'][0] == ''",
    "# What pandas gives:  df['issue_description'][0] == NaN  (float!)",
])

body(
    "NaN is a float value in Python (from the numpy library). When pandas reads a CSV column "
    "where ALL values are NaN, the column's inferred type can become float64 instead of object (string). "
    "DuckDB then creates the bronze table column as DOUBLE (float) rather than VARCHAR (string)."
)
body(
    "Then in the silver transform, COALESCE(issue_description, '') tried to coalesce a DOUBLE column "
    "with the string literal ''. DuckDB could not reconcile these types and tried to convert "
    "the empty string '' to INT32 — which is impossible."
)

h3("The fix")
body("Add keep_default_na=False to the pandas read call in csv_loader.py:")
code_block([
    "# Before — empty cells become NaN:",
    "df = pd.read_csv(file_path, dtype=str)",
    "",
    "# After — empty cells stay as empty strings:",
    "df = pd.read_csv(file_path, dtype=str, keep_default_na=False)",
])
body(
    "keep_default_na=False tells pandas: 'do not replace empty cells with NaN'. "
    "Empty cells become '' (empty string) in the DataFrame. DuckDB then creates a VARCHAR "
    "column, and COALESCE(VARCHAR, '') works correctly."
)

callout(
    "The lesson:",
    "keep_default_na=False is now a standard part of how you read CSV files for data pipelines. "
    "Without it, empty cells silently become NaN (a float), which causes unexpected type "
    "inference downstream. You will hit this bug on every real client dataset. Remember it.",
    bg="FCE4EC",
)

add_table(
    ["Scenario", "dtype=str (without fix)", "dtype=str, keep_default_na=False (with fix)"],
    [
        ["Empty cell in CSV",        "NaN (float)",             "'' (empty string)"],
        ["Column with all empty",    "float64 dtype",           "object (string) dtype"],
        ["COALESCE(col, '')",        "ConversionException",     "Returns ''"],
        ["Column type in DuckDB",    "DOUBLE or error",         "VARCHAR"],
    ],
    col_widths=[2.2, 2.2, 2.8],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — TESTS
# ══════════════════════════════════════════════════════════════════════════════

h1("6. File: tests/test_transforms.py")

body(
    "Six tests covering both transform layers. Each test is self-contained: "
    "it creates its own CSV, loads it into bronze, runs the transforms, "
    "then queries the database to verify the result."
)

h2("6.1 Test Architecture")

body("Every test follows this pattern:")
code_block([
    "1. set_test_db fixture runs automatically (autouse=True)",
    "   → DB_PATH is overridden to a temp DuckDB file",
    "   → each test gets a completely fresh, empty database",
    "",
    "2. Test creates a CSV row with the specific values it needs",
    "   → uses dict(SAMPLE_PROJECTS[0]) as a base, then overrides specific fields",
    "",
    "3. Test loads CSV → bronze → silver (→ gold if testing gold)",
    "",
    "4. Test queries the result table with SQL",
    "",
    "5. assert checks the value matches expectations",
])

h2("6.2 The Six Tests")

h3("Test 1 — test_silver_date_type")
body("Verifies that planned_finish_date becomes a DATE column in silver (not VARCHAR or TIMESTAMP).")
code_block([
    "col_types = {row[0]: row[1] for row in con.execute('DESCRIBE silver_project_clean').fetchall()}",
    "assert 'DATE' in col_types.get('planned_finish_date', '')",
])
body(
    "DESCRIBE returns column names and types. The assertion checks that 'DATE' appears in the type string. "
    "This test caught Bug 1 (TRY_STRPTIME returning TIMESTAMP instead of DATE)."
)

h3("Test 2 — test_silver_status_normalization")
body("Verifies that 'On Track' (with capital letters and a space) becomes 'on_track' in silver.")
code_block([
    "row['status'] = 'On Track'",
    "# ... load and transform ...",
    "result = con.execute('SELECT status FROM silver_project_clean LIMIT 1').fetchone()[0]",
    "assert result == 'on_track'",
])
body(
    "This test simulates real client data where the status column has inconsistent formatting. "
    "The CASE statement must handle this correctly."
)

h3("Test 3 — test_silver_null_issue_description")
body("Verifies that an empty issue_description survives the transform as an empty string, not NULL.")
code_block([
    "row['issue_description'] = ''",
    "# ... load and transform ...",
    "result = con.execute('SELECT issue_description FROM silver_project_clean LIMIT 1').fetchone()[0]",
    "assert result == ''",
])
body(
    "This test caught Bug 2 (empty strings becoming NaN). After the fix, an empty "
    "issue_description in the CSV correctly becomes an empty string in silver."
)

h3("Test 4 — test_gold_budget_variance_math")
body("Verifies the KPI math with a known input where you can calculate the expected output by hand.")
code_block([
    "row['planned_cost'] = 100000",
    "row['actual_cost']  = 115000",
    "# Expected: variance = 15000, pct = 15.0, is_over_budget = True",
    "",
    "assert row_result[0] == 15000.0   # budget_variance",
    "assert row_result[1] == 15.0      # budget_variance_pct",
    "assert row_result[2] is True      # is_over_budget (15% > 10% threshold)",
])
body(
    "Using round numbers (100000 and 115000) makes the expected values easy to calculate "
    "without floating-point surprises. 115000 - 100000 = 15000. 15000 / 100000 × 100 = 15.0."
)

h3("Test 5 — test_gold_is_late_flag")
body("Verifies that a project finishing 31 days late gets is_late = True and schedule_delay_days = 31.")
code_block([
    "row['planned_finish_date'] = '2024-01-01'",
    "row['actual_finish_date']  = '2024-02-01'   # 31 days later",
    "",
    "assert result[0] is True   # is_late",
    "assert result[1] == 31     # schedule_delay_days",
])
body(
    "January has 31 days, so 2024-01-01 to 2024-02-01 is exactly 31 days. "
    "Using a known calendar gap makes the test deterministic — no approximations."
)

h3("Test 6 — test_gold_on_time_not_late")
body("Verifies that a project finishing on its exact planned date is NOT late.")
code_block([
    "row['planned_finish_date'] = '2024-01-01'",
    "row['actual_finish_date']  = '2024-01-01'   # same date",
    "",
    "assert result[0] is False  # is_late",
    "assert result[1] == 0      # schedule_delay_days",
])
body(
    "This is the boundary condition test. The rule is: 'strictly after planned date = late'. "
    "Same date = on time. This test confirms the boundary is correct."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — DATA FLOW AFTER DAY 2
# ══════════════════════════════════════════════════════════════════════════════

h1("7. How the Data Looks After Day 2")

body("You can inspect all three tables yourself with these commands:")

h3("See the bronze table (raw strings)")
code_block([
    "python -c \"",
    "from src.storage.db import get_connection",
    "from dotenv import load_dotenv",
    "load_dotenv()",
    "con = get_connection()",
    "print(con.execute('SELECT project_name, planned_cost, status FROM bronze_project_raw LIMIT 5').fetchdf())",
    "con.close()\"",
])
body("planned_cost and status are strings: '250000' and 'delayed'.")

h3("See the silver table (cleaned types)")
code_block([
    "python -c \"",
    "from src.storage.db import get_connection",
    "from dotenv import load_dotenv",
    "load_dotenv()",
    "con = get_connection()",
    "print(con.execute('SELECT project_name, planned_cost, planned_finish_date, status FROM silver_project_clean LIMIT 5').fetchdf())",
    "con.close()\"",
])
body("planned_cost is now a float (250000.0), planned_finish_date is a DATE, status is 'on_track' etc.")

h3("See the gold table (with KPI columns)")
code_block([
    "python -c \"",
    "from src.storage.db import get_connection",
    "from dotenv import load_dotenv",
    "load_dotenv()",
    "con = get_connection()",
    "print(con.execute('SELECT project_name, budget_variance_pct, schedule_delay_days, is_over_budget, is_late FROM gold_project_kpi ORDER BY budget_variance_pct DESC LIMIT 5').fetchdf())",
    "con.close()\"",
])
body("You will see the top 5 most over-budget projects with their calculated KPIs.")
body("To run these commands, first generate the tables: python demo/sample_data.py, then run a pipeline manually (Day 3 will automate this).")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — KEY CONCEPTS
# ══════════════════════════════════════════════════════════════════════════════

h1("8. Key Concepts to Remember")

add_table(
    ["Concept", "What it means", "Why it matters"],
    [
        ["Medallion architecture",  "Bronze = raw, Silver = clean, Gold = KPIs",        "Industry standard — clients recognize it immediately"],
        ["TRY_CAST vs CAST",        "TRY returns NULL on bad input; CAST crashes",       "Use TRY at silver layer — client data is always messy"],
        ["TRY_STRPTIME + CAST DATE","Parse date string → force to DATE type",            "TIMESTAMP and DATE are different — be explicit"],
        ["COALESCE(val, default)",  "Return first non-NULL value",                       "Prevents NULL propagation into AI prompts and reports"],
        ["CASE normalization",       "Map many spellings → one controlled value",         "Every client system uses different status labels"],
        ["keep_default_na=False",   "Preserve empty CSV cells as '' not NaN",            "Prevents float type inference on empty columns"],
        ["SELECT * in gold",         "Carry all silver columns forward automatically",    "New silver columns propagate to gold for free"],
        ["Boundary test (on time)", "Test the exact edge condition of a rule",           "Off-by-one errors in date logic are common and subtle"],
    ],
    col_widths=[1.8, 2.4, 2.4],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — TEST RESULTS
# ══════════════════════════════════════════════════════════════════════════════

h1("9. Final Test Results")

body("After the two bug fixes, all 9 tests across Day 1 and Day 2 pass:")

code_block([
    "$ pytest tests/ -v",
    "",
    "tests/test_ingestion.py::test_load_happy_path                PASSED",
    "tests/test_ingestion.py::test_missing_required_column         PASSED",
    "tests/test_ingestion.py::test_empty_project_name_rows_skipped PASSED",
    "tests/test_transforms.py::test_silver_date_type               PASSED",
    "tests/test_transforms.py::test_silver_status_normalization     PASSED",
    "tests/test_transforms.py::test_silver_null_issue_description   PASSED",
    "tests/test_transforms.py::test_gold_budget_variance_math      PASSED",
    "tests/test_transforms.py::test_gold_is_late_flag               PASSED",
    "tests/test_transforms.py::test_gold_on_time_not_late           PASSED",
    "",
    "9 passed in 1.12s",
])

body("Note: The two warnings about dateutil and Jupyter are from your Anaconda environment — they do not affect this project.")

h2("Changes Made to Fix the Bugs")

add_table(
    ["File", "Change", "Why"],
    [
        ["src/ingestion/csv_loader.py",        "Added keep_default_na=False to pd.read_csv()",              "Prevents empty cells from becoming NaN (float)"],
        ["src/transforms/bronze_to_silver.py", "Wrapped TRY_STRPTIME in CAST(... AS DATE)",                  "Forces DATE type instead of TIMESTAMP"],
    ],
    col_widths=[2.5, 2.8, 2.3],
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — WHAT DAY 3 BUILDS ON THIS
# ══════════════════════════════════════════════════════════════════════════════

h1("10. What Day 3 Builds On This")

body("Day 2 left you with gold_project_kpi — 15 rows with all the KPI columns computed. Day 3 will:")

bullet("Read gold_project_kpi and aggregate it into portfolio-level statistics")
bullet("Compute: total projects, % on track / at risk / delayed, top 3 risks, total budget variance")
bullet("Return a Python dictionary (KPISummary) that the AI prompt builder will use")
bullet("Wire all four stages (ingest → silver → gold → aggregate) into a single pipeline.run_pipeline() call")
bullet("Checkpoint: python src/pipeline.py prints '15'")

body("Day 3 is where the system starts feeling like a real product — one command runs the entire data pipeline end to end.")

divider()

p_final = doc.add_paragraph()
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_final.paragraph_format.space_before = Pt(20)
run_final = p_final.add_run("Day 2 complete — 9 tests passing, medallion architecture in place.")
set_font(run_final, bold=True, size=12, color=(39, 174, 96))

# ── Save ───────────────────────────────────────────────────────────────────────
out_path = "docs/day-2-learning-guide.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
