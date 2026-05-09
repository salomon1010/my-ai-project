"""Run this script to generate consulting-playbook.docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Style helpers ──────────────────────────────────────────────────────────────

def set_font(run, bold=False, italic=False, size=None, color=None, font_name=None):
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, bold=True, size=20, color=(44, 62, 80))
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, bold=True, size=15, color=(44, 62, 80))
    return p


def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, bold=True, size=12, color=(52, 73, 94))
    return p


def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def italic_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_font(run, italic=True, size=11, color=(80, 80, 80))
    return p


def code_block(lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 30, 30)
        if i < len(lines) - 1:
            run.add_break()
    doc.add_paragraph()


def quote_block(text, attribution=""):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EBF5FB")
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)
    r1 = p.add_run("“" + text + "”")
    r1.italic = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(41, 128, 185)
    if attribution:
        r1.add_break()
        r2 = p.add_run("    — " + attribution)
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph()


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def numbered(text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def callout(label, text, bg="FFF3CD"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg)
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    r1 = p.add_run(f"{label}  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    doc.add_paragraph()


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "2C3E50")
        tcPr.append(shd)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        fill = "F8F9FA" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
    if col_widths:
        for ri, row in enumerate(t.rows):
            for ci, width in enumerate(col_widths):
                row.cells[ci].width = Inches(width)
    doc.add_paragraph()
    return t


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def section_banner(text, bg="2980B9"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg)
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.15)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(255, 255, 255)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(30)
p.paragraph_format.space_after = Pt(6)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("The AI Consultant Playbook")
set_font(run, bold=True, size=30, color=(44, 62, 80))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(6)
run2 = p2.add_run("From First Meeting to Delivered System")
set_font(run2, bold=True, size=18, color=(41, 128, 185))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(4)
run3 = p3.add_run("A Complete Field Guide for Starting Your AI Reporting Consultancy")
set_font(run3, italic=True, size=13, color=(100, 100, 100))

doc.add_paragraph()
callout("Who this guide is for:",
        "You have built the AI Executive Reporting System. Now you need to walk into a real "
        "company, earn their trust, sell them the engagement, deliver the project on time, "
        "and leave them with something that actually works. This guide covers every step — "
        "from how you shake hands to how you run your sprint demos.",
        bg="D6EAF8")

doc.add_paragraph()

add_table(
    ["Chapter", "Topic", "Pages"],
    [
        ["1", "The Real-World Scenario — Meet Meridian Corp", "~5"],
        ["2", "Your Identity as an AI Consultant", "~4"],
        ["3", "Before the First Meeting — Research & Preparation", "~3"],
        ["4", "Discovery Phase — Questions to Ask Every Manager", "~6"],
        ["5", "Your PowerPoint Deck — Slide by Slide", "~6"],
        ["6", "The Technical Solution — How You Solve Their Problems", "~4"],
        ["7", "Building Your Consulting Team", "~4"],
        ["8", "Agile Delivery — Scrum Ceremonies End to End", "~7"],
        ["9", "Sprint Demo Ceremonies — Every Two Weeks", "~5"],
        ["10", "Stakeholder Communication Cadence", "~4"],
        ["11", "Commercial Model — Pricing, Proposals & Contracts", "~4"],
        ["12", "Handling Objections", "~3"],
        ["13", "Mindset & Confidence — The Inner Game of Consulting", "~3"],
        ["14", "90-Day Launch Plan for Your Company", "~3"],
    ],
    col_widths=[0.5, 4.0, 0.8],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 1: THE REAL-WORLD SCENARIO
# ══════════════════════════════════════════════════════════════════════════════

section_banner("CHAPTER 1   The Real-World Scenario — Meet Meridian Corp")

h1("The Company You Just Walked Into")

body("You have been referred to Meridian Corp — a mid-sized industrial services company "
     "with 1,200 employees across three divisions: IT, Finance, and Operations. They run "
     "30+ active projects at any given time, and their reporting process is a nightmare. "
     "Every Monday morning the CFO receives 12 different spreadsheets from 12 different "
     "people, each formatted differently, each telling a different story.")

body("The COO complained to a mutual contact: 'I have no idea which projects are actually "
     "in trouble until it's too late to fix anything.' Your mutual contact said: "
     "'You should talk to my friend — they just built something that solves exactly this.'")

body("You have a 45-minute intro call with the CFO next Tuesday.")

h2("The 15 Projects You Already Know About")

body("Before that call, their project coordinator sent you a CSV export from their "
     "project tracking tool. You've already run it through your system. Here is what "
     "the data tells you — and this is your secret weapon walking in:")

add_table(
    ["Project", "Dept", "Status", "Budget Variance", "Days Late", "Risk"],
    [
        ["Cloud Migration Phase 1", "IT", "Delayed", "+24%", "45 days", "High"],
        ["ERP Upgrade", "IT", "Completed", "-3%", "On time", "Low"],
        ["Security Audit Program", "IT", "Delayed", "+9.5%", "20 days", "Medium"],
        ["Data Warehouse Rebuild", "IT", "On Track", "-7%", "On time", "Low"],
        ["Mobile App Launch", "IT", "At Risk", "+32%", "77 days", "CRITICAL"],
        ["Annual Budget Reconciliation", "Finance", "Completed", "-2%", "On time", "Low"],
        ["Regulatory Compliance Report", "Finance", "Delayed", "+13%", "10 days", "Medium"],
        ["Financial System Integration", "Finance", "At Risk", "+21%", "76 days", "High"],
        ["Cost Reduction Initiative", "Finance", "Completed", "-3%", "On time", "Low"],
        ["Vendor Contract Renegotiation", "Finance", "On Track", "+8%", "5 days", "Low"],
        ["Supply Chain Optimization", "Operations", "Delayed", "+10%", "30 days", "High"],
        ["Warehouse Automation", "Operations", "On Track", "-3%", "On time", "Medium"],
        ["Quality Management System", "Operations", "Delayed", "+14%", "25 days", "Medium"],
        ["Logistics Partner Integration", "Operations", "On Track", "-2%", "On time", "Low"],
        ["Fleet Management Upgrade", "Operations", "Delayed", "+26%", "31 days", "High"],
    ],
    col_widths=[2.0, 0.8, 0.8, 1.0, 0.8, 0.7],
)

h2("What the Data Tells You Before You Walk In")

body("Your AI system has already generated this analysis. You know things before anyone "
     "has told you a single word:")

callout("Portfolio headline:",
        "8 of 15 projects are late. 7 are over budget. The IT division has one critical-risk "
        "project (Mobile App Launch) with a terminated vendor and +32% cost overrun. "
        "Finance has a systemic integration failure. Operations has 4 delayed projects "
        "all citing supply chain and procurement issues — suggesting a shared root cause.",
        bg="FADBD8")

add_table(
    ["What you know", "What it means for your meeting"],
    [
        ["53% of projects are late", "Leadership knows something is wrong but may not know the scale"],
        ["Mobile App Launch: vendor terminated at +32%", "This is the burning issue — someone's job may be on the line"],
        ["3 Operations delays cite procurement", "Systemic problem, not individual PM failure — points to process gap"],
        ["Finance has 2 at-risk projects running simultaneously", "CFO is sitting on undisclosed risk right now"],
        ["4 projects completed on/under budget", "There are bright spots — identify what those teams did right"],
    ],
    col_widths=[2.5, 3.8],
)

h2("The Narrative You Will Walk In With")

body("You are not walking in to sell software. You are walking in to tell them what is "
     "happening in their own company — using data they already have but cannot read. "
     "That is your opening power.")

quote_block(
    "Before our call today, your team shared a data export with me. I ran it through "
    "our analysis engine last night. I found eight projects currently behind schedule, "
    "seven over budget, and one in a critical situation I'd like to discuss privately. "
    "I can show you all of this in a dashboard in about ten minutes. Would that be useful?",
    "Your opening line to the CFO"
)

callout("Key insight:",
        "You already have the answers. Your first job is to make them feel understood — "
        "not to impress them with technology. The demo comes second. Empathy comes first.",
        bg="D5F5E3")

divider()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 2: YOUR IDENTITY AS AN AI CONSULTANT
# ══════════════════════════════════════════════════════════════════════════════

section_banner("CHAPTER 2   Your Identity as an AI Consultant")

h1("How to Present Yourself")

body("The most important thing to understand: you are not a software developer who builds "
     "tools. You are a business advisor who uses AI to surface insights that help leaders "
     "make better decisions. The technology is the means; the outcome is the product.")

h2("Your One-Line Positioning Statement")

callout("Your pitch:",
        "I help mid-sized companies turn their raw project and operational data into "
        "AI-generated executive insights — so leaders can see what's really happening "
        "and act before small problems become expensive ones.",
        bg="D6EAF8")

body("Memorize this. Say it the same way every time. It answers: what you do, who you do "
     "it for, and why it matters. It does not mention Python, DuckDB, or Streamlit.")

h2("The Three Hats You Wear")

add_table(
    ["Hat", "When you wear it", "What you say"],
    [
        ["Business analyst",
         "Discovery phase — understanding their pain",
         "'Walk me through your current Monday morning reporting ritual.'"],
        ["Solution architect",
         "Proposal phase — designing the fix",
         "'Here is the data pipeline we will build and why each layer exists.'"],
        ["Delivery manager",
         "Execution phase — running the project",
         "'This sprint we are delivering the KPI dashboard. Demo is Friday at 2pm.'"],
    ],
    col_widths=[1.4, 2.0, 2.9],
)

h2("Your Elevator Pitch for Different Audiences")

h3("To the CFO / COO (5 seconds)")
quote_block("I turn your project data into AI reports that tell you what's at risk before "
            "it becomes a crisis. Most clients see full ROI within 90 days.")

h3("To the IT Director (30 seconds)")
quote_block("We build a lightweight data pipeline — CSV or API ingestion into DuckDB, "
            "a medallion architecture with bronze, silver, and gold layers, and a Claude AI "
            "reporting layer on top. The whole stack runs on a single server. No cloud "
            "vendor lock-in, no Databricks licence fees.")

h3("To a Project Manager (30 seconds)")
quote_block("Instead of filling out status reports that nobody reads, your data flows "
            "automatically into a dashboard. Your manager sees exactly what you're managing "
            "without you having to write a single PowerPoint slide on a Friday afternoon.")

h2("What You Are Not")

body("Be explicit about what you are NOT, because it builds trust:")
bullet("Not a staffing agency — you deliver a working system, not warm bodies")
bullet("Not a software vendor selling perpetual licences — you build and transfer ownership")
bullet("Not a data scientist — you are a delivery-focused consultant who uses AI pragmatically")
bullet("Not a replacement for their team — you augment what they already have")

h2("How to Dress and Carry Yourself")

add_table(
    ["Situation", "Dress code", "Body language tip"],
    [
        ["First discovery meeting", "Business casual — one level above their standard",
         "Arrive 10 minutes early. Have a printed one-page capability summary."],
        ["Executive presentation", "Business formal — jacket, no tie unless they wear one",
         "Stand when presenting. Use a clicker. Speak slower than feels natural."],
        ["Technical workshop", "Smart casual — same level as their tech team",
         "Sit at the table, not the head. Draw architecture diagrams by hand first."],
        ["Sprint demo", "Business casual",
         "Demo on your laptop, not projected. Let them reach for the keyboard."],
    ],
    col_widths=[1.7, 1.8, 2.8],
)

divider()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 3: BEFORE THE FIRST MEETING
# ══════════════════════════════════════════════════════════════════════════════

section_banner("CHAPTER 3   Before the First Meeting — Research & Preparation")

h1("Do This Before Every Client Meeting")

body("Unprepared consultants ask questions that Google could answer. Prepared consultants "
     "ask questions that only an insider can answer. The difference is 2 hours of research.")

h2("The Pre-Meeting Research Checklist")

add_table(
    ["Research area", "Where to look", "What you're looking for"],
    [
        ["Company financials", "Annual report, LinkedIn, Crunchbase",
         "Revenue size, growth trajectory, recent acquisitions"],
        ["Industry context", "Industry trade publications, Google News",
         "Regulatory pressures, competitive threats, sector trends"],
        ["Technology stack", "LinkedIn job postings, their careers page",
         "What systems they run (SAP, Salesforce, Oracle), what they're hiring for"],
        ["Leadership background", "LinkedIn profiles of your contacts",
         "Where they came from, what they care about, how long they've been there"],
        ["Recent news", "Google News: company name",
         "M&A activity, cost-cutting announcements, new strategic initiatives"],
        ["Their data / CSV", "The file they sent you",
         "Run it through your system. Know the numbers before they do."],
    ],
    col_widths=[1.6, 1.7, 3.0],
)

h2("Preparing Your One-Page Company Summary")

body("Before the meeting, write a single page (for yourself, not to hand out) that answers:")
bullet("What does this company do and how do they make money?")
bullet("What is their biggest challenge right now based on public information?")
bullet("Who is in the room and what do they care about individually?")
bullet("What is one insight from their data that will surprise them?")
bullet("What is my desired outcome from this meeting? (Be specific: discovery call, "
       "proposal approval, or pilot sign-off)")

h2("Preparing Your Tech for the Meeting")

code_block([
    "Pre-meeting checklist:",
    "",
    "[ ] Run demo_reset.py — loads clean sample data",
    "[ ] Run streamlit app — verify it opens in < 5 seconds",
    "[ ] Generate AI report — verify 5 sections appear correctly",
    "[ ] Export PDF — verify file downloads properly",
    "[ ] Switch COMPANY_NAME in .env to 'Meridian Corp'",
    "[ ] Charge laptop, bring charger, have HDMI + USB-C adapters",
    "[ ] Have offline backup: screenshots of the dashboard on your phone",
    "[ ] Close all other browser tabs and apps",
])

callout("Pro tip:",
        "Load their actual CSV data (the one they sent you) into the system before "
        "the meeting. When you show them their own projects in the dashboard instead "
        "of sample data, the room changes. It becomes real.",
        bg="D5F5E3")

divider()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 4: DISCOVERY PHASE — QUESTIONS TO ASK
# ══════════════════════════════════════════════════════════════════════════════

section_banner("CHAPTER 4   Discovery Phase — Questions to Ask Every Manager")

h1("The Art of the Discovery Interview")

body("Discovery is not a questionnaire. It is a conversation where you help the client "
     "articulate a problem they already feel but cannot yet name. Your job is to listen "
     "80% of the time, ask questions that open doors, and reflect back what you hear "
     "in cleaner language than they used.")

quote_block("The quality of your solution is directly proportional to the quality of your "
            "questions.", "Classic consulting principle")

h2("Questions for the CFO / COO (Executive Sponsor)")

h3("Opening questions — build rapport and context")
bullet("How do you currently get visibility into your project portfolio on a weekly basis?")
bullet("Walk me through your Monday morning reporting ritual. What lands on your desk "
       "and from how many sources?")
bullet("If you could wave a magic wand and fix one thing about how you receive information "
       "from your teams, what would it be?")

h3("Problem-depth questions — understand the pain")
bullet("When was the last time a project surprised you — something you didn't see coming "
       "until it was already a problem?")
bullet("How much time does your leadership team collectively spend each week just "
       "gathering and consolidating status information?")
bullet("What does it cost you when a project is 30 days late and you find out after the fact?")
bullet("How confident are you, right now, in the accuracy of the project status "
       "information you have?")

h3("Outcome questions — understand what success looks like")
bullet("If I showed you the health of your entire project portfolio in a single screen "
       "in under 60 seconds, what would that change for you?")
bullet("What would you need to see to believe this was working?")
bullet("Who else in the company would benefit from this kind of visibility?")

h2("Questions for the IT Director / CTO")

h3("Technical landscape questions")
bullet("What systems are your project teams currently using to track status? "
       "(Jira, Microsoft Project, spreadsheets, custom tools?)")
bullet("Where does your project data actually live today — what format, what frequency "
       "is it updated?")
bullet("Do you have a data team or data engineer on staff, or does all data work fall "
       "to the IT team?")
bullet("What is your current analytics or reporting infrastructure? Any BI tools, "
       "data warehouses, or dashboards already in place?")
bullet("What security and compliance requirements do we need to know about for "
       "handling project data?")

h3("Integration questions")
bullet("Would the data come to us as CSV exports, database queries, or API calls?")
bullet("Who owns the data pipeline today — IT, Finance, or individual project managers?")
bullet("How often does the source data change? Is it updated daily, weekly, real-time?")

h2("Questions for the Project Management Office (PMO) Lead")

h3("Process questions")
bullet("How do project managers currently submit their status updates? What is the template?")
bullet("What percentage of PMs submit their reports on time vs. late?")
bullet("What information do you currently capture that you wish more people actually read?")
bullet("What information do you wish you were capturing but aren't?")
bullet("Have you tried automated reporting before? What happened?")

h3("Adoption questions")
bullet("Who are the two or three project managers who would be your biggest champions "
       "for a new system?")
bullet("Who would be the most resistant, and why?")
bullet("What would a PM need to see in the first two weeks to believe this was better "
       "than what they have now?")

h2("Questions for Individual Project Managers")

h3("Day-to-day reality questions")
bullet("How long does it take you each week to prepare your status report?")
bullet("Have you ever submitted a status that said 'green' when you privately thought "
       "it should be 'amber'? What made you do that?")
bullet("What information do you wish your manager had about your project that the "
       "current report doesn't capture well?")
bullet("What's the most frustrating part of the current reporting process?")

h3("Impact questions")
bullet("If your manager had real-time visibility into your project budget and schedule "
       "every day, would that help you or create more pressure?")
bullet("How quickly do you typically get feedback after submitting a status report?")

h2("Listening for Red Flags")

body("During discovery, certain phrases signal deeper organizational problems. When you "
     "hear these, probe further:")

add_table(
    ["What they say", "What it might mean", "Follow-up question"],
    [
        ["'We already tried something like this before'",
         "Previous initiative failed — technical, political, or adoption issue",
         "'What happened with that? What would you do differently this time?'"],
        ["'I'm not sure IT will approve it'",
         "Shadow IT tension — tech team may feel bypassed",
         "'Who would need to be involved from IT for this to move forward?'"],
        ["'We just need better spreadsheets'",
         "They don't yet see the value of AI — need education before sale",
         "'What would better spreadsheets give you that you don't have today?'"],
        ["'The data quality is really poor'",
         "Data governance problem — pipeline must include cleaning layer",
         "'On a scale of 1-10, how reliable is the data that comes in?'"],
        ["'We don't really have budget right now'",
         "Budget objection — either real or they need stronger ROI argument",
         "'If I could show you a clear ROI within 90 days, would budget free up?'"],
    ],
    col_widths=[1.8, 2.0, 2.5],
)

h2("How to Take Notes in Discovery")

body("Do NOT open a laptop and type. Use a physical notebook. It signals presence and "
     "respect. After the meeting, transcribe your notes into structured sections:")
bullet("Their words (exact quotes — these become your proposal language)")
bullet("Pain points (ranked by how much emotion they showed when describing them)")
bullet("Constraints (budget, timeline, tech, political)")
bullet("Champions (who wants this) vs. skeptics (who will resist)")
bullet("Data available (what systems, what format, what quality)")
bullet("Open questions (what you still need to find out)")

divider()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 5: YOUR POWERPOINT DECK — SLIDE BY SLIDE
# ══════════════════════════════════════════════════════════════════════════════

section_banner("CHAPTER 5   Your PowerPoint Deck — Slide by Slide")

h1("The Executive Presentation")

body("You will use two different decks at different stages. Never use the same deck for "
     "both purposes — they have different audiences and different objectives.")

add_table(
    ["Deck", "When", "Length", "Objective"],
    [
        ["Discovery deck", "First meeting — before you know their problem",
         "8-10 slides", "Get them to talk. Book the next meeting."],
        ["Proposal deck", "Second meeting — after discovery",
         "15-20 slides", "Get them to say yes and sign."],
    ],
    col_widths=[1.4, 2.5, 0.8, 1.6],
)

h2("Deck 1: The Discovery Deck (First Meeting)")

h3("Slide 1 — Title slide")
code_block([
    "Title:     Turning Project Data Into Executive Intelligence",
    "Subtitle:  A 30-Minute Conversation About Your Reporting Challenge",
    "Visual:    Clean, minimal. Your logo + their logo if you have it.",
    "Note:      Don't show this long. Move to Slide 2 within 30 seconds.",
])

h3("Slide 2 — The Problem (Mirror Their World)")
code_block([
    "Title:     What We Hear From Leaders Like You",
    "",
    "3 pain points as large quotes (use your research / their data):",
    "  'By the time I see a red project, it's already too late to fix it.'",
    "  'I get 8 different status reports every Monday — none of them match.'",
    "  'My project managers spend Friday afternoons writing reports",
    "   instead of managing projects.'",
    "",
    "Note: If they nod at any of these, pause and say:",
    "      'Is that something you experience here too?'",
])

h3("Slide 3 — The Cost of the Problem")
code_block([
    "Title:     What Late Visibility Costs",
    "",
    "3 columns:",
    "  TIME:     Avg 4-8 hours/week per manager in reporting overhead",
    "  MONEY:    Projects caught late cost 3-5x more to fix than projects",
    "            caught early",
    "  TRUST:    Executives who don't trust their data make slower decisions",
    "",
    "Note: Make at least one of these numbers feel personal to their company.",
    "      'Based on the data you shared, your team has lost an estimated",
    "       X days of schedule across 8 projects this year.'",
])

h3("Slide 4 — What We Do (One Sentence)")
code_block([
    "Title:     We connect your raw data to AI-generated insights",
    "",
    "Simple diagram:",
    "  [Their CSV / Database]  →  [AI Pipeline]  →  [Executive Dashboard + PDF]",
    "",
    "Underneath: 'No new tools for your project managers.'",
    "             'No IT infrastructure changes.'",
    "             'First report in 30 days.'",
])

h3("Slide 5 — 60-Second Live Demo")
code_block([
    "Title:     Let me show you what your data looks like right now",
    "",
    "Action:    Close PowerPoint. Open your Streamlit dashboard.",
    "           Show their own data (or sample with their company name).",
    "           Click 'Generate Report'. Let the AI write the summary.",
    "           Export the PDF.",
    "",
    "Duration:  3-5 minutes maximum.",
    "",
    "Closing:   'This is what we could deliver for your team in 30 days.'",
])

h3("Slide 6 — How It Works (Simple Architecture)")
code_block([
    "Title:     How The System Works",
    "",
    "Three boxes left to right:",
    "  Box 1 — DATA IN",
    "  Your project data (CSV, Excel, or direct database connection)",
    "  Updated weekly or daily — your choice",
    "",
    "  Box 2 — AI PIPELINE",
    "  Cleans and standardises the data",
    "  Calculates KPIs: budget variance, schedule delay, risk level",
    "  Sends structured data to Claude AI",
    "",
    "  Box 3 — INSIGHTS OUT",
    "  Live dashboard for executives",
    "  AI-written report: risks, budget concerns, recommended actions",
    "  PDF export for board presentations",
    "",
    "Footer note: 'No hallucinations — every number in the AI report",
    "              traces directly back to your data.'",
])

h3("Slides 7–8 — Why Us / Why Now")
code_block([
    "Slide 7 — Why Us:",
    "  3 differentiators (pick the ones most relevant to this client):",
    "  - We deliver working software in 30 days, not 6 months",
    "  - Your team owns the system when we're done — no vendor lock-in",
    "  - We use Claude AI, the industry-leading model for structured analysis",
    "",
    "Slide 8 — Why Now:",
    "  Their own data: '8 projects currently behind schedule.'",
    "  'Every week without visibility is another week of unmanaged risk.'",
    "  'The pilot is $7,500. The cost of a single delayed project is more.'",
])

h3("Slide 9 — The Ask")
code_block([
    "Title:     The Next Step",
    "",
    "One ask only. Do not give them options — give them a decision:",
    "",
    "  'I'd like to run a 30-day pilot using your live project data.",
    "   Investment: $7,500. Deliverable: working dashboard + first AI report.",
    "   If you love it, we expand. If not, you keep everything we built.'",
    "",
    "Then stop talking. Wait for their response.",
])

h2("Deck 2: The Proposal Deck (After Discovery)")

h3("Structure overview")
add_table(
    ["Slide #", "Title", "Purpose"],
    [
        ["1", "Cover — [Client Name] AI Reporting Pilot", "Personalization signal"],
        ["2", "What We Heard (from discovery)", "Prove you listened"],
        ["3", "Your Current Situation (their data)", "Show you've done the analysis"],
        ["4", "The Root Cause", "Reframe the problem professionally"],
        ["5", "Our Solution — Architecture Diagram", "Build technical confidence"],
        ["6", "What You Will Have in 30 Days", "Concrete deliverables"],
        ["7", "The Delivery Timeline — Sprint by Sprint", "Show rigor and process"],
        ["8", "Demo Ceremony Schedule", "Set expectations for ongoing visibility"],
        ["9", "Your Team", "Introduce the people they're hiring"],
        ["10", "Investment & ROI", "The commercial case"],
        ["11", "Terms & Next Steps", "Close"],
        ["12–15", "Appendix: Technical Architecture, Data Security, Sample Reports", "For IT due diligence"],
    ],
    col_widths=[0.6, 2.5, 3.2],
)

h2("Presentation Techniques That Work in Boardrooms")

bullet("The pregnant pause: after asking a question, do NOT fill the silence. Wait 10 full seconds.")
bullet("Mirroring: repeat the last 3 words of what they say as a question. "
       "('...we can't trust the data.' → 'Can't trust the data?') They will elaborate.")
bullet("Hypothesis-led framing: 'Based on what I've seen so far, my hypothesis is X. "
       "Am I on the right track?' This is more powerful than asking 'What do you think?'")
bullet("The visual pause: when you put a new slide up, say nothing for 5 seconds. "
       "Let them read it before you explain it.")
bullet("Kill the laser pointer: if you're using a clicker, move to what you're pointing "
       "at. Standing still and pointing with a laser is passive.")

divider()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 6: THE TECHNICAL SOLUTION
# ══════════════════════════════════════════════════════════════════════════════

section_banner("CHAPTER 6   The Technical Solution — How You Solve Their Problems")

h1("Translating Technology Into Business Value")

body("Every technical choice you made has a business reason. Learn to speak both "
     "languages fluently — you will face technical scrutiny from IT and business "
     "scrutiny from leadership in the same meeting.")

h2("The System Explained to Three Different Audiences")

h3("To the CFO")
quote_block("Your project data goes in one end. An AI-written executive briefing comes "
            "out the other. The whole process takes about 60 seconds, runs automatically, "
            "and the report is ready before your Monday morning meeting. You don't need to "
            "change how your teams work — we connect to the data they're already creating.")

h3("To the IT Director")
quote_block("We ingest CSV exports or query your existing database directly via ODBC. "
            "The pipeline runs in Python, stores data in DuckDB using a medallion "
            "architecture — bronze for raw, silver for cleaned, gold for KPI-computed. "
            "The AI layer calls Claude via the Anthropic API with a structured prompt that "
            "injects only verified data, eliminating hallucination risk. The dashboard is "
            "Streamlit. The whole stack deploys to a single Linux VM or Docker container. "
            "No Databricks, no Snowflake, no cloud-native lock-in.")

h3("To a Project Manager")
quote_block("You keep tracking projects exactly the way you do now. Every week, we pull "
            "your data and build a report automatically. Your manager sees the same "
            "information you see — there are no surprises on either side.")

h2("The Anti-Hallucination Architecture — Why It Matters")

body("When you present AI to executives, the most common fear is: 'Will the AI make "
     "things up?' This is a legitimate technical concern. Here is how you explain your "
     "solution in non-technical terms:")

callout("What to say:",
        "Every number in the AI report traces directly back to a row in their data. "
        "The AI does not analyse or reason from scratch — it narrates what the data "
        "already shows. We inject the KPI summary into the prompt as structured facts, "
        "and the model writes in clear English around those facts. "
        "If a project name doesn't appear in the data, it cannot appear in the report.",
        bg="D6EAF8")

h2("Problem → Solution Mapping for Meridian Corp")

add_table(
    ["Their Problem", "Your Solution", "Delivered by"],
    [
        ["8 different spreadsheets every Monday",
         "Single dashboard with live data, one source of truth",
         "Sprint 1 — Day 14"],
        ["No early warning on at-risk projects",
         "Traffic-light KPI cards update every time data refreshes",
         "Sprint 1 — Day 14"],
        ["Managers spend 4 hrs/week on status reports",
         "AI generates the narrative automatically; PMs enter data, not write",
         "Sprint 2 — Day 28"],
        ["Mobile App Launch crisis invisible until too late",
         "Critical-risk flag on dashboard, AI escalation section in report",
         "Sprint 1 — Day 14"],
        ["Board presentations require manual slide building",
         "One-click PDF export with cover page, KPIs, AI analysis",
         "Sprint 3 — Day 42"],
        ["No visibility into cross-department patterns",
         "Issues-by-department aggregation in the gold layer",
         "Sprint 2 — Day 28"],
    ],
    col_widths=[2.1, 2.5, 1.5],
)

divider()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 7: BUILDING YOUR CONSULTING TEAM
# ══════════════════════════════════════════════════════════════════════════════

section_banner("CHAPTER 7   Building Your Consulting Team")

h1("You Cannot Do This Alone (Eventually)")

body("In the beginning, you are the entire company. You sell, you deliver, you invoice, "
     "you support. That is fine for the first one or two clients. By the third engagement, "
     "you will need help — or you will deliver late and damage your reputation.")

h2("The Lean Team Model — First 12 Months")

body("Start with a core of 2–3 people maximum. Hire only for a confirmed engagement, "
     "not in anticipation of one.")

add_table(
    ["Role", "What they do", "Hire when", "Profile to look for"],
    [
        ["You (Founder / Engagement Lead)",
         "Sell, architect, run discovery, present, manage client relationship",
         "From day one",
         "You"],
        ["Junior Python Developer",
         "Build and maintain the pipeline, fix bugs, write tests, deploy",
         "When you have signed your first paid client",
         "2-3 yrs Python, data engineering experience, self-directed"],
        ["Business Analyst / Project Coordinator",
         "Run discovery interviews, document requirements, manage sprint ceremonies",
         "When you have 2 simultaneous clients",
         "Strong communicator, comfortable with data, organised"],
        ["Part-time Designer",
         "PowerPoint decks, dashboard styling, PDF template design",
         "When you pitch to enterprise (>500 employees)",
         "Freelancer or agency — not a full hire"],
    ],
    col_widths=[1.6, 2.2, 1.4, 1.8],
)

h2("Where to Find Your First Hire")

bullet("LinkedIn: search 'junior data engineer' + your city, filter 1–3 years experience")
bullet("Upwork: for the first project, consider a contract freelancer before full hire")
bullet("University connections: data science / computer science final-year students "
       "are motivated and cost-effective for junior pipeline work")
bullet("Your own network: post on LinkedIn — 'Looking for a Python developer who wants "
       "to work on AI consulting projects.' You will get referrals.")

h2("How to Interview a Junior Developer for This Role")

h3("Technical screening question (give them 48 hours)")
code_block([
    "Assignment: Write a Python function that reads a CSV file, computes the",
    "following metrics, and returns them as a dictionary:",
    "  - Total number of projects",
    "  - Number of projects where actual_cost > planned_cost",
    "  - Average of (actual_cost - planned_cost) across all projects",
    "",
    "Bonus: Write one pytest test for your function.",
    "",
    "Use only pandas and the Python standard library.",
])

body("What you're looking for: clean code (good variable names, no magic numbers), "
     "the test exists and uses a fixture, no hardcoded file paths. You are NOT looking "
     "for perfection — you are looking for someone who can learn quickly.")

h3("Culture / fit questions")
bullet("'Describe the last time you had to figure something out with no documentation. "
       "What did you do?'")
bullet("'If a test fails and you don't know why, walk me through your debugging process.'")
bullet("'How do you feel about showing half-finished work to a client in a demo?'")

h2("The Subcontractor Model")

body("Before hiring your first full-time employee, consider the subcontractor model. "
     "You win the client, bring in a vetted freelancer for the technical delivery, "
     "and manage the relationship. Your margin is the difference between what the "
     "client pays you and what you pay the freelancer. This is how most solo consultants "
     "scale without employment risk.")

callout("Tax and legal note:",
        "In most countries, a subcontractor arrangement requires clear contracts, "
        "proper invoicing, and no day-to-day control of how they work. Speak to an "
        "accountant before structuring this. Get a standard consultancy agreement "
        "template from a business lawyer.",
        bg="FADBD8")

h2("How Long Does Building a Team Take?")

add_table(
    ["Month", "Team composition", "Capacity"],
    [
        ["Months 1–3", "You alone", "1 small client (30-day pilot)"],
        ["Months 4–6", "You + 1 junior developer (contract)", "1-2 clients in parallel"],
        ["Months 7–9", "You + 1 junior dev + 1 BA / coordinator", "2-3 clients"],
        ["Months 10–12", "Above + part-time designer, subcontractor on demand", "3-5 clients"],
        ["Year 2+", "Permanent junior staff, senior dev hire if needed", "5+ clients"],
    ],
    col_widths=[1.2, 2.8, 2.3],
)

divider()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 8: AGILE DELIVERY — SCRUM CEREMONIES END TO END
# ══════════════════════════════════════════════════════════════════════════════

section_banner("CHAPTER 8   Agile Delivery — Scrum Ceremonies End to End")

h1("Why Agile and Why Scrum")

body("Traditional consulting delivery is waterfall: you disappear for 3 months and "
     "return with a system. If anything was misunderstood, you find out at the end "
     "when it's expensive to fix. Agile delivery flips this: you show something "
     "working every two weeks and course-correct continuously.")

body("For AI consulting specifically, Agile is non-negotiable because:")
bullet("The client doesn't know exactly what they want until they see it")
bullet("Data quality problems are discovered only when you actually connect to their data")
bullet("Executive priorities shift — what was critical in Week 1 may be irrelevant by Week 8")
bullet("Demos every two weeks maintain trust and prevent 'scope creep ambushes'")

h2("Your Sprint Structure")

callout("Standard sprint length for this type of project:",
        "2-week sprints (10 working days). Long enough to deliver something meaningful. "
        "Short enough to catch problems before they compound.",
        bg="D6EAF8")

h2("The Five Scrum Ceremonies — How You Run Each One")

h3("Ceremony 1: Sprint Planning (Day 1 of each sprint, 2 hours)")

body("Purpose: Decide exactly what will be built this sprint and who will do it.")

add_table(
    ["Time", "Activity", "Who speaks"],
    [
        ["0:00–0:20", "Review the product backlog — top priority items from last sprint review",
         "You (Scrum Master / Product Owner combined)"],
        ["0:20–0:50", "Break selected items into tasks — specific, estimable, testable",
         "You + developer"],
        ["0:50–1:20", "Estimate effort (use T-shirt sizing: S/M/L or story points 1/2/3/5/8)",
         "Developer leads, you facilitate"],
        ["1:20–1:50", "Commit to a sprint goal: one sentence describing what success looks like",
         "Agreed by all"],
        ["1:50–2:00", "Write sprint goal on the board / shared doc. Done.",
         "You document"],
    ],
    col_widths=[0.7, 3.0, 2.6],
)

callout("Sprint goal example:",
        "'By end of sprint 1, the client can log into the dashboard and see all 15 projects "
        "in a traffic-light table with budget variance and schedule delay data.'",
        bg="FFF3CD")

h3("Ceremony 2: Daily Standup (Every weekday, 15 minutes, same time)")

body("Purpose: Synchronize the team. Surface blockers before they become delays.")

body("Format — each person answers only three questions:")
bullet("What did I complete since yesterday?")
bullet("What will I complete today?")
bullet("What is blocking me? (If nothing, say 'no blockers')")

body("Rules for running it well:")
bullet("Stand up — physically. It keeps it short.")
bullet("Start on time, even if people are missing. This trains punctuality.")
bullet("Do NOT solve problems in the standup. 'Parking lot' any discussion "
       "that needs more than 30 seconds: 'Let's take that offline.'")
bullet("For a 2-person team (you + 1 developer), this takes 5 minutes.")

callout("Client standups:",
        "For a consulting engagement, you do NOT include the client in daily standups. "
        "They attend Sprint Planning and Sprint Review. Too much access creates "
        "micromanagement and slows delivery. Give them a weekly written summary instead.",
        bg="FFF3CD")

h3("Ceremony 3: Sprint Review / Demo (Last day of each sprint, 1 hour)")

body("Purpose: Show the client what was built. Get feedback. Update the backlog.")
body("This is covered in depth in Chapter 9.")

h3("Ceremony 4: Sprint Retrospective (After the review, 45 minutes, internal only)")

body("Purpose: Improve how the team works, not what is being built. This is for "
     "the delivery team — not the client.")

code_block([
    "Retrospective format (use a physical or virtual board):",
    "",
    "Column 1 — WENT WELL",
    "  What should we keep doing?",
    "  What made us fast or effective?",
    "",
    "Column 2 — COULD IMPROVE",
    "  What slowed us down?",
    "  What caused rework or confusion?",
    "",
    "Column 3 — ACTION ITEMS",
    "  One or two specific changes to make next sprint",
    "  Each action item has an owner and a deadline",
    "",
    "Rule: Never raise an issue without proposing a solution.",
])

h3("Ceremony 5: Backlog Refinement (Mid-sprint, 1 hour)")

body("Purpose: Prepare the next sprint before planning day. Items at the top of the "
     "backlog should be well-defined, estimated, and ready to pull into a sprint.")

body("Typical agenda:")
bullet("Review upcoming user stories — are acceptance criteria clear?")
bullet("Break large stories into smaller tasks if needed")
bullet("Remove items that are no longer a priority")
bullet("Add new items that came up in the current sprint")

h2("The Product Backlog — Your Master Task List")

body("Your product backlog is a prioritised list of everything the system needs to do. "
     "You own this list. The client informs it. No one else adds to it without your review.")

h3("Backlog structure for the Meridian Corp engagement:")

add_table(
    ["Priority", "User Story", "Acceptance Criteria", "Sprint"],
    [
        ["1", "As a CFO, I can see all 15 projects in a single dashboard",
         "Dashboard loads in < 5s, shows status, budget variance, schedule delay",
         "Sprint 1"],
        ["2", "As a CFO, I can see which projects are at critical risk",
         "Critical/high risk projects shown with red indicator, sorted to top",
         "Sprint 1"],
        ["3", "As a COO, I receive an AI-generated portfolio summary",
         "5-section report generated in < 30s, references real project names",
         "Sprint 2"],
        ["4", "As an exec, I can export the report as a professional PDF",
         "PDF includes cover, KPI table, AI analysis, page numbers",
         "Sprint 2"],
        ["5", "As a PM, my status data flows in without manual re-entry",
         "CSV from project tool imported automatically via scheduled job",
         "Sprint 3"],
        ["6", "As an IT director, I can see system audit logs",
         "Log file captures every data load with timestamp and row counts",
         "Sprint 3"],
    ],
    col_widths=[0.6, 2.2, 2.0, 0.7],
)

divider()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 9: SPRINT DEMO CEREMONIES
# ══════════════════════════════════════════════════════════════════════════════

section_banner("CHAPTER 9   Sprint Demo Ceremonies — Every Two Weeks")

h1("The Demo Is Your Most Important Meeting")

body("Every two weeks, you show the client something that works. Not mockups. Not "
     "wireframes. Not slides. Working software with their data.")

body("The sprint demo ceremony is the single most powerful tool you have for:")
bullet("Maintaining trust ('I can see it's real and it's progressing')")
bullet("Catching misalignments before they become expensive ('That's not what I meant')")
bullet("Managing scope ('This is what we built — shall we add X to next sprint?')")
bullet("Building the case for expanding the engagement ('I want to show the CEO next week')")

h2("The Demo Ceremony Agenda (60 minutes)")

add_table(
    ["Time", "Activity", "Who leads"],
    [
        ["0:00–0:05", "Welcome and sprint goal recap — what we committed to build",
         "You"],
        ["0:05–0:30", "Live demo — show the working software with their data",
         "You (or developer)"],
        ["0:30–0:40", "Structured feedback — three questions (see below)",
         "You facilitate, client answers"],
        ["0:40–0:50", "Backlog review — what's planned for next sprint",
         "You"],
        ["0:50–0:58", "Blocker / risk discussion — anything that could delay sprint 2",
         "You"],
        ["0:58–1:00", "Confirm next demo date and who attends",
         "You"],
    ],
    col_widths=[0.7, 3.2, 2.4],
)

h2("The Three Feedback Questions to Ask at Every Demo")

body("Do not ask 'What do you think?' — it produces vague answers. "
     "Ask these three specific questions:")

bullet("'Does this show you the information you actually need — or is there something "
       "important that's missing?'")
bullet("'Is there anything here that works differently than you expected?'")
bullet("'Looking at next sprint's plan, is the priority order correct — or should "
       "we move anything up or down?'")

body("Write their answers in the meeting. Feed them directly into the next sprint planning session.")

h2("How to Demo Like a Professional")

h3("Before the demo")
code_block([
    "[ ] Run demo_reset.py — start from a known clean state",
    "[ ] Load their actual data (not sample data if you have theirs)",
    "[ ] Test every click you plan to make — no surprises on screen",
    "[ ] Have a backup: screenshots on your phone in case the laptop fails",
    "[ ] Set screen resolution to 125% — text must be readable from the back of the room",
    "[ ] Close Slack, email, and all notifications",
    "[ ] Have the app already loaded in the browser — no 'loading...' in front of clients",
])

h3("During the demo")
bullet("Narrate what you're doing and why: 'I'm clicking Load Sample Data, which "
       "triggers the full data pipeline...'")
bullet("Point out features that were specifically requested: 'Last demo you asked for "
       "a way to filter by department — here it is in the sidebar.'")
bullet("When something goes wrong (it will), stay calm: 'Interesting — I haven't "
       "seen that before. Let me note it and we'll investigate after the meeting.' Then move on.")
bullet("Do not apologise excessively for imperfections. You are showing "
       "working software, not a finished product.")
bullet("End the demo before you run out of impressive things to show. "
       "Leave them wanting the next sprint.")

h3("After the demo — the follow-up email (send within 2 hours)")
code_block([
    "Subject: Sprint [N] Demo — Summary & Next Steps — Meridian Corp",
    "",
    "Hi [Name],",
    "",
    "Thank you for attending today's demo. Here is a summary:",
    "",
    "WHAT WE SHOWED:",
    "  - [Feature 1]: [brief description]",
    "  - [Feature 2]: [brief description]",
    "",
    "YOUR FEEDBACK:",
    "  - [Exact quote from their feedback]",
    "  - [Action item that came from it]",
    "",
    "SPRINT [N+1] PLAN:",
    "  - [Top 3 items from backlog]",
    "",
    "NEXT DEMO: [Date] at [Time] — same link / location.",
    "",
    "Let me know if you have any questions before then.",
    "",
    "[Your name]",
])

h2("Common Demo Mistakes and How to Avoid Them")

add_table(
    ["Mistake", "What happens", "How to avoid it"],
    [
        ["Showing mockups instead of working software",
         "Client loses trust — 'Where is the real thing?'",
         "Never demo anything that isn't actually running"],
        ["Demoing on production data in front of execs for the first time",
         "Unexpected data quality issues become a distraction",
         "Always test with their data privately before the demo"],
        ["Running too long — showing everything",
         "Audience loses attention, feedback session gets cut",
         "Demo only what was in the sprint goal. 25 minutes max."],
        ["Asking 'Any questions?' and getting silence",
         "No useful feedback, you don't know if they're happy",
         "Ask the three specific feedback questions instead"],
        ["Skipping the demo when a sprint was slow",
         "Missed chance to reset expectations early",
         "Always hold the demo — even to show what you started"],
    ],
    col_widths=[1.6, 1.9, 2.8],
)

divider()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 10: STAKEHOLDER COMMUNICATION CADENCE
# ══════════════════════════════════════════════════════════════════════════════

section_banner("CHAPTER 10   Stakeholder Communication Cadence")

h1("Who Needs to Know What and When")

body("Poor communication is the number one reason consulting engagements fail — "
     "not technical problems. Your client's leadership team needs to feel informed "
     "at all times, especially when things are difficult.")

h2("The Communication Rhythm")

add_table(
    ["Frequency", "Format", "Audience", "Purpose"],
    [
        ["Daily", "Standup (internal, 15 min)", "Your team only",
         "Synchronise, surface blockers"],
        ["Weekly", "Written status update email (1 page)", "Client project sponsor",
         "What happened, what's next, any risks"],
        ["Bi-weekly", "Sprint demo (1 hour)", "All client stakeholders",
         "Show progress, get feedback, update backlog"],
        ["Monthly", "Steering committee presentation (30 min)", "CFO / COO / IT Director",
         "Big picture: are we on track, any scope changes, budget status"],
        ["Ad hoc", "Phone or Slack message (< 5 min)", "Project sponsor",
         "Urgent blockers, decisions needed, scope questions"],
        ["End of engagement", "Final delivery presentation (1 hour)", "All stakeholders",
         "What was built, how to use it, next steps"],
    ],
    col_widths=[0.8, 1.8, 1.7, 2.0],
)

h2("The Weekly Status Update Email")

body("This is your most important communication habit. Send it every Friday before 5pm, "
     "every single week, without being asked. Clients who feel informed do not micromanage.")

code_block([
    "Subject: Meridian Corp — Week [N] Status Update",
    "",
    "STATUS: ON TRACK  /  AT RISK  /  NEEDS ATTENTION  (pick one — be honest)",
    "",
    "COMPLETED THIS WEEK:",
    "  - [Specific item delivered]",
    "  - [Specific item delivered]",
    "",
    "IN PROGRESS (sprint [N]):",
    "  - [Item, % complete]",
    "  - [Item, % complete]",
    "",
    "NEXT WEEK:",
    "  - [What you plan to start/complete]",
    "",
    "RISKS & BLOCKERS:",
    "  - [Risk: description | Impact: High/Med/Low | Mitigation: what you're doing]",
    "  OR: None this week.",
    "",
    "DECISIONS NEEDED FROM CLIENT:",
    "  - [Be specific — vague asks get no response]",
    "  OR: None this week.",
    "",
    "NEXT DEMO: [Date, Time, Location/Link]",
])

callout("Rule:",
        "Never send a status email that says 'everything is fine' when something is wrong. "
        "Surface problems early, with a proposed solution. "
        "Surprises destroy trust. Early warnings build it.",
        bg="FADBD8")

h2("Managing the Difficult Stakeholder")

body("Every engagement has one person who is skeptical, territorial, or actively resistant. "
     "Usually this is the IT Director who wasn't consulted before the engagement started, "
     "or a PM who feels their job is being automated.")

add_table(
    ["Behaviour", "What's behind it", "How to respond"],
    [
        ["'Why wasn't I consulted?'",
         "Feels bypassed — status threat",
         "Schedule a 1:1. Ask for their input on the technical architecture. "
         "Make them a contributor, not an observer."],
        ["'This will never work with our data'",
         "Legitimate concern or self-protection",
         "'Let's test it with a sample of your real data and see what happens.' "
         "Evidence beats argument every time."],
        ["'We already tried this and it failed'",
         "Prior bad experience — burned before",
         "Ask them to tell you what went wrong. Acknowledge it. "
         "Explain specifically how your approach is different."],
        ["Slow email responses / missed meetings",
         "Not a priority for them — or passive resistance",
         "Move decisions up one level. CC the exec sponsor on items "
         "that are blocking delivery."],
    ],
    col_widths=[1.5, 1.7, 3.1],
)

divider()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 11: COMMERCIAL MODEL
# ══════════════════════════════════════════════════════════════════════════════

section_banner("CHAPTER 11   Commercial Model — Pricing, Proposals & Contracts")

h1("How to Price Your Work")

body("Pricing is the most uncomfortable part of starting a consultancy. Here is the "
     "fundamental rule: price on value delivered, not on hours worked. A report that "
     "saves a CFO 5 hours per week is worth far more than the 40 hours it took you to build.")

h2("The Three Pricing Models")

add_table(
    ["Model", "How it works", "When to use", "Typical range"],
    [
        ["Fixed-price project",
         "You quote a flat fee for a defined scope and deliverables",
         "Well-defined projects with clear requirements",
         "$5,000 – $50,000"],
        ["Time and materials",
         "You charge a day rate; client pays for actual time spent",
         "Exploratory or unclear scope; long-term ongoing work",
         "$1,500 – $3,000 / day"],
        ["Retainer",
         "Monthly fee for guaranteed availability and ongoing support",
         "After initial project is live; for maintenance and expansion",
         "$2,000 – $8,000 / month"],
    ],
    col_widths=[1.3, 2.2, 1.7, 1.1],
)

h2("Your Engagement Tiers")

add_table(
    ["Tier", "Name", "Price", "What's included", "Timeline"],
    [
        ["Starter", "AI Reporting Pilot",
         "$7,500",
         "Working dashboard + 1 AI report + PDF export + 30-day support",
         "30 days"],
        ["Standard", "AI Reporting System",
         "$15,000 – $25,000",
         "Full pipeline + dashboard + PDF + data integration + 90-day support + training",
         "60–90 days"],
        ["Enterprise", "AI Reporting Platform",
         "$40,000 – $75,000",
         "Multi-department rollout + custom models + API integration + 6-month support + team training",
         "90–180 days"],
    ],
    col_widths=[0.7, 1.5, 1.0, 3.0, 0.8],
)

h2("The Pilot-First Strategy")

body("The single most effective commercial strategy for a new consultancy is the "
     "paid pilot. Here's why:")
bullet("It lowers the client's perceived risk ('It's only $7,500')")
bullet("It gives you a real engagement to build a case study from")
bullet("Almost every pilot converts to a full engagement if the demo impresses")
bullet("If it doesn't convert, you still got paid and learned about their data")

quote_block("The pilot is not a discount. It is a scoped first phase. "
            "Make sure the client understands: the pilot delivers a working product — "
            "the full engagement expands it.", "Positioning the pilot correctly")

h2("Your Proposal Structure")

code_block([
    "Section 1 — Understanding of your situation (2 paragraphs)",
    "  Use their exact words from discovery. Prove you listened.",
    "",
    "Section 2 — Our proposed approach (1 page)",
    "  3 phases, each with clear deliverable and dates.",
    "",
    "Section 3 — Deliverables (table)",
    "  Exactly what they will receive, with acceptance criteria.",
    "",
    "Section 4 — Timeline (Gantt or sprint table)",
    "  Sprint-by-sprint, with demo dates visible.",
    "",
    "Section 5 — Investment (table)",
    "  Fixed fee per phase. Optional: ROI calculation.",
    "",
    "Section 6 — Our team (bios, 2-3 sentences each)",
    "",
    "Section 7 — Terms (payment, IP ownership, confidentiality)",
    "  Simple: 50% upfront, 50% on delivery. Net 30.",
    "",
    "Appendix — Technical architecture diagram",
])

h2("The Two Non-Negotiable Contract Clauses")

callout("IP ownership:",
        "The client owns all code, data, and outputs delivered under the engagement. "
        "You retain the right to reuse your general methodologies, frameworks, and "
        "non-client-specific components in future engagements.",
        bg="D5F5E3")

callout("Change management:",
        "Any change to scope, timeline, or deliverables must be agreed in writing before "
        "work begins. This protects both parties. Use a simple Change Request template.",
        bg="FADBD8")

divider()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 12: HANDLING OBJECTIONS
# ══════════════════════════════════════════════════════════════════════════════

section_banner("CHAPTER 12   Handling Objections")

h1("Every Objection Is a Question in Disguise")

body("An objection is not a rejection. It is the client telling you what they need "
     "to hear before they can say yes. Each objection has an emotional source and "
     "a rational source. Address both.")

add_table(
    ["Objection", "What they're really asking", "How to respond"],
    [
        ["'It's too expensive'",
         "Is this worth the risk for me?",
         "'What does one delayed project cost your company? Our pilot is $7,500. "
         "If it saves one project from going off the rails, it pays for itself ten times over.'"],
        ["'We don't have the budget right now'",
         "Can this wait?",
         "'When does your budget cycle reset? I'd like to put a proposal in front of you "
         "before that closes. Can we schedule a 30-minute call in the next two weeks?'"],
        ["'We tried AI before and it didn't work'",
         "Can I trust you not to waste my time again?",
         "'Tell me what happened — I want to understand exactly what went wrong. "
         "Our approach is different in [specific ways]. Can I show you in 10 minutes?'"],
        ["'We'd need to involve IT and they're already stretched'",
         "Will this create more work for my team?",
         "'We run everything — your IT team just needs to give us read access to the data. "
         "No infrastructure changes. Typical IT involvement is 2 hours in week one.'"],
        ["'We need to think about it'",
         "I'm not convinced yet — or I need to sell it internally",
         "'Absolutely. What would you need to see to feel confident? "
         "And who else needs to be part of this decision?'"],
        ["'Can we start smaller?'",
         "I want a proof of concept before I commit",
         "'That's exactly what our pilot is designed for. $7,500, 30 days, "
         "working product. If it's not what you expected, you owe us nothing more.'"],
    ],
    col_widths=[1.4, 1.7, 3.2],
)

h2("The 'Not Right Now' Objection")

body("When a client says 'not right now,' most consultants give up. This is a mistake. "
     "'Not right now' means their pain is not yet intense enough — or their calendar is full. "
     "The correct response is a nurture sequence:")

numbered("Send a follow-up email 3 days later with one insight from their data "
         "('I noticed in your data that...')")
numbered("Connect on LinkedIn and share one relevant article per month about AI in reporting")
numbered("Check in 90 days: 'Hi [Name], I wanted to see if anything has changed "
         "on your side...'")
numbered("When their industry gets a major news event (a competitor's project failure, "
         "a new regulation), send a personalised note connecting it to what you do")

callout("Reality:",
        "Many of your best clients will say 'not right now' twice before they say yes. "
        "The ones who close quickly were already at the decision point. The others need "
        "to see the problem get worse before they act. Your job is to be there when they're ready.",
        bg="D6EAF8")

divider()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 13: MINDSET & CONFIDENCE
# ══════════════════════════════════════════════════════════════════════════════

section_banner("CHAPTER 13   Mindset & Confidence — The Inner Game of Consulting")

h1("The Confidence You Need Is Built, Not Born")

body("Every successful consultant you have ever admired felt exactly what you feel now: "
     "imposter syndrome. The difference is they acted anyway. Confidence is not the "
     "absence of doubt — it is the decision to move forward in spite of it.")

h2("What You Actually Know That Others Don't")

body("Stop and inventory what you have built:")
bullet("A working data pipeline: CSV ingestion, medallion architecture, DuckDB")
bullet("An AI reporting engine using Claude, with anti-hallucination design")
bullet("A live dashboard that a non-technical executive can use in 60 seconds")
bullet("A PDF export that looks professional enough for board meetings")
bullet("A demo reset that works cold and reliably")
bullet("21 automated tests that prove the system works")
bullet("The ability to explain your system to a CFO, an IT director, and a PM "
       "in language each of them understands")

callout("Reframe:",
        "Most companies have none of this. Most of their data sits in spreadsheets "
        "nobody reads. You have built something real. You are ahead of 95% of people "
        "who call themselves 'AI consultants.'",
        bg="D5F5E3")

h2("Managing Imposter Syndrome in Client Meetings")

add_table(
    ["Moment", "What fear says", "What truth says"],
    [
        ["Before the first meeting",
         "'They'll see I don't have 20 years of experience'",
         "'They hired me because of what I can do, not my title'"],
        ["When asked a question you can't answer",
         "'I should know this — I'm supposed to be the expert'",
         "'Great question. I want to give you an accurate answer — "
         "let me research and get back to you by end of day.'"],
        ["When the demo glitches",
         "'This is embarrassing, they'll lose confidence'",
         "'Every real system has rough edges in early demos. "
         "What matters is how you handle it.'"],
        ["When they push back on price",
         "'Maybe I'm charging too much'",
         "'They're negotiating. That's normal. Stay calm and anchor to value.'"],
    ],
    col_widths=[1.4, 2.0, 2.9],
)

h2("Three Habits That Build Confidence Over Time")

h3("1. Keep a wins journal")
body("Every day, write one thing that went well. A good question you asked. A demo that "
     "landed. A client who replied positively to your email. After 90 days, you will "
     "have a record of 90 real wins to draw on when doubt hits.")

h3("2. Do a debrief after every client meeting")
body("Three questions: What went well? What would I do differently? What did I learn? "
     "This converts experience into skill. Consultants who debrief improve. "
     "Consultants who don't repeat the same mistakes.")

h3("3. Teach what you know")
body("Write a LinkedIn post about something you built. Record a 5-minute video "
     "explaining medallion architecture in plain English. Run a free webinar. "
     "Teaching forces clarity, and clarity builds confidence. "
     "It also generates inbound leads.")

divider()


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 14: 90-DAY LAUNCH PLAN
# ══════════════════════════════════════════════════════════════════════════════

section_banner("CHAPTER 14   90-Day Launch Plan for Your Consulting Company")

h1("Your First 90 Days")

body("You do not need a perfect company before you can start. You need a product, "
     "a pitch, and one meeting. Here is the sequence that gets you from zero to first "
     "paying client in 90 days.")

h2("Month 1 — Foundation (Days 1–30)")

add_table(
    ["Week", "Action", "Outcome"],
    [
        ["Week 1",
         "Register your business entity. Open a business bank account. "
         "Get a professional email (yourname@yourcompany.com).",
         "You are legally a company"],
        ["Week 1",
         "Set up LinkedIn company page. Write your personal LinkedIn headline: "
         "'AI Reporting Consultant | I turn project data into executive insights'",
         "Professional presence"],
        ["Week 2",
         "Write your one-page capability summary (PDF). "
         "Design a simple logo (Canva is fine). Build a 1-page website.",
         "You have leave-behind materials"],
        ["Week 2",
         "List 20 people in your network who manage projects or lead teams. "
         "These are your first outreach targets.",
         "Prospect list"],
        ["Week 3",
         "Reach out to the top 10 on your list. Not a sales pitch — "
         "a 'coffee chat' or 'I built something I'd love your feedback on.'",
         "5 conversations booked"],
        ["Week 4",
         "Run your first discovery conversations. Show the demo. "
         "Even if they're not a buyer, ask: 'Who do you know who deals with this problem?'",
         "Referrals and one qualified lead"],
    ],
    col_widths=[0.7, 3.5, 2.1],
)

h2("Month 2 — First Client (Days 31–60)")

add_table(
    ["Week", "Action", "Outcome"],
    [
        ["Week 5",
         "Send a personalized proposal to your qualified lead. "
         "Follow up every 3 business days until you get a decision.",
         "Proposal submitted"],
        ["Week 6",
         "Close your first paid pilot ($7,500). "
         "Send the contract. Collect 50% upfront.",
         "First revenue"],
        ["Week 7–8",
         "Deliver sprint 1 of the pilot. Run your first demo ceremony. "
         "Send your first weekly status update.",
         "Delivery underway"],
    ],
    col_widths=[0.7, 3.5, 2.1],
)

h2("Month 3 — Expand and Systematise (Days 61–90)")

add_table(
    ["Week", "Action", "Outcome"],
    [
        ["Week 9",
         "Complete the pilot. Collect the final 50%. "
         "Run the final demo. Ask for a testimonial and a referral.",
         "First engagement complete"],
        ["Week 10",
         "Convert pilot to full engagement — or close second client from referral.",
         "Pipeline building"],
        ["Week 11–12",
         "Document your delivery process. Build reusable templates: "
         "proposal template, status update template, sprint planning agenda. "
         "What took 10 hours to do the first time should take 2 hours the second time.",
         "Scalable process"],
    ],
    col_widths=[0.7, 3.5, 2.1],
)

h2("The Only KPIs That Matter in Month 1")

add_table(
    ["KPI", "Target", "Why it matters"],
    [
        ["Discovery conversations held", "5 per month",
         "Pipeline is built on conversations, not proposals"],
        ["Proposals sent", "2 per month",
         "You can't close what you haven't proposed"],
        ["Demo sessions delivered", "3 per month",
         "The demo closes more deals than any sales technique"],
        ["LinkedIn posts published", "2 per week",
         "Inbound leads take 3-6 months to build — start now"],
        ["Follow-ups sent", "100% of open proposals",
         "Most consultants stop following up too soon"],
    ],
    col_widths=[2.0, 1.4, 2.9],
)

h2("Final Thought")

quote_block(
    "The gap between where you are and where you want to be is not skill. "
    "You have built the system. The gap is reps. Every meeting you take, "
    "every demo you run, every proposal you send makes the next one easier. "
    "The only way to get your first client is to go get your first client.",
    "Your reminder on hard days"
)

callout("You are ready.",
        "You have a working product, a clear pitch, a delivery methodology, "
        "and this playbook. Most people who call themselves AI consultants have none of these. "
        "Go take the first meeting.",
        bg="D5F5E3")

doc.add_paragraph()

add_table(
    ["What you have built", "Status"],
    [
        ["Bronze/silver/gold data pipeline", "Done"],
        ["DuckDB storage layer", "Done"],
        ["KPI aggregation engine", "Done"],
        ["Anti-hallucination AI reporting", "Done"],
        ["Streamlit executive dashboard", "Done"],
        ["ReportLab PDF export", "Done"],
        ["Demo reset (cold-start reliable)", "Done"],
        ["21 automated tests", "Done"],
        ["Complete daily learning guides (Days 1–7)", "Done"],
        ["This consulting playbook", "Done"],
        ["Your first client meeting", "Next step"],
    ],
    col_widths=[3.5, 2.8],
)

# ── Save ──────────────────────────────────────────────────────────────────────

output_path = "docs/consulting-playbook.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
