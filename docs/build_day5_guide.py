"""Run this script to generate day-5-learning-guide.docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Style helpers ──────────────────────────────────────────────────────────────

def set_font(run, bold=False, italic=False, size=None, color=None, font_name=None):
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, bold=True, size=20, color=(44, 62, 80))
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, bold=True, size=15, color=(44, 62, 80))
    return p


def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, bold=True, size=12, color=(52, 73, 94))
    return p


def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def code_block(lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 30, 30)
        if i < len(lines) - 1:
            run.add_break()
    doc.add_paragraph()


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def callout(label, text, bg="FFF3CD"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg)
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    r1 = p.add_run(f"{label}  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    doc.add_paragraph()


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "2C3E50")
        tcPr.append(shd)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        fill = "F8F9FA" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("─" * 80)
    run.font.color.rgb = RGBColor(189, 195, 199)
    run.font.size = Pt(8)


# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run("AI Executive Reporting System")
set_font(run, bold=True, size=26, color=(44, 62, 80))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Day 5 Learning Guide")
set_font(run2, bold=True, size=20, color=(52, 152, 219))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Streamlit UI — Session State, KPI Cards, Project Table, AI Summary")
set_font(run3, italic=True, size=13, color=(127, 140, 141))

doc.add_paragraph()
divider()
doc.add_paragraph()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run(
    "This guide explains how Streamlit's execution model works, how session_state persists\n"
    "data between interactions, and how each UI component was built."
)
set_font(run4, size=11, color=(52, 73, 94))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — WHAT WE BUILT
# ══════════════════════════════════════════════════════════════════════════════

h1("1. What Day 5 Built")

body(
    "Days 1–4 built a complete data and AI engine running entirely in Python. "
    "Day 5 adds the browser-based interface that makes all of it accessible without "
    "writing a line of code or opening a terminal."
)

body("The browser UI after Day 5:")

code_block([
    "Browser",
    "  ├── Sidebar: Upload CSV  or  Load Sample Data button",
    "  ├── KPI Cards row: Total Projects | On Track% | At Risk/Delayed% | Budget Variance",
    "  ├── Project Table: all 15 projects, sortable, with checkboxes for is_late/is_over_budget",
    "  ├── AI Summary: Generate Report button + markdown output",
    "  └── Export: PDF download button (wired on Day 6)",
])

body("Four new files were created:")
bullet("app/components/kpi_cards.py — four st.metric() cards")
bullet("app/components/project_table.py — styled st.dataframe() with column config")
bullet("app/components/ai_summary.py — Generate Report button + markdown display")
bullet("app/main.py — Streamlit entry point, session state, page layout")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — STREAMLIT EXECUTION MODEL
# ══════════════════════════════════════════════════════════════════════════════

h1("2. How Streamlit Works — The Execution Model")

body(
    "Streamlit's execution model is different from every other web framework. "
    "Understanding it is the key to writing correct Streamlit apps."
)

h2("2.1 Re-runs Top to Bottom")

callout(
    "The core rule:",
    "Every time the user interacts with any widget (clicks a button, uploads a file, "
    "changes a slider), Streamlit re-runs your entire Python script from the very top. "
    "Every variable is reset. Every import re-executes. The whole file starts fresh.",
    bg="E3F2FD",
)

code_block([
    "# This is your entire Streamlit app — simplified:",
    "",
    "x = 0                            # x is reset to 0 on EVERY run",
    "if st.button('Add 1'):",
    "    x = x + 1                    # This only runs during the click run",
    "st.write(x)                      # Always shows 1 — never accumulates!",
    "",
    "# On first load:    button shown, x=0 displayed",
    "# On button click:  whole script runs again, x starts at 0, becomes 1, displayed",
    "# On second click:  whole script runs AGAIN, x starts at 0 again, becomes 1 again",
    "# Result:           x is always 1, never 2 — it resets every run",
])

body(
    "This is why a counter built with a plain variable never works in Streamlit. "
    "You need session_state to persist values between runs."
)

h2("2.2 st.session_state — Persistence Between Runs")

body(
    "st.session_state is a dictionary-like object that survives across re-runs "
    "for the duration of the browser session. Values you store in it are still there "
    "the next time the script re-runs."
)

code_block([
    "# Correct pattern for a counter:",
    "",
    "if 'count' not in st.session_state:",
    "    st.session_state['count'] = 0    # Initialize only if not set",
    "",
    "if st.button('Add 1'):",
    "    st.session_state['count'] += 1   # Increment persists between runs",
    "",
    "st.write(st.session_state['count'])  # Shows accumulated count correctly",
    "",
    "# On first load:    count initialized to 0, shown",
    "# On first click:   count becomes 1, shown",
    "# On second click:  count becomes 2, shown (did not reset!)",
])

body(
    "In our app, st.session_state stores the pipeline result — the entire dict "
    "containing kpi_summary and ai_report. Once the pipeline runs, the result "
    "persists in session_state so the KPI cards and table remain visible even "
    "when the user clicks 'Generate Report' (which triggers another re-run)."
)

code_block([
    "# How we use it in main.py:",
    "",
    "# Initialize:",
    "if 'pipeline_result' not in st.session_state:",
    "    st.session_state['pipeline_result'] = None",
    "",
    "# Store result after pipeline runs:",
    "st.session_state['pipeline_result'] = run_pipeline(csv_path)",
    "",
    "# Read in display section:",
    "if st.session_state['pipeline_result']:",
    "    result = st.session_state['pipeline_result']",
    "    # render KPI cards, table, etc.",
])

h2("2.3 st.rerun() — Force a Fresh Render")

body(
    "After the pipeline runs and stores its result in session_state, we call st.rerun(). "
    "This immediately triggers another top-to-bottom re-run of the script, "
    "which causes Streamlit to render the KPI cards and table with the new data."
)

code_block([
    "if load_sample:",
    "    st.session_state['pipeline_result'] = run_pipeline('data/input/project_data.csv')",
    "    st.rerun()    ← forces immediate re-render to show the results",
    "",
    "# Without st.rerun():",
    "#   The pipeline result is stored in session_state",
    "#   But the current run continues executing below the if block",
    "#   The 'if st.session_state[pipeline_result]' check happens in the SAME run",
    "#   so results ARE shown — but the sidebar success message lags by one click",
    "#",
    "# With st.rerun():",
    "#   Execution stops immediately",
    "#   A fresh run starts from the top",
    "#   session_state already has the result",
    "#   Everything renders cleanly in the fresh run",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — app/main.py STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════

h1("3. File: app/main.py — The Entry Point")

body(
    "main.py is the only file Streamlit runs directly. "
    "It orchestrates the four UI sections: sidebar, KPI cards, project table, AI summary."
)

h2("3.1 Page Configuration")

code_block([
    "st.set_page_config(",
    "    page_title='AI Executive Reporting',",
    "    page_icon='📊',",
    "    layout='wide'",
    ")",
])

body(
    "st.set_page_config() must be the first Streamlit call in the file. "
    "layout='wide' uses the full browser width instead of the default narrow column, "
    "which is essential for the four-column KPI cards to look right."
)

h2("3.2 The Sidebar")

code_block([
    "with st.sidebar:",
    "    st.header('Data Source')",
    "    uploaded = st.file_uploader('Upload CSV', type='csv')",
    "    load_sample = st.button('Load Sample Data', type='primary')",
    "",
    "    if st.session_state.get('pipeline_result'):",
    "        st.success('Data loaded')",
    "        total = st.session_state['pipeline_result']['kpi_summary']['total_projects']",
    "        st.caption(f'{total} projects in pipeline')",
])

body(
    "Everything inside 'with st.sidebar:' renders in the left panel. "
    "st.file_uploader() returns None when no file is uploaded, or a file-like object when one is. "
    "st.button() returns True on the run triggered by the click, False on all other runs. "
    "This is why 'if load_sample:' correctly triggers only once per click."
)

h2("3.3 Handling File Upload with tempfile")

code_block([
    "if uploaded is not None:",
    "    with tempfile.NamedTemporaryFile(delete=False, suffix='.csv') as f:",
    "        f.write(uploaded.read())",
    "        tmp_path = f.name",
    "    st.session_state['pipeline_result'] = run_pipeline(tmp_path)",
    "    st.rerun()",
])

body(
    "uploaded is a Streamlit UploadedFile object — it has a .read() method but is not a file path. "
    "The pipeline's csv_loader expects a file path string, not a file object. "
    "tempfile.NamedTemporaryFile creates a real file on disk with a generated name, "
    "writes the uploaded bytes into it, and gives us the path. "
    "delete=False prevents the file from being deleted when the with-block exits — "
    "the pipeline needs to open it after the block closes."
)

callout(
    "Why not just pass the UploadedFile directly?",
    "csv_loader.py uses pd.read_csv(file_path), which expects a string path. "
    "We could rewrite it to accept file objects, but using a temp file is simpler "
    "and keeps csv_loader unchanged. This is a clean boundary: the UI adapter "
    "(main.py) handles the conversion, and the pipeline code stays pure.",
    bg="E3F2FD",
)

h2("3.4 Lazy Imports Inside Conditionals")

code_block([
    "if load_sample:",
    "    from src.pipeline import run_pipeline   ← imported only when needed",
    "    with st.spinner('Running data pipeline...'):",
    "        st.session_state['pipeline_result'] = run_pipeline(...)",
])

body(
    "Imports inside if-blocks are called lazy imports. "
    "In Streamlit, every re-run re-executes every top-level import. "
    "If run_pipeline was imported at the top of the file, it would be imported "
    "on every single re-run — even when just scrolling or resizing the browser. "
    "By importing inside the if-block, the import only executes when that branch runs."
)
body(
    "For heavy modules, lazy imports can meaningfully speed up page load. "
    "For this project the difference is small, but it is the correct pattern."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — kpi_cards.py
# ══════════════════════════════════════════════════════════════════════════════

h1("4. File: app/components/kpi_cards.py")

body(
    "This file renders the four metric cards at the top of the page. "
    "st.metric() is Streamlit's purpose-built widget for showing a KPI with a delta indicator."
)

h2("4.1 Four-Column Layout")

code_block([
    "col1, col2, col3, col4 = st.columns(4)",
    "",
    "with col1:",
    "    st.metric(label='Total Projects', value=kpi_summary['total_projects'])",
])

body(
    "st.columns(4) creates four equal-width columns across the page. "
    "Anything rendered inside 'with col1:' appears in the first column. "
    "The four columns are created once and the content fills them left to right."
)

h2("4.2 st.metric() Parameters")

add_table(
    ["Parameter", "What it does", "Example"],
    [
        ["label",       "Title above the number",                  "'Total Projects'"],
        ["value",       "The main large number displayed",         "15  or  '40.0%'"],
        ["delta",       "Small indicator below the value",         "'+13.5% vs plan'"],
        ["delta_color", "Color of the delta: normal/inverse/off",  "'inverse' for budget overrun"],
    ],
    col_widths=[1.4, 2.8, 2.4],
)

h2("4.3 delta_color='inverse' — When Up is Bad")

code_block([
    "# Budget variance card:",
    "st.metric(",
    "    label='Budget Variance',",
    "    value=f'${abs(variance):,.0f}',",
    "    delta=f'{sign}{variance_pct:.1f}% vs plan',",
    "    delta_color='inverse' if variance > 0 else 'normal'",
    ")",
])

body(
    "By default, a positive delta shows in green (good) and negative in red (bad). "
    "For budget variance, this logic is inverted: being over budget (+) is bad, "
    "under budget (-) is good. delta_color='inverse' flips the colors: "
    "positive delta shows red, negative delta shows green."
)

add_table(
    ["delta_color value", "Positive delta color", "Negative delta color", "Use when"],
    [
        ["'normal'",  "Green (good)",  "Red (bad)",    "Higher is better: revenue, on-track %"],
        ["'inverse'", "Red (bad)",     "Green (good)", "Lower is better: risk %, budget overrun"],
        ["'off'",     "Gray",          "Gray",         "Neutral: no judgment on direction"],
    ],
    col_widths=[1.5, 1.5, 1.5, 2.1],
)

body(
    "In the On Track card, delta_color='normal' is used because higher on-track percentage is good. "
    "In the At Risk/Delayed card, delta_color='inverse' is used because a higher at-risk percentage is bad."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — project_table.py
# ══════════════════════════════════════════════════════════════════════════════

h1("5. File: app/components/project_table.py")

body(
    "This file renders the project portfolio as a styled, sortable dataframe. "
    "st.dataframe() with column_config gives precise control over how each column appears."
)

h2("5.1 Column Selection and Ordering")

code_block([
    "display_cols = [",
    "    'project_name', 'department', 'status', 'risk_level',",
    "    'budget_variance_pct', 'schedule_delay_days',",
    "    'is_over_budget', 'is_late', 'owner'",
    "]",
    "display_cols = [c for c in display_cols if c in df.columns]",
    "df = df[display_cols]",
])

body(
    "The gold table has more columns than we want to show (planned_cost, actual_cost, etc.). "
    "This pattern defines the display order and filters to only existing columns. "
    "The list comprehension [c for c in display_cols if c in df.columns] guards against "
    "missing columns — if the schema changes, the table still renders rather than crashing."
)

h2("5.2 column_config — Custom Column Display")

code_block([
    "st.dataframe(",
    "    df,",
    "    use_container_width=True,",
    "    column_config={",
    "        'project_name':        st.column_config.TextColumn('Project'),",
    "        'budget_variance_pct': st.column_config.NumberColumn(",
    "            'Budget Var %', format='%.1f%%'",
    "        ),",
    "        'schedule_delay_days': st.column_config.NumberColumn('Delay (days)'),",
    "        'is_over_budget':      st.column_config.CheckboxColumn('Over Budget'),",
    "        'is_late':             st.column_config.CheckboxColumn('Late'),",
    "    },",
    "    hide_index=True,",
    ")",
])

add_table(
    ["Column config type", "What it renders", "Key parameters"],
    [
        ["TextColumn",     "Plain text",                          "'label' renames the header"],
        ["NumberColumn",   "Numeric with optional format string", "format='%.1f%%' → '32.1%'"],
        ["CheckboxColumn", "Green checkmark (True) or empty (False)", "'label' renames header"],
        ["SelectboxColumn","Dropdown (editable tables)",           "options=[] for choices"],
    ],
    col_widths=[1.8, 2.4, 2.4],
)

body(
    "CheckboxColumn is what makes is_late and is_over_budget appear as visual checkmarks "
    "instead of 'True'/'False' text. Clients immediately understand a checkmark column — "
    "they do not need to read a boolean."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — ai_summary.py
# ══════════════════════════════════════════════════════════════════════════════

h1("6. File: app/components/ai_summary.py")

body(
    "This file renders the AI report section: a Generate button and the report text below it. "
    "The on_generate callback pattern keeps the component decoupled from the pipeline."
)

h2("6.1 The Full File")

code_block([
    "def render_ai_summary(ai_report: str, on_generate) -> None:",
    "    col1, col2 = st.columns([1, 4])",
    "",
    "    with col1:",
    "        if st.button('Generate Report', type='primary'):",
    "            on_generate()",
    "",
    "    if not ai_report or ai_report.startswith('['):",
    "        st.info('Click Generate Report to create an AI-written executive summary.')",
    "    else:",
    "        from datetime import datetime",
    "        st.caption(f'Generated at {datetime.now().strftime(\"%Y-%m-%d %H:%M\")}')",
    "        st.markdown(ai_report)",
])

h2("6.2 The Callback Pattern")

body(
    "on_generate is a function passed in from main.py as a parameter. "
    "When the button is clicked, ai_summary.py calls on_generate() — "
    "it does not know or care what that function does."
)

code_block([
    "# In main.py — the callback is defined here:",
    "def on_generate():",
    "    from src.llm.prompt_builder import build_prompt",
    "    from src.llm.claude_client import generate_report",
    "    with st.spinner('Generating AI report... (5-10 seconds)'):",
    "        prompt = build_prompt(kpi)",
    "        st.session_state['pipeline_result']['ai_report'] = generate_report(prompt)",
    "",
    "# Passed in when calling the component:",
    "render_ai_summary(result['ai_report'], on_generate)",
])

body(
    "Why use a callback instead of calling generate_report inside ai_summary.py? "
    "Decoupling. ai_summary.py is a UI component — it should not know about the LLM layer. "
    "By passing on_generate as a parameter, you can test ai_summary.py with a mock callback "
    "and swap the AI implementation without touching the component."
)

h2("6.3 Detecting the Fallback String")

code_block([
    "if not ai_report or ai_report.startswith('['):",
    "    st.info('Click Generate Report...')",
])

body(
    "ai_report is None when the pipeline has not generated a report yet. "
    "It starts with '[' when the fallback message was returned (e.g. '[Report generation unavailable...]'). "
    "Both cases show the 'Click Generate Report' info box. "
    "Any other string is rendered as markdown — the five-section AI report."
)

h2("6.4 Columns for Button Layout")

code_block([
    "col1, col2 = st.columns([1, 4])",
    "",
    "# [1, 4] means:",
    "# col1 gets 1/5 of the width (the button)",
    "# col2 gets 4/5 of the width (available for future content)",
])

body(
    "st.columns([1, 4]) creates two columns with a 1:4 width ratio. "
    "The button fits neatly in the narrow column without stretching across the full page width. "
    "Proportional column widths are specified as a list of integers (ratios) rather than pixel values."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — st.spinner AND USER FEEDBACK
# ══════════════════════════════════════════════════════════════════════════════

h1("7. User Feedback — st.spinner()")

body(
    "The pipeline takes about 0.1 seconds. The AI report generation takes 3–8 seconds. "
    "Without feedback, the user sees a frozen screen and does not know if anything is happening. "
    "st.spinner() shows a rotating animation with a message while a block of code runs."
)

code_block([
    "with st.spinner('Running data pipeline...'):",
    "    st.session_state['pipeline_result'] = run_pipeline(csv_path)",
    "",
    "# The spinner shows during run_pipeline()",
    "# Disappears when the with-block exits",
])

body(
    "st.spinner() is a context manager. "
    "The spinner appears when the with-block starts and disappears when it exits, "
    "regardless of whether the code succeeded or raised an exception."
)

add_table(
    ["Widget", "What it shows", "When to use"],
    [
        ["st.spinner(msg)", "Rotating animation + message (blocks UI)", "Long operations: pipeline, API calls"],
        ["st.progress(n)",  "Progress bar 0.0–1.0",                     "When you can report partial progress"],
        ["st.info(msg)",    "Blue info box",                             "Static message: 'click X to start'"],
        ["st.success(msg)", "Green success box",                         "After an operation completes"],
        ["st.error(msg)",   "Red error box",                             "When something fails"],
        ["st.warning(msg)", "Yellow warning box",                        "Non-fatal issues the user should know"],
    ],
    col_widths=[1.8, 2.4, 2.4],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — THE FULL PAGE LAYOUT FLOW
# ══════════════════════════════════════════════════════════════════════════════

h1("8. The Complete Page Flow")

body("Here is the exact sequence main.py executes on every re-run:")

code_block([
    "RUN START (triggered by any user interaction)",
    "",
    "1. load_dotenv()                          — load .env every run",
    "2. st.set_page_config(...)                — set browser tab title + layout",
    "3. st.title(...), st.caption(...)         — render page title",
    "4. with st.sidebar: ...                   — render sidebar widgets",
    "   → uploaded = st.file_uploader(...)     — returns None or file object",
    "   → load_sample = st.button(...)         — returns True on click run, else False",
    "5. session_state initialization            — set 'pipeline_result' = None if not set",
    "6. if load_sample: ...                    — run pipeline with sample CSV",
    "7. if uploaded is not None: ...           — run pipeline with uploaded file",
    "8. if st.session_state['pipeline_result']:",
    "   → render_kpi_cards(kpi)               — KPI cards",
    "   → render_project_table(...)           — project table",
    "   → render_ai_summary(..., on_generate) — AI section",
    "   → Export PDF button",
    "9. else:",
    "   → st.info('Use the sidebar...')       — empty state message",
    "",
    "RUN END",
])

body(
    "On the very first page load: steps 1–5 run, load_sample=False, uploaded=None, "
    "pipeline_result=None, so step 9 shows the empty state message."
)
body(
    "When 'Load Sample Data' is clicked: steps 1–5 run, load_sample=True, "
    "step 6 runs the pipeline, stores result in session_state, calls st.rerun(). "
    "st.rerun() restarts from step 1. Now pipeline_result exists, so step 8 renders everything."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — KEY CONCEPTS
# ══════════════════════════════════════════════════════════════════════════════

h1("9. Key Concepts to Remember")

add_table(
    ["Concept", "What it means", "Why it matters"],
    [
        ["Streamlit re-run model",   "Entire script reruns on every interaction",             "Variables reset — cannot use plain Python state"],
        ["st.session_state",         "Dict that persists across re-runs in a session",        "The only way to store pipeline results between runs"],
        ["st.rerun()",               "Force an immediate fresh re-run",                       "Ensures UI renders with updated session_state"],
        ["st.columns([1, 4])",       "Create columns with proportional widths",               "Controls button width and layout proportions"],
        ["delta_color='inverse'",    "Flip green/red for metrics where higher is worse",      "Budget overrun and at-risk% should show red when high"],
        ["CheckboxColumn",           "Render True/False as visual checkmarks",                "Clients read checkmarks faster than boolean text"],
        ["Lazy imports",             "Import inside if-blocks, not at top level",             "Avoids re-importing heavy modules on every re-run"],
        ["tempfile for uploads",     "Write UploadedFile bytes to a real temp file",          "csv_loader expects a path string, not a file object"],
        ["Callback parameter",       "Pass on_generate as a function to the component",       "Decouples UI from AI layer for testability"],
        ["st.spinner() context mgr", "Show spinner for the duration of the with-block",       "Essential UX: users need feedback during 5s AI call"],
    ],
    col_widths=[1.9, 2.3, 2.4],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — MANUAL VERIFICATION
# ══════════════════════════════════════════════════════════════════════════════

h1("10. How to Verify the UI Manually")

body("Run the app, then verify each step:")

code_block([
    "# Make sure sample data exists:",
    "python demo/sample_data.py",
    "",
    "# Start the UI:",
    "streamlit run app/main.py",
    "# Opens http://localhost:8501 automatically",
])

add_table(
    ["Step", "Action", "Expected result"],
    [
        ["1", "Page loads",                      "Empty state message: 'Use the sidebar to load data'"],
        ["2", "Click 'Load Sample Data'",         "Spinner shows, then KPI cards appear: 15 projects"],
        ["3", "Check KPI card values",            "On Track ~26.7%, At Risk/Delayed ~73.3%, Budget Variance ~$183K"],
        ["4", "Scroll project table",             "15 rows visible, is_late and is_over_budget show checkmarks"],
        ["5", "Click column header to sort",      "Table re-sorts by that column"],
        ["6", "Click 'Generate Report'",          "Spinner shows '5-10 seconds', then AI report appears below"],
        ["7", "Check report has 5 sections",      "## Executive Summary through ## Recommended Actions visible"],
        ["8", "Click 'Export PDF'",               "Error (expected — Day 6 builds pdf_generator.py)"],
        ["9", "Upload a CSV via file uploader",   "Pipeline reruns with uploaded file, KPI cards update"],
    ],
    col_widths=[0.4, 2.4, 3.8],
)

h1("11. What Day 6 Builds On This")

body("Day 5 left one stub: the 'Export PDF' button calls src/reporting/pdf_generator.py which does not exist yet. Day 6 will:")

bullet("Create src/reporting/pdf_generator.py — ReportLab 5-section PDF: header, executive summary, KPI table, risk table, AI recommendations")
bullet("Create demo/demo_reset.py — drops DuckDB tables, regenerates sample CSV, reruns data pipeline; completes in under 30 seconds")
bullet("Wire the PDF download button so clicking Export PDF produces a downloadable file")
bullet("Checkpoint: click Export PDF → PDF file downloads with all 5 sections")

body(
    "After Day 6, the demo is complete: data in → KPI dashboard → AI report → PDF export → reset for next client meeting."
)

divider()

p_final = doc.add_paragraph()
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_final.paragraph_format.space_before = Pt(20)
run_final = p_final.add_run("Day 5 complete — browser UI running, 19 tests still passing.")
set_font(run_final, bold=True, size=12, color=(39, 174, 96))

# ── Save ───────────────────────────────────────────────────────────────────────
out_path = "docs/day-5-learning-guide.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
