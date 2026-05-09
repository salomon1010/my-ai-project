"""Run this script to generate day-6-learning-guide.docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Style helpers ──────────────────────────────────────────────────────────────

def set_font(run, bold=False, italic=False, size=None, color=None, font_name=None):
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, bold=True, size=20, color=(44, 62, 80))
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, bold=True, size=15, color=(44, 62, 80))
    return p


def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, bold=True, size=12, color=(52, 73, 94))
    return p


def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def code_block(lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 30, 30)
        if i < len(lines) - 1:
            run.add_break()
    doc.add_paragraph()


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def callout(label, text, bg="FFF3CD"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg)
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    r1 = p.add_run(f"{label}  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    doc.add_paragraph()


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "2C3E50")
        tcPr.append(shd)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        fill = "F8F9FA" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("─" * 80)
    run.font.color.rgb = RGBColor(189, 195, 199)
    run.font.size = Pt(8)


# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run("AI Executive Reporting System")
set_font(run, bold=True, size=26, color=(44, 62, 80))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Day 6 Learning Guide")
set_font(run2, bold=True, size=20, color=(52, 152, 219))

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("PDF Export + Demo Reset — ReportLab, Flowables, and Demo Reliability")
set_font(run3, italic=True, size=13, color=(127, 140, 141))

doc.add_paragraph()
divider()
doc.add_paragraph()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run(
    "This guide explains ReportLab's flowable/story architecture, every TableStyle command,\n"
    "the AI section parser, canvas page-number callbacks, and the demo reset script."
)
set_font(run4, size=11, color=(52, 73, 94))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — WHAT WE BUILT
# ══════════════════════════════════════════════════════════════════════════════

h1("1. What Day 6 Built")

body(
    "Day 5 left the 'Export PDF' button wired to a file that did not exist yet. "
    "Day 6 completes the loop: clicking that button produces a professionally formatted, "
    "multi-page PDF that a client can read without ever opening the web app."
)

body("The complete system after Day 6:")

code_block([
    "Click 'Export PDF' in the browser",
    "  → generate_pdf(kpi_summary, ai_report, output_path)",
    "      → _parse_ai_sections(ai_report)  — split AI text into 5 keyed sections",
    "      → Build story list: cover, exec summary, KPI table, risk table, AI sections",
    "      → doc.build(story, onFirstPage=page_num, onLaterPages=page_num)",
    "      → PDF saved to outputs/reports/report_YYYYMMDD.pdf",
    "  → st.download_button() — browser downloads the file",
])

body("Two new files were created:")
bullet("src/reporting/pdf_generator.py — ReportLab 3-page PDF with 5 content sections")
bullet("demo/demo_reset.py — drops DuckDB tables, regenerates CSV, reruns data pipeline in < 1s")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — REPORTLAB'S FLOWABLE / STORY PATTERN
# ══════════════════════════════════════════════════════════════════════════════

h1("2. ReportLab's Flowable / Story Pattern")

body(
    "ReportLab is a pure-Python PDF library. Its core concept is the 'story': "
    "a Python list of flowable objects. Each flowable knows how to draw itself and "
    "how much space it needs. The document engine places them on pages, "
    "automatically handling page breaks when content overflows."
)

h2("2.1 What a Story Looks Like")

code_block([
    "story = []",
    "",
    "story.append(Paragraph('AI Executive Portfolio Report', header_style))",
    "story.append(Paragraph('Acme Corp · May 08, 2026', caption_style))",
    "story.append(Spacer(1, 0.2 * inch))",
    "story.append(Paragraph('Executive Summary', subheader_style))",
    "story.append(Paragraph('Portfolio text here...', body_style))",
    "story.append(PageBreak())",
    "story.append(Paragraph('KPI Dashboard', subheader_style))",
    "story.append(kpi_table)   # a Table object",
    "",
    "doc.build(story)          # ReportLab places all flowables onto pages",
])

body(
    "You build the entire document as a list, then call doc.build() once. "
    "ReportLab handles pagination: if a Paragraph does not fit at the bottom of a page, "
    "it automatically continues on the next page. PageBreak() forces a new page explicitly."
)

h2("2.2 The Flowable Types We Used")

add_table(
    ["Flowable", "What it renders", "Key parameters"],
    [
        ["Paragraph(text, style)", "A block of styled text", "text can contain basic HTML like <br/>, <b>"],
        ["Spacer(width, height)",  "Empty vertical space",   "Spacer(1, 0.2*inch) — width ignored, height matters"],
        ["Table(data, colWidths)", "A grid of cells",        "data = list of lists; colWidths in inches"],
        ["PageBreak()",           "Forces a new page",       "No parameters needed"],
    ],
    col_widths=[2.3, 2.3, 2.0],
)

callout(
    "The key insight:",
    "You never tell ReportLab where things go on the page. "
    "You just append flowables to the story list in the order you want them. "
    "The engine figures out the layout. This is fundamentally different from HTML/CSS "
    "where you specify positions.",
    bg="E8F5E9",
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — PARAGRAPH STYLES
# ══════════════════════════════════════════════════════════════════════════════

h1("3. Paragraph Styles — ParagraphStyle and getSampleStyleSheet")

body(
    "Every Paragraph needs a style that defines font size, color, spacing, and alignment. "
    "ReportLab provides a set of built-in styles via getSampleStyleSheet(), "
    "and you create custom styles by inheriting from those base styles."
)

h2("3.1 getSampleStyleSheet()")

code_block([
    "styles = getSampleStyleSheet()",
    "",
    "# Built-in styles available:",
    "styles['Normal']     — base style, 10pt Helvetica",
    "styles['BodyText']   — like Normal but with paragraph spacing",
    "styles['Heading1']   — large bold heading",
    "styles['Heading2']   — medium bold heading",
    "styles['Heading3']   — smaller bold heading",
    "styles['Code']       — monospace font",
])

h2("3.2 Creating Custom Styles with ParagraphStyle")

code_block([
    "header_style = ParagraphStyle(",
    "    'Header',                          # unique name (required)",
    "    parent=styles['Heading1'],         # inherit all Heading1 defaults",
    "    fontSize=20,                       # override font size",
    "    textColor=colors.HexColor('#2c3e50'),  # override color",
    "    spaceAfter=6,                      # 6 points of space below",
    ")",
    "",
    "body_style = ParagraphStyle(",
    "    'Body',",
    "    parent=styles['BodyText'],",
    "    fontSize=10,",
    "    leading=14,    # line height: 14pt (must be > fontSize to avoid overlap)",
    "    spaceAfter=6,",
    ")",
])

add_table(
    ["ParagraphStyle attribute", "What it controls", "Example value"],
    [
        ["fontSize",    "Font size in points",               "10, 14, 20"],
        ["textColor",   "Text color",                        "colors.HexColor('#2c3e50')  or  colors.grey"],
        ["leading",     "Line height (space between lines)",  "14  (should be ~1.4× fontSize)"],
        ["spaceAfter",  "Space below the paragraph in points","6, 8, 20"],
        ["spaceBefore", "Space above the paragraph",         "12, 16"],
        ["parent",      "Inherit all settings from this style","styles['BodyText']"],
    ],
    col_widths=[2.2, 2.4, 2.0],
)

body(
    "leading must be larger than fontSize or lines will overlap. "
    "A typical ratio is leading = fontSize × 1.4. "
    "For fontSize=10, leading=14 gives comfortable line spacing."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — TABLESTYLE COMMANDS
# ══════════════════════════════════════════════════════════════════════════════

h1("4. TableStyle — Every Command Explained")

body(
    "TableStyle defines the visual appearance of a Table. "
    "It takes a list of commands, each as a tuple of (command_name, start_cell, end_cell, value). "
    "Cell coordinates are (col, row), starting from (0, 0) at the top-left. "
    "(-1, -1) means the bottom-right cell (the entire table)."
)

h2("4.1 Cell Coordinate System")

code_block([
    "# A 3-row, 3-column table:",
    "#",
    "#   (0,0)  (1,0)  (2,0)   ← header row",
    "#   (0,1)  (1,1)  (2,1)   ← first data row",
    "#   (0,2)  (1,2)  (2,2)   ← last data row",
    "#",
    "# Shortcuts:",
    "#   (-1, 0)   = last column, first row",
    "#   (0, -1)   = first column, last row",
    "#   (-1, -1)  = last column, last row",
    "#",
    "# Range (0,0) to (-1,0) = entire header row",
    "# Range (0,0) to (-1,-1) = entire table",
])

h2("4.2 Every Command Used in pdf_generator.py")

add_table(
    ["Command", "Parameters after cell range", "What it does"],
    [
        ["BACKGROUND",      "color",                  "Fill cell background with solid color"],
        ["TEXTCOLOR",       "color",                  "Set text color for cells in range"],
        ["FONTNAME",        "'Helvetica-Bold'",        "Font family name (must be a PDF standard font)"],
        ["FONTSIZE",        "point_size",              "Font size in points"],
        ["ROWBACKGROUNDS",  "[color1, color2]",        "Alternate row fill colors (repeating pattern)"],
        ["GRID",            "line_width, color",       "Draw border lines on all cell edges"],
        ["TOPPADDING",      "points",                  "Space between cell top edge and text"],
        ["BOTTOMPADDING",   "points",                  "Space between text and cell bottom edge"],
        ["LEFTPADDING",     "points",                  "Space between left edge and text"],
    ],
    col_widths=[1.9, 1.9, 2.8],
)

h2("4.3 The Header Row Pattern")

code_block([
    "kpi_table.setStyle(TableStyle([",
    "    # Dark header row:",
    "    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),",
    "    ('TEXTCOLOR',  (0, 0), (-1, 0), colors.white),",
    "    ('FONTNAME',   (0, 0), (-1, 0), 'Helvetica-Bold'),",
    "    ('FONTSIZE',   (0, 0), (-1, 0), 11),",
    "",
    "    # Alternating data rows:",
    "    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),",
    "    ('FONTSIZE',       (0, 1), (-1, -1), 10),",
    "",
    "    # Grid lines on entire table:",
    "    ('GRID',           (0, 0), (-1, -1), 0.5, colors.HexColor('#dee2e6')),",
    "",
    "    # Padding on all cells:",
    "    ('TOPPADDING',    (0, 0), (-1, -1), 6),",
    "    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),",
    "    ('LEFTPADDING',   (0, 0), (-1, -1), 10),",
    "]))",
])

body(
    "ROWBACKGROUNDS takes a list of two colors and alternates them across data rows. "
    "Row 1 (first data row, index 1) gets colors.white. "
    "Row 2 gets HexColor('#f8f9fa') (very light gray). "
    "Row 3 gets white again. This gives the 'zebra stripe' pattern common in data tables."
)

h2("4.4 repeatRows=1 — Headers That Repeat Across Pages")

code_block([
    "risk_table = Table(risk_rows, colWidths=col_widths, repeatRows=1)",
    "",
    "# repeatRows=1 means: repeat the first 1 row on every new page",
    "# If the risk table spans 2 pages, the header row prints at the top of page 2 also",
    "# Without repeatRows=1, the second page would show data with no column headers",
])

body(
    "The risk analysis table has 15 project rows and can span a page boundary. "
    "repeatRows=1 ensures the column headers (Project, Dept, Budget Var %, etc.) "
    "appear at the top of every continuation page."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — _parse_ai_sections()
# ══════════════════════════════════════════════════════════════════════════════

h1("5. Parsing the AI Report — _parse_ai_sections()")

body(
    "The AI report is a single string with five sections separated by ## headers. "
    "The PDF generator needs to place each section in a different part of the document. "
    "_parse_ai_sections() splits the string into a dictionary keyed by section name."
)

h2("5.1 The Input and Expected Output")

code_block([
    "# Input (from generate_report()):",
    "ai_report = '''",
    "## Executive Summary",
    "The portfolio is under pressure with 40% of projects delayed...",
    "",
    "## Key Risks",
    "Mobile App Launch carries critical risk at +32.1% over budget...",
    "",
    "## Budget Concerns",
    "Seven of 15 projects exceed the 10% variance threshold...",
    "'''",
    "",
    "# Output from _parse_ai_sections():",
    "{",
    "    'Executive Summary': 'The portfolio is under pressure with 40% of projects delayed...',",
    "    'Key Risks': 'Mobile App Launch carries critical risk at +32.1% over budget...',",
    "    'Budget Concerns': 'Seven of 15 projects exceed the 10% variance threshold...',",
    "}",
])

h2("5.2 The Parser — Line by Line")

code_block([
    "def _parse_ai_sections(ai_report: str) -> dict[str, str]:",
    "    sections = {}",
    "    current_key = None",
    "    current_lines = []",
    "",
    "    for line in ai_report.split('\\n'):         # iterate every line",
    "        if line.startswith('## '):              # new section header found",
    "            if current_key:                    # save the previous section",
    "                sections[current_key] = '\\n'.join(current_lines).strip()",
    "            current_key = line[3:].strip()     # extract header text (skip '## ')",
    "            current_lines = []                 # reset accumulator",
    "        else:",
    "            current_lines.append(line)         # accumulate body lines",
    "",
    "    if current_key:                            # save the last section",
    "        sections[current_key] = '\\n'.join(current_lines).strip()",
    "",
    "    return sections",
])

body(
    "The parser uses a state machine pattern. current_key tracks which section we are inside. "
    "current_lines accumulates the body lines of the current section. "
    "When a new ## header is found, the accumulated lines are saved and the accumulators reset."
)

callout(
    "Why not use split('## ')?",
    "str.split('## ') would work for a simple case, but it strips the delimiter "
    "and makes it harder to reconstruct the section names. "
    "The line-by-line state machine is more explicit and easier to debug: "
    "you can add a print() anywhere to see exactly what the parser is seeing.",
    bg="E3F2FD",
)

body(
    "The '_' prefix in _parse_ai_sections and _add_page_number signals that these functions "
    "are private helpers — they exist only to support generate_pdf() and should not be "
    "called from outside the module. Python does not enforce this, but the underscore "
    "is a widely recognized convention."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — CANVAS CALLBACKS AND PAGE NUMBERS
# ══════════════════════════════════════════════════════════════════════════════

h1("6. Canvas Callbacks — Adding Page Numbers")

body(
    "The flowable story system handles content. For running headers and footers "
    "(things that appear on every page), ReportLab uses canvas callbacks. "
    "These are functions called directly after each page is rendered."
)

h2("6.1 The Callback Function")

code_block([
    "def _add_page_number(canvas, doc):",
    "    canvas.saveState()",
    "    canvas.setFont('Helvetica', 9)",
    "    canvas.setFillColor(colors.grey)",
    "    canvas.drawRightString(",
    "        letter[0] - 0.5 * inch,   # x position: right margin",
    "        0.4 * inch,               # y position: 0.4 inch from bottom",
    "        f'Page {doc.page}'        # text to draw",
    "    )",
    "    canvas.restoreState()",
])

body(
    "canvas is ReportLab's low-level drawing surface. Unlike flowables, "
    "canvas operations use absolute coordinates measured from the bottom-left corner of the page. "
    "saveState() and restoreState() are like a stack: saveState() preserves the current "
    "drawing settings (font, color, transform), and restoreState() restores them. "
    "This ensures the callback does not accidentally affect the main content rendering."
)

code_block([
    "# PDF coordinate system (bottom-left origin):",
    "#",
    "#  (0, 11*inch)  ──────────  (8.5*inch, 11*inch)   ← top of page",
    "#       |                              |",
    "#       |          content             |",
    "#       |                              |",
    "#  (0, 0)  ────────────────  (8.5*inch, 0)          ← bottom of page",
    "#",
    "# letter[0] = 8.5*inch (page width)",
    "# letter[1] = 11*inch  (page height)",
    "#",
    "# y=0.4*inch = 0.4 inches from bottom",
    "# x = letter[0] - 0.5*inch = right-aligned with 0.5 inch margin",
])

h2("6.2 Wiring the Callback to doc.build()")

code_block([
    "doc.build(",
    "    story,",
    "    onFirstPage=_add_page_number,   # called after first page renders",
    "    onLaterPages=_add_page_number,  # called after every subsequent page",
    ")",
])

body(
    "onFirstPage and onLaterPages let you use different callbacks for the first page "
    "(e.g. no page number on a cover page) versus subsequent pages. "
    "In our case we pass the same function to both — page numbers appear on every page."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — THE THREE-PAGE PDF STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════

h1("7. The Three-Page PDF Structure")

body("The generated PDF has three pages, separated by explicit PageBreak() calls:")

add_table(
    ["Page", "Content", "ReportLab flowables"],
    [
        ["Page 1", "Report title, company, date, Executive Summary text, KPI summary table",
         "Paragraph × 3, Spacer × 2, Paragraph (exec summary), Table (KPI)"],
        ["Page 2", "Risk Analysis — full project table with all 15 rows",
         "PageBreak, Paragraph (heading), Table (risk rows, repeatRows=1)"],
        ["Page 3", "AI Recommendations — Key Risks, Budget Concerns, Schedule Delays, Actions",
         "PageBreak, Paragraph × 4 headings, Paragraph × multiple body"],
    ],
    col_widths=[0.6, 2.8, 3.2],
)

h2("7.1 Multi-Paragraph Text Handling")

code_block([
    "# The AI report section may contain multiple paragraphs separated by blank lines:",
    "summary_text = ai_sections.get('Executive Summary', 'Not available.')",
    "",
    "for para in summary_text.split('\\n\\n'):   # split on double newline",
    "    if para.strip():",
    "        story.append(Paragraph(para.strip(), body_style))",
])

body(
    "The AI model sometimes outputs paragraphs separated by blank lines (double newlines). "
    "Splitting on '\\n\\n' and creating a separate Paragraph for each gives correct "
    "paragraph spacing. A single Paragraph with embedded newlines would render as one "
    "block with less whitespace."
)

h2("7.2 Rendering Bullet Points from AI Text")

code_block([
    "# AI text like:",
    "# - Mobile App Launch: 77 days late | Owner: Lisa Wang",
    "# - Financial System Integration: 76 days late | Owner: David Brown",
    "",
    "# Is rendered with .replace('\\n', '<br/>'):",
    "story.append(Paragraph(para.strip().replace('\\n', '<br/>'), body_style))",
])

body(
    "ReportLab's Paragraph supports a small set of HTML-like tags. "
    "<br/> is a line break within a paragraph. "
    "The replace() converts bare newlines into <br/> tags so each bullet point "
    "appears on its own line while remaining in the same Paragraph flow."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — demo_reset.py
# ══════════════════════════════════════════════════════════════════════════════

h1("8. File: demo/demo_reset.py — Demo Reliability")

body(
    "A consulting demo is only as good as its reliability. "
    "If the data is stale, the database is corrupted, or the pipeline was run with different data, "
    "you cannot trust what the client sees. demo_reset.py is your guarantee."
)

h2("8.1 What It Does")

code_block([
    "python demo/demo_reset.py",
    "",
    "Resetting demo environment...",
    "  Database cleared",
    "  Sample data written → data/input/project_data.csv",
    "  Pipeline complete — 15 projects loaded",
    "",
    "Reset complete in 0.9s",
    "",
    "Demo is ready. Run:",
    "  streamlit run app/main.py",
])

body("Three operations in sequence:")
bullet("1. DROP TABLE IF EXISTS — deletes all three DuckDB tables")
bullet("2. write_sample_csv() — regenerates the CSV from the SAMPLE_PROJECTS constant")
bullet("3. Run data pipeline stages 1–4 (no LLM) — repopulates all three tables")

h2("8.2 Why DROP TABLE, Not DELETE?")

code_block([
    "# DROP TABLE removes the table entirely:",
    "con.execute('DROP TABLE IF EXISTS gold_project_kpi')",
    "",
    "# DELETE FROM removes rows but keeps the table:",
    "# con.execute('DELETE FROM gold_project_kpi')  ← not used",
    "",
    "# Why DROP is better here:",
    "# - If the table schema changed (new column added), DROP + recreate picks it up",
    "# - DELETE leaves an empty table that could confuse the pipeline",
    "# - DROP IF EXISTS is safe even if the table was never created (cold start)",
])

body(
    "IF EXISTS prevents an error if the table does not exist yet. "
    "This is what makes the reset work on a fresh machine with no database at all — "
    "the cold start case."
)

h2("8.3 sys.path.insert for Running as a Script")

code_block([
    "# demo/demo_reset.py needs to import from src/ and demo/",
    "# When run as 'python demo/demo_reset.py', Python adds 'demo/' to sys.path",
    "# But 'src' and 'demo' packages are one level up (in the project root)",
    "",
    "import sys",
    "import os",
    "sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))",
    "",
    "# __file__ = '/path/to/project/demo/demo_reset.py'",
    "# os.path.dirname(__file__) = '/path/to/project/demo'",
    "# os.path.dirname(os.path.dirname(__file__)) = '/path/to/project'  ← project root",
    "",
    "# Now 'from src.storage.db import get_connection' finds the right package",
])

body(
    "This is a more precise version of conftest.py's sys.path fix. "
    "conftest.py uses os.path.dirname(__file__) (one level up from conftest.py, which is at the root). "
    "demo_reset.py uses os.path.dirname(os.path.dirname(__file__)) (two levels up: demo/ → root)."
)

h2("8.4 Why Skip the LLM in Demo Reset?")

code_block([
    "# demo_reset.py runs stages 1–4 only (data pipeline):",
    "load_csv_to_bronze(csv_path)",
    "bronze_to_silver()",
    "silver_to_gold()",
    "summary = compute_kpi_summary()",
    "",
    "# It does NOT call:",
    "# build_prompt(summary)",
    "# generate_report(prompt)",
])

body(
    "The LLM call costs money and takes 3–8 seconds. The reset is meant to be run "
    "before every client meeting — potentially several times. Skipping the LLM keeps "
    "the reset fast (under 1 second) and free. The client generates the AI report live "
    "by clicking 'Generate Report' in the UI — that is part of the demo."
)

callout(
    "Demo reliability principle:",
    "The demo reset gives you a known good state before every meeting. "
    "If something went wrong during testing, if you accidentally loaded different data, "
    "if the database got corrupted — one command brings you back to a clean, verified state. "
    "Never walk into a client meeting without running it.",
    bg="E8F5E9",
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — ENVIRONMENT CUSTOMIZATION
# ══════════════════════════════════════════════════════════════════════════════

h1("9. Customizing the PDF for Different Clients")

body(
    "The PDF generator reads COMPANY_NAME from the .env file. "
    "Before a demo with a real client, you can customize the report header "
    "by changing a single line in .env — no code change needed."
)

code_block([
    "# .env",
    "COMPANY_NAME=Acme Corporation",
    "",
    "# PDF header becomes:",
    "# AI Executive Portfolio Report",
    "# Acme Corporation · May 08, 2026",
    "",
    "# Change for next client:",
    "COMPANY_NAME=GlobalTech Industries",
    "",
    "# PDF header becomes:",
    "# AI Executive Portfolio Report",
    "# GlobalTech Industries · May 08, 2026",
])

body(
    "This is the same env-var pattern used for MODEL_ID. "
    "Client-specific configuration belongs in .env, not in code. "
    "It also means you can show a client their own name on the report during the demo — "
    "a small detail that makes the demo feel more real."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — KEY CONCEPTS
# ══════════════════════════════════════════════════════════════════════════════

h1("10. Key Concepts to Remember")

add_table(
    ["Concept", "What it means", "Why it matters"],
    [
        ["Story / flowable pattern",    "Append objects to a list; engine handles layout",    "Never specify coordinates manually"],
        ["Paragraph + ParagraphStyle",  "Text block with font/color/spacing config",          "All text in PDF goes through a style"],
        ["TableStyle commands",         "(command, start_cell, end_cell, value) tuples",      "Controls all visual aspects of tables"],
        ["ROWBACKGROUNDS",              "List of two colors alternating across data rows",     "Zebra stripes for readable tables"],
        ["repeatRows=1",               "Repeat header row on multi-page tables",              "Column headers visible on every page"],
        ["leading",                    "Line height — must exceed fontSize",                  "Overlapping text if leading < fontSize"],
        ["canvas callbacks",           "Functions called after each page renders",            "Used for running headers, footers, page numbers"],
        ["saveState / restoreState",   "Preserve/restore canvas drawing settings",           "Callbacks must not pollute main content state"],
        ["Bottom-left origin",         "PDF coordinates start at (0,0) bottom-left",         "y=0.4*inch is near bottom, not near top"],
        ["_parse_ai_sections()",       "State machine splits AI text by ## headers",         "Lets PDF place each section independently"],
        ["DROP TABLE IF EXISTS",       "Removes table safely even if it does not exist",     "Makes demo reset work on cold start"],
        ["sys.path double dirname",    "os.path.dirname(dirname(__file__)) = project root",  "Scripts in subdirs need project root on path"],
    ],
    col_widths=[2.0, 2.3, 2.3],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — VERIFICATION
# ══════════════════════════════════════════════════════════════════════════════

h1("11. Verifying Day 6")

h2("Test the PDF generator directly")

code_block([
    "PYTHONPATH=. python -c \"",
    "from dotenv import load_dotenv; load_dotenv()",
    "from src.pipeline import run_pipeline",
    "from src.reporting.pdf_generator import generate_pdf",
    "",
    "result = run_pipeline('data/input/project_data.csv')",
    "kpi = result['kpi_summary']",
    "mock_report = '''## Executive Summary",
    "Test summary.",
    "",
    "## Key Risks",
    "Test risks.",
    "",
    "## Budget Concerns",
    "Test budget.",
    "",
    "## Schedule Delays",
    "Test delays.",
    "",
    "## Recommended Actions",
    "Test actions.",
    "'''",
    "path = generate_pdf(kpi, mock_report, 'outputs/reports/test.pdf')",
    "import os",
    "print(f'PDF: {path} ({os.path.getsize(path):,} bytes)')",
    "\" 2>&1 | grep PDF",
    "",
    "# Expected: PDF: outputs/reports/test.pdf (5,000+ bytes)",
])

h2("Test the demo reset")

code_block([
    "# Simulate cold start:",
    "rm -f data/db/reporting.duckdb",
    "",
    "# Reset should complete under 30 seconds:",
    "PYTHONPATH=. python demo/demo_reset.py",
    "",
    "# Expected:",
    "# Resetting demo environment...",
    "#   Database cleared",
    "#   Sample data written → data/input/project_data.csv",
    "#   Pipeline complete — 15 projects loaded",
    "#",
    "# Reset complete in 0.9s",
])

h2("Full test suite (still 19 passing)")

code_block([
    "PYTHONPATH=. pytest tests/ -v",
    "# Expected: 19 passed",
])

h1("12. What Day 7 Builds On This")

body("Day 7 is the final day — polish, two extra tests, README, and demo rehearsal:")

bullet("Add test_pipeline_stages_called_in_order — verifies all 6 pipeline functions are called")
bullet("Add test_pipeline_survives_llm_failure — verifies pipeline completes even with fallback AI report")
bullet("Create README.md — 5-minute setup guide, demo flow, architecture overview, consulting offer")
bullet("Cold start rehearsal: rm database → demo_reset.py → streamlit run → full UI test")
bullet("Practice the 3-minute demo script out loud")

body("After Day 7, the system is production-quality: every layer tested, cold start verified, and a professional README that any developer can follow to run it from scratch.")

divider()

p_final = doc.add_paragraph()
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_final.paragraph_format.space_before = Pt(20)
run_final = p_final.add_run("Day 6 complete — PDF exporting, demo reset working in under 1 second.")
set_font(run_final, bold=True, size=12, color=(39, 174, 96))

# ── Save ───────────────────────────────────────────────────────────────────────
out_path = "docs/day-6-learning-guide.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
