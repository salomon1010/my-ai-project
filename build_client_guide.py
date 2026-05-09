"""Generates the AI Executive Reporting – Client & Cost Guide .docx"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Helpers ───────────────────────────────────────────────────────────────────
BRAND_BLUE  = RGBColor(0x1A, 0x56, 0xDB)   # deep blue
BRAND_DARK  = RGBColor(0x11, 0x18, 0x27)   # near-black
ACCENT_TEAL = RGBColor(0x0E, 0x9F, 0x6E)   # green/teal
LIGHT_GRAY  = RGBColor(0xF9, 0xFA, 0xFB)
MID_GRAY    = RGBColor(0x6B, 0x72, 0x80)

def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = BRAND_BLUE
    p.runs[0].font.size = Pt(18)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = BRAND_DARK
    p.runs[0].font.size = Pt(14)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def heading3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = ACCENT_TEAL
    p.runs[0].font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent  = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after  = Pt(3)
    run = p.runs[0]
    run.font.size = Pt(11)
    return p

def callout(text, label="NOTE"):
    """Shaded callout paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.color.rgb = BRAND_BLUE
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p

def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p

def divider():
    doc.add_paragraph("─" * 80).runs[0].font.color.rgb = MID_GRAY

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # shade header
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "1A56DB")
        cell._tc.get_or_add_tcPr().append(shading)
    # data rows
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        fill = "F3F4F6" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), fill)
            cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run("\n\nAI Executive Reporting System")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = BRAND_BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Client & Cost Transparency Guide")
r2.font.size = Pt(16)
r2.font.color.rgb = MID_GRAY

doc.add_paragraph()
tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = tagline.add_run(
    "Everything you need to know about how this system works,\n"
    "what it costs, and how to discuss it confidently with any client."
)
r3.font.size = Pt(12)
r3.italic = True
r3.font.color.rgb = BRAND_DARK

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 – WHAT WE BUILT
# ══════════════════════════════════════════════════════════════════════════════
heading1("1. What We Built — System Overview")

body(
    "The AI Executive Reporting System transforms raw project data (CSV files) into "
    "professional, AI-written executive reports automatically. Here is the complete flow:"
)

add_table(
    ["Step", "What Happens", "Technology"],
    [
        ["1 – Ingest",   "Raw CSV uploaded or loaded from disk",          "Python / DuckDB"],
        ["2 – Bronze",   "Raw data stored as-is in the database",         "DuckDB"],
        ["3 – Silver",   "Data cleaned, validated, normalised",           "pandas / DuckDB"],
        ["4 – Gold",     "KPIs computed (budget variance, risk, delays)", "pandas / DuckDB"],
        ["5 – AI Report","Claude reads KPIs and writes the summary",      "Anthropic API"],
        ["6 – Dashboard","Results shown in a web browser",                "Streamlit"],
        ["7 – PDF",      "One-click professional PDF export",             "ReportLab"],
    ],
    col_widths=[1.0, 3.2, 1.8],
)

callout(
    "The client never touches code. They open a browser, click 'Load Sample Data', "
    "then 'Generate Report'. That is the entire workflow.",
    label="CLIENT VIEW"
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 – STEP-BY-STEP SETUP WE DID
# ══════════════════════════════════════════════════════════════════════════════
heading1("2. Step-by-Step Setup (What We Did Together)")

heading2("Step 1 — Create an Anthropic Account & Buy API Credits")
bullet("Went to https://console.anthropic.com")
bullet("Created an account with the email address")
bullet("Navigated to Billing → Added a payment method")
bullet("Purchased $20 in API credits (credits, not a subscription)")
callout(
    "$20 = $20 of usable API credit. There is no monthly fee, no auto-renewal "
    "unless you set one up. You only pay for what you use.",
    label="IMPORTANT"
)

heading2("Step 2 — Generate an API Key")
bullet("In the Anthropic console: API Keys → Create Key")
bullet("Named the key (e.g., 'my-ai-project-dev')")
bullet("Copied the key immediately — it is only shown once")
bullet("Key format: sk-ant-api03-…")
callout(
    "Treat the API key like a password. Anyone who has it can spend your credits. "
    "Never share it in emails, screenshots, or commit it to git.",
    label="SECURITY"
)

heading2("Step 3 — Configure the Project")
bullet("Opened the project folder: /Users/…/my-ai-project")
bullet("Copied .env.example → .env")
bullet("Added the API key to ANTHROPIC_API_KEY in .env")
bullet("The .env file is listed in .gitignore so it is never accidentally uploaded")

heading2("Step 4 — Install Dependencies")
code_block("pip install -r requirements.txt")
body("All libraries were already present in the Anaconda environment (zero new installs needed).")

heading2("Step 5 — Load Sample Data")
code_block("python demo/demo_reset.py")
body("This seeded 15 sample projects into the local DuckDB database through the full pipeline (1.6 seconds).")

heading2("Step 6 — Launch the Dashboard")
code_block("streamlit run app/main.py --browser.gatherUsageStats false")
body("App available at http://localhost:8501. The first run required dismissing Streamlit's one-time email prompt, handled by creating ~/.streamlit/credentials.toml.")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 – HOW THE ANTHROPIC API WORKS
# ══════════════════════════════════════════════════════════════════════════════
heading1("3. How the Anthropic API Works")

body(
    "Claude is Anthropic's AI model. Instead of a chat interface, the API lets your "
    "software send instructions and data programmatically and receive a response."
)

heading2("The Core Concept: Tokens")
body(
    "Everything sent to or received from Claude is measured in tokens. "
    "One token ≈ 0.75 words (roughly 4 characters). The API charges based on "
    "how many tokens go IN (input) and how many come OUT (output)."
)

add_table(
    ["Direction", "What It Contains", "Billed As"],
    [
        ["Input  →", "Your CSV data summary + instructions to Claude", "Input tokens"],
        ["← Output", "The executive summary Claude writes",           "Output tokens"],
    ],
    col_widths=[1.2, 3.5, 1.5],
)

heading2("How a Report Is Generated (Under the Hood)")
bullet("Pipeline runs: CSV → Bronze → Silver → Gold KPIs")
bullet("prompt_builder.py formats the KPI data into a structured prompt")
bullet("claude_client.py sends the prompt to the Anthropic API")
bullet("Claude reads the prompt and writes a polished executive summary")
bullet("The summary is displayed in the browser and embedded in the PDF")

code_block(
    "# Simplified view of what happens in claude_client.py\n"
    "client = anthropic.Anthropic(api_key=os.getenv('ANTHROPIC_API_KEY'))\n"
    "message = client.messages.create(\n"
    "    model='claude-haiku-4-5',\n"
    "    max_tokens=1024,\n"
    "    messages=[{'role': 'user', 'content': prompt}]\n"
    ")\n"
    "return message.content[0].text"
)

callout(
    "The API call takes 3–10 seconds. The rest of the pipeline (data loading, KPI calculation) "
    "takes under 2 seconds. Total time from click to report: under 15 seconds.",
    label="PERFORMANCE"
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 – COST BREAKDOWN
# ══════════════════════════════════════════════════════════════════════════════
heading1("4. Complete Cost Breakdown")

heading2("Credit Structure")
bullet("You buy CREDITS, not a subscription")
bullet("$20 = $20 of API credits (no markup by Anthropic)")
bullet("Credits expire 1 year after purchase")
bullet("Pay-as-you-go: you are only charged for tokens processed")
bullet("Unused credits roll over until expiry")

heading2("Model Pricing (as of May 2026)")

add_table(
    ["Model", "Best For", "Input / 1M tokens", "Output / 1M tokens", "Relative Cost"],
    [
        ["Claude Haiku 4.5",   "Routine summaries, high volume",   "$0.80",  "$4.00",  "1× (cheapest)"],
        ["Claude Sonnet 4.6",  "Balanced quality + cost (default)","$3.00",  "$15.00", "4–5×"],
        ["Claude Opus 4.7",    "Complex analysis, highest quality","$15.00", "$75.00", "20×"],
    ],
    col_widths=[1.5, 2.2, 1.4, 1.5, 1.4],
)

heading2("Cost Per Report — This Project")

body("Scenario: 1 CSV file → 1 executive summary report")

add_table(
    ["Item", "Tokens", "Haiku 4.5", "Sonnet 4.6", "Opus 4.7"],
    [
        ["Input (data + prompt)", "~2,000", "$0.0016", "$0.006", "$0.030"],
        ["Output (summary)",      "~1,000", "$0.0040", "$0.015", "$0.075"],
        ["TOTAL per report",      "~3,000", "$0.006",  "$0.021", "$0.105"],
        ["Reports per $20",       "—",      "~3,300",  "~950",   "~190"],
    ],
    col_widths=[2.0, 1.2, 1.2, 1.2, 1.2],
)

callout(
    "For most clients, Haiku 4.5 produces excellent routine summaries at ~$0.006 each. "
    "Reserve Sonnet for reports requiring deeper narrative or nuance. "
    "Opus is rarely needed for structured data reporting.",
    label="RECOMMENDATION"
)

heading2("Scaling Estimates — What Will a Client Actually Spend?")

add_table(
    ["Usage Pattern", "Reports/Month", "Model", "Monthly Cost", "Annual Cost"],
    [
        ["Small team, weekly reports",    "4",   "Sonnet", "~$0.08",  "~$1.00"],
        ["Department, daily reports",     "22",  "Sonnet", "~$0.46",  "~$5.52"],
        ["Enterprise, hourly batch",      "720", "Haiku",  "~$4.32",  "~$51.84"],
        ["Large portfolio, complex data", "100", "Opus",   "~$10.50", "~$126.00"],
    ],
    col_widths=[2.2, 1.2, 1.0, 1.2, 1.2],
)

body(
    "In almost every realistic scenario, Anthropic API costs are negligible compared "
    "to the consulting fee or the time saved by automation."
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 – MODEL SELECTION GUIDE
# ══════════════════════════════════════════════════════════════════════════════
heading1("5. Choosing the Right Model for Each Client")

add_table(
    ["Client Profile", "Recommended Model", "Why"],
    [
        ["Startup, cost-sensitive, standard reports",
         "Haiku 4.5",
         "4× cheaper, fast, great for structured KPI data"],
        ["Mid-market, mixed complexity",
         "Sonnet 4.6",
         "Best balance of quality and cost for most use cases"],
        ["Enterprise, board-level reporting, nuanced narrative",
         "Opus 4.7",
         "Highest reasoning, best writing — cost is not a concern"],
        ["High-volume batch processing (100+ reports/day)",
         "Haiku 4.5 + caching",
         "Lowest cost; prompt caching cuts input cost by 50–90%"],
        ["R&D / experimental features",
         "Sonnet 4.6",
         "Good capability at a manageable cost while iterating"],
    ],
    col_widths=[2.0, 1.5, 3.0],
)

heading2("Prompt Caching — The Hidden Cost Reducer")
body(
    "If your system prompt (the instructions you send to Claude every time) is the same "
    "across many requests, Anthropic caches it. Cached tokens cost 90% less on input. "
    "For a project with a fixed system prompt, this alone can cut API costs in half."
)
bullet("Enable with: cache_control={'type': 'ephemeral'} in the API call")
bullet("Cache TTL: 5 minutes (refreshed on each hit)")
bullet("Savings: up to 90% on input tokens for the cached portion")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 – CLIENT CONVERSATION GUIDE
# ══════════════════════════════════════════════════════════════════════════════
heading1("6. How to Talk to Clients About This")

heading2("During the Project (Discovery & Build Phase)")

heading3("When the client asks: 'What does the AI cost?'")
body(
    "Use this framing:"
)
callout(
    '"The AI component costs between $0.006 and $0.02 per report generated — '
    'that is less than one cent to two cents. With a $20 top-up, you can generate '
    'roughly 1,000 reports. For most businesses, the API cost is essentially zero '
    'compared to the time it replaces."',
    label="SCRIPT"
)

heading3("When the client asks: 'Is this a subscription?'")
callout(
    '"No subscription. You buy credits directly from Anthropic — the company that makes '
    'Claude. You only spend credits when a report is generated. If you generate zero '
    'reports in a month, you spend zero. Credits last one year."',
    label="SCRIPT"
)

heading3("When the client asks: 'Who holds our data?'")
callout(
    '"Your data is sent to Anthropic\'s API only at the moment a report is generated, '
    'and only the KPI summary — not the raw CSV. Anthropic does not use API data to '
    'train their models (by default). Everything else stays on your infrastructure."',
    label="SCRIPT"
)

heading2("After Go-Live (Ongoing Transparency)")

heading3("Monthly cost reporting")
body("Share a simple summary each month:")
add_table(
    ["Month", "Reports Generated", "Model Used", "API Cost", "Cost per Report"],
    [
        ["May 2026",  "47",  "Haiku 4.5",   "$0.28",  "$0.006"],
        ["June 2026", "63",  "Haiku 4.5",   "$0.38",  "$0.006"],
        ["July 2026", "120", "Mixed",        "$1.10",  "$0.009"],
    ],
    col_widths=[1.2, 1.5, 1.5, 1.2, 1.5],
)
callout(
    "Proactively share this table every month. Clients appreciate transparency. "
    "When API costs are this low, showing them only builds trust.",
    label="BEST PRACTICE"
)

heading3("Budget alerts")
bullet("Set a billing alert in console.anthropic.com (e.g., notify at $10 spent)")
bullet("Share the alert threshold with the client so they feel in control")
bullet("Offer to set up a monthly credit auto-reload if they prefer predictability")

heading2("Handling the 'Can We Switch Models?' Conversation")
body(
    "Clients sometimes ask to upgrade or downgrade. Here is how to frame each direction:"
)
add_table(
    ["Request", "What It Means", "Your Response"],
    [
        ["'Make it cheaper'",
         "Switch from Sonnet → Haiku",
         "Cost drops 4×; quality drops slightly for complex data. "
         "Recommend a 2-week trial on Haiku with side-by-side comparison."],
        ["'Make it smarter'",
         "Switch from Haiku/Sonnet → Opus",
         "Cost increases 5–20×; reserve for board-level or investor reports. "
         "Most clients do not need Opus for structured KPI data."],
        ["'Can we batch everything at night?'",
         "Scheduled batch processing",
         "Yes — add a cron job. Same cost, just shifted to off-hours. "
         "Also enables prompt caching savings."],
    ],
    col_widths=[1.5, 2.0, 3.0],
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 – MONITORING & SAFETY
# ══════════════════════════════════════════════════════════════════════════════
heading1("7. Monitoring, Safety & Best Practices")

heading2("Tracking Usage")
bullet("Dashboard: https://console.anthropic.com/settings/usage")
bullet("Shows tokens used, cost by day, cost by model")
bullet("Set email alerts at 50% and 90% of your credit balance")
bullet("Export usage CSV monthly for client billing reconciliation")

heading2("Rate Limits (with $20 credits)")
add_table(
    ["Limit Type", "Value", "Impact on This Project"],
    [
        ["Tokens per minute",   "50,000 TPM",  "~16 reports/minute — far above typical need"],
        ["Requests per minute", "50 RPM",      "50 concurrent report requests — more than enough"],
        ["Tokens per day",      "1,000,000 TPD","~333 reports/day on Sonnet; ~5,000/day on Haiku"],
    ],
    col_widths=[1.8, 1.5, 3.2],
)

heading2("Security Checklist")
bullet(".env file is in .gitignore — API key never enters version control")
bullet("Use separate API keys for development and production")
bullet("Rotate the API key immediately if it is ever exposed (console → API Keys → Revoke)")
bullet("Do not log the raw prompt in production if it contains sensitive client data")
bullet("Consider storing the API key in a secrets manager (AWS Secrets Manager, 1Password) for production")

heading2("Cost Control Techniques")
add_table(
    ["Technique", "Savings", "How to Implement"],
    [
        ["Use Haiku for routine reports",          "75% vs Sonnet",  "Set model='claude-haiku-4-5' for standard reports"],
        ["Prompt caching",                          "50–90% on input","Add cache_control to system prompt in API call"],
        ["Summarise CSV before sending",            "50–80% tokens",  "Send KPI summary, not raw rows, to Claude"],
        ["Set max_tokens limit",                    "Prevents spikes","max_tokens=1024 caps output length and cost"],
        ["Batch off-peak with cron",                "Better cache hits","Schedule batch jobs; warm cache reused across batch"],
    ],
    col_widths=[2.2, 1.3, 3.0],
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 – CONSULTING OFFER SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
heading1("8. The Consulting Offer at a Glance")

add_table(
    ["Item", "Details"],
    [
        ["Service",       "AI Reporting Automation Pilot"],
        ["Price range",   "$5,000 – $10,000 (one-time)"],
        ["Timeline",      "30 days"],
        ["Deliverables",  "Data pipeline, KPI model, AI report, dashboard, executive demo"],
        ["API cost to client", "~$20–$50/year for typical usage (paid directly to Anthropic)"],
        ["Ongoing support",    "Optional retainer for model updates, new KPI fields, UI changes"],
    ],
    col_widths=[2.0, 4.5],
)

callout(
    "The system pays for itself the first month it replaces a manual reporting process. "
    "A weekly executive report that takes 3 hours to prepare manually = $300–$600/month "
    "in staff time. This system generates the same report in 15 seconds.",
    label="ROI FRAMING"
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 – QUICK REFERENCE CARD
# ══════════════════════════════════════════════════════════════════════════════
heading1("9. Quick Reference Card")

heading2("Commands")
add_table(
    ["Task", "Command"],
    [
        ["Reset demo data",     "python demo/demo_reset.py"],
        ["Start dashboard",     "streamlit run app/main.py --browser.gatherUsageStats false"],
        ["Run all tests",       "pytest tests/ -v"],
        ["Open dashboard",      "http://localhost:8501"],
        ["Check API usage",     "https://console.anthropic.com/settings/usage"],
    ],
    col_widths=[2.5, 4.0],
)

heading2("Environment Variables (.env)")
add_table(
    ["Variable", "Example Value", "Purpose"],
    [
        ["ANTHROPIC_API_KEY",   "sk-ant-api03-…",       "Your Anthropic API key"],
        ["MODEL_ID",            "claude-haiku-4-5",     "Model used for report generation"],
        ["DB_PATH",             "data/db/reporting.duckdb", "Local database path"],
        ["REPORTS_OUTPUT_DIR",  "outputs/reports",      "Where PDFs are saved"],
        ["COMPANY_NAME",        "Acme Corporation",     "Appears in report header"],
    ],
    col_widths=[1.8, 2.0, 2.7],
)

heading2("Model IDs for .env")
add_table(
    ["Model", "MODEL_ID value", "Use When"],
    [
        ["Haiku 4.5",  "claude-haiku-4-5",            "Cost-sensitive / high volume"],
        ["Sonnet 4.6", "claude-sonnet-4-6",            "Default — best balance"],
        ["Opus 4.7",   "claude-opus-4-7",              "Premium quality needed"],
    ],
    col_widths=[1.5, 2.5, 2.5],
)

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run("AI Executive Reporting System — Client & Cost Guide\n")
fr.font.size = Pt(10)
fr.font.color.rgb = MID_GRAY
fr2 = footer_p.add_run("Built with Python · Streamlit · DuckDB · Anthropic Claude API")
fr2.font.size = Pt(9)
fr2.font.color.rgb = MID_GRAY

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = "outputs/AI_Executive_Reporting_Client_Cost_Guide.docx"
os.makedirs("outputs", exist_ok=True)
doc.save(out_path)
print(f"Saved → {out_path}")
